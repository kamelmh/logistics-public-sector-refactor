from docx import Document
import re

doc = Document(r'C:\Users\Admin\Downloads\Wordvice_EDITOR SAMPLE TEST_20200526.docx')

# Track all replacements
replacements = [
    # Technology Passage Title
    ('People ased on', 'People Based on'),
    
    # Abstract - Technology
    ('social related to aging', 'social issues related to aging'),
    ('The governments want to reduce', 'Governments want to reduce'),
    ('health conditions  reports', 'health condition reports'),
    ('self-reports are often', 'Self-reports are often'),
    ('To fix the problems', 'To address these problems'),
    ('In this paper, we a data-driven', 'In this paper, we propose a data-driven'),
    ('web technologies for connecting', 'web technologies to connect'),
    ('proposed ystem', 'proposed system'),
    ('provides a method elderly', 'provides a method for monitoring elderly'),
    ('and to register recognizable', 'and registers recognizable'),
    ('controls actuators in the home', 'It controls actuators in the home'),
    ('summary of activities to .', 'summary of activities.'),
    
    # Introduction - Technology
    ('in life expectancy the proportion', 'With increases in life expectancy, the proportion'),
    ('rapidly increased [1]', 'have rapidly increased [1]'),
    ('population globally and the social cost', 'population globally, and the social costs'),
    ('related to aging is difficult', 'related to aging make it difficult'),
    ('elderly people are exposed', 'Elderly people are exposed'),
    ('elderly people need the help', 'As elderly people need the help'),
    ('are used such as', 'are used, such as'),
    ('is a way  people', 'is a way of assessing people'),
    ('tend do every day', 'tend to do every day'),
    ('IADL) not necessary', 'IADLs) are not necessary'),
    ('and an independent lifestyle', 'but are necessary for an independent lifestyle'),
    ('reported  these methods', 'reported through these methods'),
    ('can be more .', 'can be more accurate.'),
    ('solve these problems', 'To solve these problems'),
    ('Internet of (IoT)', 'Internet of Things (IoT)'),
    ('with the advent of', 'With the advent of'),
    
    # Biology Passage
    ('HCCtumors  possessing', 'HCC tumors possessing'),
    ('promoter may directed', 'promoter may be directed'),
    ('promoter dna', 'promoter DNA'),
    ('to their promoters.(AGCG?)', 'to their promoters.'),
    ('associated with an decrease', 'associated with a decrease'),
    ('coupled to another SNPs', 'coupled to other SNPs'),
    
    # Electrical Engineering
    ('at the reciever side', 'at the receiver side'),
    ('has received great attentions', 'has received great attention'),
    ('cannot be avoid', 'cannot avoid'),
    ('we propose the simplified', 'we propose a simplified'),
    ('that is the same compared', 'that is comparable to'),
    
    # Materials Engineering
    ('To make modelling', 'To create a model'),
    ('thicknessin Fig', 'thickness (Fig'),
    ('connected on each', 'connected at each'),
    ('bind together at the topend', 'bound together at the top end'),
    
    # Economics/FTPL
    ('why we are worried about this attempted', 'why we are concerned about this attempted'),
    
    # Humanities
    ('Country music men often were negotiating', 'Country music men often negotiated'),
]

count = 0
for para in doc.paragraphs:
    for run in para.runs:
        for old, new in replacements:
            if old in run.text:
                run.text = run.text.replace(old, new)
                count += 1

# Also check tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    for old, new in replacements:
                        if old in run.text:
                            run.text = run.text.replace(old, new)
                            count += 1

out = r'C:\Users\Admin\Downloads\Wordvice_SAMPLE_TEST_FIXED.docx'
doc.save(out)
print(f'Done! {count} corrections applied.')
print(f'Saved: {out}')

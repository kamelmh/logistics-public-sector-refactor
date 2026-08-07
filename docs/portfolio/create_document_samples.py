from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_editing_sample():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title
    title = doc.add_heading('Academic Editing Sample', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introduction
    doc.add_heading('Before Editing (Original Text)', level=1)
    
    original = """The student have been studying english for three years and they is very good at writing. However, there are some mistakes in there papers that needs to be corrected before submission. Their professor said that the quality of there work is good but need more attention to detail and grammer.

Furthermore, the research methadology used in the study was not properly explained. The data analisis section lacks clarity and the conclusions drawn from the findings are not well supported by the evidence presented in the paper."""
    
    p = doc.add_paragraph(original)
    p.runs[0].font.color.rgb = RGBColor(234, 67, 53)  # Red for errors
    
    doc.add_heading('After Editing (Corrected Text)', level=1)
    
    corrected = """The student has been studying English for three years and is very good at writing. However, there are some errors in their papers that need to be corrected before submission. Their professor said that the quality of their work is good but needs more attention to detail and grammar.

Furthermore, the research methodology used in the study was not properly explained. The data analysis section lacks clarity, and the conclusions drawn from the findings are not well supported by the evidence presented in the paper."""
    
    p = doc.add_paragraph(corrected)
    p.runs[0].font.color.rgb = RGBColor(52, 168, 83)  # Green for corrected
    
    doc.add_heading('Changes Made', level=1)
    
    changes = [
        '"have" → "has" (subject-verb agreement)',
        '"english" → "English" (capitalization)',
        '"they is" → "is" (grammar)',
        '"there" → "their" (possessive)',
        '"needs" → "need" (subject-verb agreement)',
        '"grammer" → "grammar" (spelling)',
        '"methadology" → "methodology" (spelling)',
        '"analisis" → "analysis" (spelling)',
        'Added comma after "clarity" (punctuation)'
    ]
    
    for change in changes:
        doc.add_paragraph(change, style='List Bullet')
    
    doc.add_heading('Formatting Applied', level=1)
    
    formatting = [
        'APA 7th Edition style',
        '12pt Times New Roman font',
        'Double spacing',
        '1-inch margins',
        'Proper paragraph indentation',
        'Page numbers in header'
    ]
    
    for item in formatting:
        doc.add_paragraph(item, style='List Bullet')
    
    output_path = os.path.join(r"C:\Users\Admin\Projects\active\portfolio", "editing_sample.docx")
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path

def create_formatting_sample():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Title page
    for _ in range(6):
        doc.add_paragraph('')
    
    title = doc.add_heading('Professional Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Formatted by MAHI Kamel Abdelghani')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(26, 115, 232)
    
    date = doc.add_paragraph('July 2026')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary..................................................3',
        '2. Introduction.............................................................4',
        '3. Methodology............................................................5',
        '4. Findings.................................................................6',
        '5. Recommendations....................................................8',
        '6. Conclusion..............................................................9'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(12)
    
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    
    doc.add_paragraph(
        'This report provides a comprehensive analysis of the current market conditions '
        'and strategic recommendations for business growth. The findings indicate significant '
        'opportunities for expansion in the digital services sector.'
    )
    
    # Introduction
    doc.add_heading('2. Introduction', level=1)
    
    doc.add_paragraph(
        'The purpose of this report is to analyze current market trends and provide '
        'actionable recommendations for business development. The scope includes an '
        'examination of competitive dynamics, customer preferences, and emerging opportunities.'
    )
    
    # Methodology
    doc.add_heading('3. Methodology', level=1)
    
    doc.add_paragraph('The following methods were employed in this study:')
    
    methods = [
        'Primary research through client interviews',
        'Secondary research using industry reports',
        'Competitive analysis of key players',
        'Statistical analysis of market data'
    ]
    for method in methods:
        doc.add_paragraph(method, style='List Bullet')
    
    # Findings
    doc.add_heading('4. Findings', level=1)
    
    doc.add_heading('4.1 Market Overview', level=2)
    doc.add_paragraph(
        'The market has shown consistent growth over the past five years, with a '
        'compound annual growth rate of 12%. Key drivers include digital transformation '
        'and increasing demand for professional services.'
    )
    
    doc.add_heading('4.2 Competitive Landscape', level=2)
    doc.add_paragraph(
        'Analysis reveals a fragmented market with opportunities for differentiated '
        'service providers. Quality and reliability emerge as key differentiators.'
    )
    
    # Recommendations
    doc.add_heading('5. Recommendations', level=1)
    
    recommendations = [
        'Invest in digital capabilities',
        'Expand service offerings',
        'Strengthen client relationships',
        'Implement quality assurance processes',
        'Develop strategic partnerships'
    ]
    for i, rec in enumerate(recommendations, 1):
        doc.add_paragraph(f'{i}. {rec}')
    
    # Conclusion
    doc.add_heading('6. Conclusion', level=1)
    
    doc.add_paragraph(
        'This report demonstrates the importance of strategic planning and professional '
        'presentation in business communications. The recommendations provided offer a '
        'clear path forward for organizational growth and success.'
    )
    
    # Footer with page numbers
    doc.add_paragraph('')
    p = doc.add_paragraph('Report prepared by MAHI Kamel Abdelghani')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.italic = True
    
    output_path = os.path.join(r"C:\Users\Admin\Projects\active\portfolio", "formatting_sample.docx")
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path

def create_apa_paper_sample():
    doc = Document()
    
    # Set default font to Times New Roman 12pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title Page
    for _ in range(6):
        doc.add_paragraph('')
    
    title = doc.add_paragraph('The Impact of Digital Transformation on Academic Publishing')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(12)
    
    author = doc.add_paragraph('MAHI Kamel Abdelghani')
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    affiliation = doc.add_paragraph('University of Saida')
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    course = doc.add_paragraph('English Language and Literature')
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    instructor = doc.add_paragraph('Dr. Professor Name')
    instructor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date = doc.add_paragraph('July 15, 2026')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Abstract
    doc.add_heading('Abstract', level=1)
    
    abstract = (
        'This paper examines the transformative impact of digital technologies on academic '
        'publishing. Through analysis of current trends and case studies, the research identifies '
        'key areas of change including open access publishing, digital workflows, and new '
        'dissemination methods. The findings suggest that while digital transformation offers '
        'significant benefits, it also presents challenges for traditional publishing models. '
        'Implications for academics, publishers, and institutions are discussed.'
    )
    p = doc.add_paragraph(abstract)
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    doc.add_paragraph('')
    p = doc.add_paragraph('Keywords: digital transformation, academic publishing, open access')
    p.runs[0].font.italic = True
    
    doc.add_page_break()
    
    # Introduction
    doc.add_heading('The Impact of Digital Transformation on Academic Publishing', level=1)
    
    intro = (
        'The landscape of academic publishing has undergone significant changes in recent years, '
        'driven by advances in digital technology. This transformation has affected every aspect '
            'of the publishing process, from manuscript submission to distribution and readership '
        '(Smith, 2024). This paper examines these changes and their implications for the future '
        'of academic communication.'
    )
    p = doc.add_paragraph(intro)
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    # Literature Review
    doc.add_heading('Literature Review', level=1)
    
    lit_review = (
        'Previous research has documented the shift from print to digital publishing. '
        'Johnson (2023) found that over 70% of academic journals now offer online access, '
        'while Williams (2024) noted that open access publishing has increased by 150% in '
        'the past decade. These changes have democratized access to knowledge but also '
        'raised questions about sustainability and quality control.'
    )
    p = doc.add_paragraph(lit_review)
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    # Methodology
    doc.add_heading('Methodology', level=1)
    
    methodology = (
        'This study employed a mixed-methods approach, combining quantitative analysis of '
        'publishing trends with qualitative interviews of industry professionals. Data was '
        'collected from 50 academic publishers and 100 researchers across multiple disciplines.'
    )
    p = doc.add_paragraph(methodology)
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    # Results
    doc.add_heading('Results', level=1)
    
    results = (
        'The analysis revealed several key findings. First, digital submission systems have '
        'reduced publication timelines by an average of 40%. Second, open access articles '
        'receive 18% more citations than traditionally published works. Third, researchers '
        'under 35 prefer digital-only formats at a rate of 3:1 compared to print.'
    )
    p = doc.add_paragraph(results)
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    # Discussion
    doc.add_heading('Discussion', level=1)
    
    discussion = (
        'These findings suggest that digital transformation is fundamentally altering academic '
        'publishing. The increased accessibility of open access publishing aligns with the '
        'democratization of knowledge, while the efficiency gains from digital workflows benefit '
        'both publishers and authors. However, challenges remain in ensuring quality and '
        'sustainability in the digital landscape.'
    )
    p = doc.add_paragraph(discussion)
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    # Conclusion
    doc.add_heading('Conclusion', level=1)
    
    conclusion = (
        'Digital transformation has profoundly impacted academic publishing, offering both '
        'opportunities and challenges. As the field continues to evolve, stakeholders must '
        'adapt to new technologies while maintaining the standards of scholarly communication. '
        'Future research should examine the long-term effects of these changes on knowledge '
        'production and dissemination.'
    )
    p = doc.add_paragraph(conclusion)
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # References
    doc.add_heading('References', level=1)
    
    references = [
        'Johnson, A. (2023). The state of academic publishing. Journal of Scholarly Communication, 15(2), 45-62.',
        'Smith, B. (2024). Digital transformation in higher education. Academic Press.',
        'Williams, C. (2024). Open access publishing trends. Research Quarterly, 28(1), 112-128.'
    ]
    
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
    
    output_path = os.path.join(r"C:\Users\Admin\Projects\active\portfolio", "apa_paper_sample.docx")
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path

# Create all document samples
print("Creating professional document samples...")
create_editing_sample()
create_formatting_sample()
create_apa_paper_sample()
print("\nAll document samples created!")

"""
Thesis Table Style Comparator
Compares current DOCX with backup v7c to check table styles/borders.
"""
import sys
from docx import Document
from collections import defaultdict

def get_table_info(doc):
    """Extract table style info from a DOCX."""
    tables = []
    for i, table in enumerate(doc.tables):
        info = {
            "index": i,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "style": str(table.style) if table.style else "None",
            "borders": [],
            "shading": [],
        }
        # Check first cell for border/shading info
        if table.rows and table.rows[0].cells:
            cell = table.rows[0].cells[0]
            tc = cell._tc
            tcPr = tc.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr")
            if tcPr is not None:
                borders = tcPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders")
                if borders is not None:
                    for child in borders:
                        tag = child.tag.split("}")[-1]
                        attrs = dict(child.attrib)
                        info["borders"].append(f"{tag}:{attrs}")
                shading = tcPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd")
                if shading is not None:
                    info["shading"].append(dict(shading.attrib))
        tables.append(info)
    return tables

def compare_tables(current_file, backup_file):
    """Compare table structures between two DOCX files."""
    print(f"\n  Comparing table styles:")
    print(f"  Current: {current_file}")
    print(f"  Backup:  {backup_file}")
    
    try:
        doc_curr = Document(current_file)
        doc_back = Document(backup_file)
    except Exception as e:
        print(f"\n  [ERROR] Cannot open DOCX: {e}")
        print(f"  Make sure python-docx is installed: pip install python-docx")
        return False
    
    tables_curr = get_table_info(doc_curr)
    tables_back = get_table_info(doc_back)
    
    print(f"\n  Table count: Current={len(tables_curr)}, Backup={len(tables_back)}")
    
    issues = []
    
    # Compare table count
    if len(tables_curr) != len(tables_back):
        issues.append(f"Table count mismatch: {len(tables_curr)} vs {len(tables_back)}")
    
    # Compare each table
    for i in range(min(len(tables_curr), len(tables_back))):
        tc = tables_curr[i]
        tb = tables_back[i]
        
        # Style
        if tc["style"] != tb["style"]:
            issues.append(f"Table {i}: Style mismatch '{tc['style']}' vs '{tb['style']}'")
        
        # Dimensions
        if tc["rows"] != tb["rows"] or tc["cols"] != tb["cols"]:
            issues.append(f"Table {i}: Dimensions {tc['rows']}x{tc['cols']} vs {tb['rows']}x{tb['cols']}")
        
        # Borders
        if tc["borders"] != tb["borders"]:
            issues.append(f"Table {i}: Border styles differ")
            if tc["borders"]:
                print(f"    Table {i} current borders: {tc['borders'][:2]}")
            if tb["borders"]:
                print(f"    Table {i} backup borders:  {tb['borders'][:2]}")
        
        # Shading
        if tc["shading"] != tb["shading"]:
            issues.append(f"Table {i}: Shading differs")
    
    # Print results
    print(f"\n  === RESULTS ===")
    if not issues:
        print(f"  [PASS] All {len(tables_curr)} tables match backup v7c")
        return True
    else:
        print(f"  [WARN] {len(issues)} issue(s) found:")
        for issue in issues:
            print(f"    - {issue}")
        return False

if __name__ == "__main__":
    from pathlib import Path
    base = str(Path(__file__).resolve().parent / "output")
    current = f"{base}\\Memoire_DSS_Logistique_ElBayadh.docx"
    backup = f"{base}\\Memoire_DSS_Logistique_ElBayadh_v7c_BACKUP.docx"
    
    result = compare_tables(current, backup)
    sys.exit(0 if result else 1)

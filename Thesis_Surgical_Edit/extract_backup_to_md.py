#!/usr/bin/env python
"""
extract_backup_to_md.py — ACADEMIX Full Reclamation
==================================================
Extracts ALL content from the old v7c_BACKUP DOCX (the TRUE golden source)
into a clean markdown file with python-docx (no pandoc bracket artifacts).

Usage:
  python Thesis_Surgical_Edit/extract_backup_to_md.py

Output:
  Thesis_Surgical_Edit/output/golden_extracted.md
"""

import re, os, sys
from lxml import etree
from docx import Document
from docx.oxml.ns import qn

# ─── PATHS ───────────────────────────────────────────────────────────
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
BACKUP_PATH = os.path.join(PROJECT_ROOT, 'backups', 'thesis-versions',
                           'Memoire_DSS_Logistique_ElBayadh_v7c_BACKUP.docx')
OUTPUT_PATH = os.path.join(PROJECT_ROOT, 'Thesis_Surgical_Edit', 'output', 'golden_extracted.md')

NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

# ─── GROUND TRUTH CORRECT DATA (from ERP v13.3) ─────────────────────
GROUND_TRUTH = {
    'D_annual': '789',
    'D_old': '1,546',
    'PU': '4,500',
    'PU_old': '400',
    'Qstar': '37',
    'Qstar_old': '176',
    'ROP': '206',
    'ROP_old': '212.4',
    'SS': '200',
    'LT': '2',
    'S': '801.45',
    'I': '20%',
    'I_float': 0.20,
    'VERSION': 'v13.3',
    'VERSION_old': 'v13.2',
    'TC_denom': '900',  # PU * I = 4500 * 0.2
    'TC_denom_old': '80',  # 400 * 0.2
}

# Data fix replacements: old_pattern -> new_value
# These are applied to the extracted text
DATA_FIXES = [
    # D: 1,546 -> 789
    (r'1[,.]?546', '789'),
    # PU: 400 -> 4,500 (but don't mess up 400 in other contexts)
    # Pattern: "400 دج" or "400 DA" or "400دج"
    (r'400\s*(دج|DA|د\.ج)', r'4,500 \1'),
    # Also fix standalone PU values that appear in tables
    (r'400\s*دينار', '4,500 دينار'),
    # Q*: 176 -> 37
    (r'\b176\b(?!\s*مليون)', '37'),  # doesn't match 176 million or similar
    # ROP: 212.4 -> 206
    (r'212[.,]?4', '206'),
    # v13.2 -> v13.3
    (r'v13\.2', 'v13.3'),
    # TC formula: denominator 80 -> 900 (from PU*I)
    (r'(?<=/)\s*80\s*(?=\)?\s*×\s*|\))', '900'),
    # TC total value needs recalculation - flag for manual check
]


def extract_footnotes(doc):
    """Extract all footnotes from the DOCX footnotes part."""
    fns = {}
    for rel in doc.part.rels.values():
        if 'footnotes' in rel.reltype:
            fp = rel.target_part
            xml = etree.fromstring(fp.blob)
            for fn in xml.findall(f'.//{NS}footnote'):
                fid_attr = fn.get(f'{NS}id', '0')
                if not fid_attr:
                    continue
                fid = int(fid_attr)
                if fid < 1:
                    continue
                texts = []
                for t in fn.findall(f'.//{NS}t'):
                    if t.text:
                        texts.append(t.text)
                fns[fid] = ''.join(texts)
    return fns


def paragraph_has_footnote(para_xml):
    """Check if a paragraph XML has footnote references and return their IDs."""
    fn_refs = para_xml.findall(f'.//{NS}footnoteReference')
    return [int(ref.get(f'{NS}id', '0')) for ref in fn_refs if ref.get(f'{NS}id')]


def get_para_text_with_footnotes(para, fns):
    """Get paragraph text, replacing footnote markers with inline ^[text]."""
    xml = para._element
    fn_ids = paragraph_has_footnote(xml)
    
    if not fn_ids:
        return para.text or ''
    
    # Build text from XML runs, inserting footnote markers
    text_parts = []
    fn_idx = 0
    
    for elem in xml.iter():
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        
        if tag == 't' and elem.text:
            text_parts.append(elem.text)
        elif tag == 'footnoteReference' and fn_idx < len(fn_ids):
            fid = fn_ids[fn_idx]
            fn_text = fns.get(fid, f'[FN{fid}]')
            # Add space before footnote if preceding text doesn't end with space
            if text_parts and not text_parts[-1].endswith(' ') and not text_parts[-1].endswith('\n'):
                text_parts.append(' ')
            text_parts.append(f'^[{fn_text}]')
            fn_idx += 1
        elif tag == 'br':
            text_parts.append('\n')
        elif tag == 'tab':
            text_parts.append('\t')
    
    return ''.join(text_parts)


def get_style_info(para):
    """Get the paragraph's style info."""
    xml = para._element
    ppr = xml.find(f'{NS}pPr')
    if ppr is None:
        return '', False
    
    # Get paragraph style name
    p_style = ppr.find(f'{NS}pStyle')
    style_name = ''
    if p_style is not None:
        style_val = p_style.get(f'{NS}val', '')
        if style_val:
            style_name = style_val
    
    # Check for bold in any run
    bold = False
    for r in xml.findall(f'.//{NS}r'):
        rpr = r.find(f'{NS}rPr')
        if rpr is not None:
            b_elem = rpr.find(f'{NS}b')
            if b_elem is not None:
                b_val = b_elem.get(f'{NS}val', 'true')
                if b_val != 'false' and b_val != '0':
                    bold = True
                    break
    
    return style_name, bold


def detect_heading_level(style_name):
    """Convert style name to heading level (0 = not a heading)."""
    lower = style_name.lower()
    if 'heading1' in lower or 'titre1' in lower or 'heading 1' in lower:
        return 1
    if 'heading2' in lower or 'titre2' in lower or 'heading 2' in lower:
        return 2
    if 'heading3' in lower or 'titre3' in lower or 'heading 3' in lower:
        return 3
    if 'heading4' in lower or 'titre4' in lower or 'heading 4' in lower:
        return 4
    if 'heading5' in lower or 'titre5' in lower or 'heading 5' in lower:
        return 5
    return 0


def is_likely_toc_para(text):
    """Detect if a paragraph is a TOC entry (numbered with page refs)."""
    # TOC entries typically have tab-separated page numbers at the end
    if '\t' in text:
        return True
    # Or end with a bare number
    if re.match(r'^[\d\u0660-\u0669]+[\.\s]', text) and re.search(r'\t?\d+$', text):
        return True
    return False


def is_likely_cover_para(style_name, text):
    """Detect cover page paragraphs."""
    cover_keywords = ['الجمهورية الجزائرية', 'وزارة التربية الوطنية', 'مديرية التربية',
                      'معهد وطني متخصص', 'مذكرة تخرج', 'لنيل شهادة']
    for kw in cover_keywords:
        if kw in text:
            return True
    return False


def cleanup_md(text):
    """Clean up extracted markdown: remove revision notes, fix spacing, etc."""
    import re as re_mod
    
    lines = text.split('\n')
    cleaned = []
    prev_real_line = ''
    
    for line in lines:
        stripped = line.strip()
        
        if not stripped:
            cleaned.append(line)
            continue
        
        # Detect revision notes: lines that start with "N:" or "**N:" where N is a digit
        # and the TEXT after "N:" matches the previous content line (duplicate)
        
        # Extract text after stripping bold markers and number prefix
        content_stripped = stripped.strip('* ')
        
        m = re_mod.match(r'^(\d+):\s+(.*)$', content_stripped)
        if m:
            num = m.group(1)
            after_num = m.group(2)
            
            # Check if this is a DUPLICATE of the previous real line
            # (the text after "N:" is similar to the previous line)
            prev_clean = prev_real_line.strip('* ')
            
            # Remove any formatting from comparison
            prev_simple = re_mod.sub(r'[\*_#]', '', prev_clean).strip()
            after_simple = re_mod.sub(r'[\*_#]', '', after_num).strip()
            
            # If it's a duplicate (same content), skip it
            if after_simple and prev_simple and (
                after_simple == prev_simple or
                after_simple in prev_simple or
                prev_simple in after_simple
            ):
                continue  # Skip duplicate revision note
            
            # If it's NOT a duplicate, keep the content but strip the "N:" prefix
            # Reconstruct with bold if needed
            if stripped.startswith('**') and stripped.endswith('**'):
                new_line = '**' + after_num + '**'
            elif stripped.startswith('**'):
                new_line = '**' + after_num
            else:
                new_line = after_num
            cleaned.append(new_line)
            prev_real_line = after_num
            continue
        
        # Regular line - keep as-is
        cleaned.append(line)
        prev_real_line = stripped
    
    # Fix multiple consecutive blank lines (max 2)
    result = '\n'.join(cleaned)
    result = re_mod.sub(r'\n{3,}', '\n\n', result)
    
    return result


def rows_to_markdown_table(rows):
    """Convert table rows to markdown table format."""
    if not rows:
        return ''
    
    lines = []
    # Header row
    header = rows[0]
    lines.append('| ' + ' | '.join(header) + ' |')
    # Separator
    lines.append('| ' + ' | '.join(['---'] * len(header)) + ' |')
    # Data rows
    for row in rows[1:]:
        # Pad row to match header length
        padded = list(row) + [''] * (len(header) - len(row))
        if any(cell.strip() for cell in padded):  # skip empty rows
            lines.append('| ' + ' | '.join(padded[:len(header)]) + ' |')
    
    return '\n'.join(lines)


def extract_all(doc, fns):
    """Main extraction: walk through body XML children (paragraphs + tables in order).
    This correctly handles tables that are siblings of paragraphs in the document body."""
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    # Build lookup: XML element id -> (index, paragraph) for body-level paragraphs
    para_lookup = {}
    for pi, dp in enumerate(doc.paragraphs):
        para_lookup[id(dp._element)] = (pi, dp)
    
    # Build lookup: XML element id -> (index, table)
    table_lookup = {}
    for ti, t in enumerate(doc.tables):
        table_lookup[id(t._element)] = (ti, t)
    
    lines = []
    
    # Track state
    in_cover = True
    in_toc = False
    in_table_list = False
    in_bibliography = False
    last_heading = ''
    consecutive_blanks = 0
    
    started_content = False
    
    for child in doc.element.body:
        # Get local tag name (strip namespace)
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        
        if tag == 'p':
            # ─── PARAGRAPH ───
            para_info = para_lookup.get(id(child))
            if para_info is None:
                continue  # not a body-level paragraph
            pi, para = para_info
        
            text_raw = para.text or ''
            text = text_raw.strip()
            style_name, bold = get_style_info(para)
            heading_level = detect_heading_level(style_name)
            
            # Get text with footnotes inline
            full_text = get_para_text_with_footnotes(para, fns)
            full_text_stripped = full_text.strip()
            
            # ─── COVER PAGE DETECTION ───
            if not started_content:
                if is_likely_cover_para(style_name, text):
                    continue
                # Check for الملخص heading or المقدمة العامة heading
                if 'الملخص' in text and heading_level == 2:
                    started_content = True
                    lines.append('')
                    lines.append('## الملخص')
                    lines.append('')
                    continue
                if 'Résumé' in text and heading_level == 2:
                    # Already past الملخص — treat as content start
                    started_content = True
                    lines.append('')
                    lines.append('## Résumé')
                    lines.append('')
                    continue
                # Include dedication and thanks sections
                if text == 'إهداء':
                    lines.append('')
                    lines.append('## إهداء')
                    lines.append('')
                    started_content = True
                    continue
                if text == 'شكر وتقدير':
                    lines.append('')
                    lines.append('## شكر وتقدير')
                    lines.append('')
                    started_content = True
                    continue
                if text in ('فهرس المحتويات', 'قائمة الجداول'):
                    continue
                # If we hit a heading that's not cover-related, start content
                if heading_level > 0:
                    started_content = True
                else:
                    continue  # Still in pre-content area
            
            # ─── SKIP TOC AND TABLE LIST ───
            if 'فهرس المحتويات' in text:
                in_toc = True
                continue
            if 'قائمة الجداول' in text:
                in_table_list = True
                continue
            if in_toc and heading_level > 0 and 'فهرس' not in text and 'قائمة' not in text:
                in_toc = False
            if in_table_list and heading_level > 0 and 'فهرس' not in text and 'قائمة' not in text:
                in_table_list = False
            if in_toc or in_table_list:
                continue
            
            # ─── HANDLE SPECIAL NON-HEADING TITLES ───
            # إهداء and شكر وتقدير are Normal style (not headings) in the old backup
            if text == 'شكر وتقدير':
                lines.append('')
                lines.append('## شكر وتقدير')
                lines.append('')
                continue
            
            # ─── HANDLE EMPTY LINES ───
            if not full_text_stripped:
                consecutive_blanks += 1
                # Add blank lines sparingly (max 2)
                if consecutive_blanks <= 2:
                    lines.append('')
                continue
            consecutive_blanks = 0
            
            # ─── HANDLE HEADINGS ───
            if heading_level > 0:
                prefix = '#' * heading_level
                # Use full_text_stripped to preserve footnotes in headings
                heading_text = full_text_stripped if full_text_stripped else text
                lines.append('')
                lines.append(f'{prefix} {heading_text}')
                lines.append('')
                last_heading = text
                continue
            
            # ─── HANDLE NORMAL TEXT ───
            # Apply bold/emphasis
            if bold and not full_text_stripped.startswith('**'):
                full_text_stripped = f'**{full_text_stripped}**'
            
            # Check if it's a list item (numbered or bulleted)
            xml = para._element
            numPr = xml.find(f'.//{NS}numPr')
            is_list_item = numPr is not None
            
            if is_list_item:
                # Get numbering info
                num_id = numPr.find(f'{NS}numId')
                ilvl = numPr.find(f'{NS}ilvl')
                indent = 0
                if ilvl is not None:
                    indent = int(ilvl.get(f'{NS}val', '0'))
                
                # Simple list representation
                prefix = '  ' * indent + '- '
                lines.append(f'{prefix}{full_text_stripped}')
            else:
                # Regular paragraph
                lines.append(full_text_stripped)
        
        elif tag == 'tbl':
            # ─── TABLE ───
            table_info = table_lookup.get(id(child))
            if table_info is None:
                continue
            ti, table = table_info
            
            # Build rows_data directly from the table
            rows_data = []
            for ri, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]
                rows_data.append(row_data)
            
            md_table = rows_to_markdown_table(rows_data)
            lines.append('')
            lines.append(md_table)
            lines.append('')
    
    return '\n'.join(lines)


def fix_data_values(text):
    """Apply data value fixes to the extracted text."""
    import re as re_mod
    
    fixes_applied = 0
    
    # D: 1,546 → 789
    text, count = re_mod.subn(r'\b1[,.]?546\b', '789', text)
    fixes_applied += count
    
    # PU: 400 DZD → 4,500 DZD (careful with context)
    text, count = re_mod.subn(r'400\s*(دج|DA|د\.ج|دينار)', r'4,500 \1', text)
    fixes_applied += count
    
    # Also catch "400" as a standalone unit price reference
    # Pattern like: "سعر الوحدة 400" or "400 دينار جزائري"
    text, count = re_mod.subn(r'سعر\s+(الوحدة|شراء)\s+400', r'سعر \1 4,500', text)
    fixes_applied += count
    
    # Q*: 176 → 37
    text, count = re_mod.subn(r'\b176\b', '37', text)
    fixes_applied += count
    
    # ROP: 212.4 → 206
    text, count = re_mod.subn(r'212[.,]?4', '206', text)
    fixes_applied += count
    
    # v13.2 → v13.3
    text, count = re_mod.subn(r'v13\.2', 'v13.3', text)
    fixes_applied += count
    
    # TC denominator: (Q/2)*PU*I where PU=400, I=20% → denominator was 80
    # Now PU=4500, I=20% → denominator should be 900
    # Pattern: /80  or  /80)  etc in TC formulas
    text, count = re_mod.subn(r'(?<=[\(/\s])80(?=\s*[×\)\s,])', '900', text)
    fixes_applied += count
    
    # Also fix TC = D/Q * S + Q/2 * PU * I calculations
    # The old TC = 14,080.xx DZD with old values
    # New TC ≈ 33,746 DZD (will flag for manual check)
    
    print(f'  [FIXES] Applied {fixes_applied} data value replacements')
    return text


def main():
    print('=' * 60)
    print('  ACADEMIX — Full Reclamation from Old Backup')
    print('=' * 60)
    
    if not os.path.exists(BACKUP_PATH):
        print(f'[ERROR] Backup not found: {BACKUP_PATH}')
        sys.exit(1)
    
    print(f'  Source: {BACKUP_PATH}')
    
    # Load DOCX
    doc = Document(BACKUP_PATH)
    print(f'  Paragraphs: {len(doc.paragraphs)}')
    print(f'  Tables: {len(doc.tables)}')
    
    # Extract footnotes
    fns = extract_footnotes(doc)
    print(f'  Footnotes: {len(fns)}')
    
    # Verify 46 footnotes
    if len(fns) != 46:
        print(f'  [WARNING] Expected 46 footnotes, found {len(fns)}')
    
    # Extract ALL content as markdown
    print('\n  Extracting content...')
    md_content = extract_all(doc, fns)
    
    # Clean up (remove revision notes, fix spacing)
    print('\n  Cleaning up...')
    md_content = cleanup_md(md_content)
    
    # Fix data values
    print('\n  Fixing data values...')
    md_content = fix_data_values(md_content)
    
    # Build YAML front matter
    yaml = '''---
title: "نظام دعم القرار لتسيير المخزونات"
author: "ماحي كمال عبد الغني"
date: "2026-05-18"
lang: ar
dir: rtl
---
'''
    
    # Write output
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    with open(OUTPUT_PATH, 'w', encoding='utf-8') as f:
        f.write(yaml)
        # Skip empty leading lines
        md_content = md_content.lstrip('\n')
        f.write(md_content)
    
    char_count = len(md_content)
    para_count = md_content.count('\n\n') + 1
    print(f'\n  Output: {OUTPUT_PATH}')
    print(f'  Characters: {char_count}')
    print(f'  Paragraphs (approx): {para_count}')
    print(f'\n  [DONE] Extraction complete.')
    
    # Compare with old backup stats
    backup_chars = sum(len(p.text or '') for p in doc.paragraphs)
    print(f'  Old backup chars: {backup_chars}')
    print(f'  Extracted chars:  {char_count}')
    ratio = (char_count / backup_chars * 100) if backup_chars else 0
    print(f'  Retention: {ratio:.1f}%')


if __name__ == '__main__':
    main()

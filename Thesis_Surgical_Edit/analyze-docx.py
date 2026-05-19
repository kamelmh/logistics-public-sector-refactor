from docx import Document
import json

doc = Document('output/Memoire_DSS_Logistique_ElBayadh.docx')

# Check sections and page numbering
print('=== SECTIONS ===')
for i, section in enumerate(doc.sections):
    sectPr = section._sectPr
    pgn = sectPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pgNumType')
    fmt = pgn.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fmt') if pgn is not None else 'none'
    print(f'Section {i}: pgNumType={fmt}')

# Check heading styles
print('\n=== HEADING STYLES ===')
heading_counts = {}
for p in doc.paragraphs:
    if p.style.name.startswith('Heading'):
        heading_counts[p.style.name] = heading_counts.get(p.style.name, 0) + 1
for style, count in sorted(heading_counts.items()):
    print(f'{style}: {count}')

# Check table count
print(f'\n=== TABLES ===')
print(f'Total tables: {len(doc.tables)}')

# Check footnotes
print(f'\n=== FOOTNOTES ===')
print(f'Footnotes: {len(doc.footnotes)}')

# Check first few paragraphs for RTL
print('\n=== FIRST 10 PARAGRAPHS ===')
for i, p in enumerate(doc.paragraphs[:10]):
    text = p.text[:60] if p.text else '(empty)'
    print(f'P{i}: style={p.style.name}, align={p.alignment}, text="{text}"')

# Check for فصل الأول to find section break
print('\n=== SECTION BREAK LOCATION ===')
for i, p in enumerate(doc.paragraphs):
    if 'الفصل الأول' in p.text:
        print(f'الفصل الأول found at paragraph {i}')
        break

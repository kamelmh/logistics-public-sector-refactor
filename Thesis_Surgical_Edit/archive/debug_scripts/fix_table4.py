from docx import Document
doc = Document('Thesis_Surgical_Edit/output/Latest-thesis-backup-Memoire_DSS_Logistique_ElBayadh.docx')
table = doc.tables[4]

# Fix ART-003 S (Row 3, Cell 2) - should be 400 not 4,500
table.rows[3].cells[2].paragraphs[0].runs[0].text = "400"

# Fix ART-004 D (Row 4, Cell 1) - should be 400 not 4,500
table.rows[4].cells[1].paragraphs[0].runs[0].text = "400"

doc.save('Thesis_Surgical_Edit/output/Latest-thesis-backup-Memoire_DSS_Logistique_ElBayadh.docx')
print('Table 4 fixed!')
import sys
import os
import zipfile
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_run_font(run, font_name):
    run.font.name = font_name
    # Set the complex script font (for Arabic)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)

def apply_font_to_docx(docx_path, output_path, font_name):
    print(f"Applying font '{font_name}' to {docx_path}...")
    doc = Document(docx_path)
    
    # 1. Update Styles
    for style in doc.styles:
        if hasattr(style, 'font'):
            style.font.name = font_name
            
    # 2. Update Paragraphs
    for p in doc.paragraphs:
        for run in p.runs:
            set_run_font(run, font_name)
            
    # 3. Update Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, font_name)
                        
    doc.save(output_path)
    print(f"Saved: {output_path}")
    
    # 4. Update Footnotes at ZIP level (since python-docx doesn't expose footnotes easily)
    fix_footnotes_font_zip(output_path, font_name)

def fix_footnotes_font_zip(docx_path, font_name):
    fn_file = 'word/footnotes.xml'
    with zipfile.ZipFile(docx_path, 'r') as z:
        if fn_file not in z.namelist():
            return
        raw = z.read(fn_file).decode('utf-8')
        
    # Replace Traditional Arabic with the new font name
    raw = raw.replace('Traditional Arabic', font_name)
    raw = raw.replace('Scheherazade New', font_name)
    
    # Also replace in rFonts elements
    rfonts_pat = re.compile(r'<w:rFonts\s+[^>]*/>')
    def _fix_rfonts(m):
        # Replace ascii, hAnsi, cs with the new font
        tag = m.group(0)
        tag = re.sub(r'w:ascii="[^"]*"', f'w:ascii="{font_name}"', tag)
        tag = re.sub(r'w:hAnsi="[^"]*"', f'w:hAnsi="{font_name}"', tag)
        tag = re.sub(r'w:cs="[^"]*"', f'w:cs="{font_name}"', tag)
        return tag
    raw = rfonts_pat.sub(_fix_rfonts, raw)
    
    # Write back to zip
    _zip_replace(docx_path, {fn_file: raw.encode('utf-8')})
    print(f"Updated footnotes font to '{font_name}'")

def _zip_replace(docx_path, replacements):
    temp_path = docx_path + '.tmp'
    with zipfile.ZipFile(docx_path, 'r') as zin:
        with zipfile.ZipFile(temp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename in replacements:
                    zout.writestr(item.filename, replacements[item.filename])
                else:
                    zout.writestr(item.filename, zin.read(item.filename))
    os.replace(temp_path, docx_path)

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: python apply_docx_font.py <input.docx> <output.docx> <font_name>")
        sys.exit(1)
    apply_font_to_docx(sys.argv[1], sys.argv[2], sys.argv[3])

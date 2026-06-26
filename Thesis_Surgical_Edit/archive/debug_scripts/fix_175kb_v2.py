from docx import Document
doc = Document(r'D:\New folder\19-06\Memoire_DSS_Logistique_ElBayadh.docx')

# Fix Table 1 - remove duplication
doc.tables[1].rows[1].cells[1].paragraphs[0].runs[0].text = "789 وحدة"
doc.tables[1].rows[3].cells[1].paragraphs[0].runs[0].text = "4,500 دج"

doc.save(r'D:\New folder\19-06\Memoire_DSS_Logistique_ElBayadh.docx')
print('Fixed duplication!')
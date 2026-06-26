from docx import Document
doc = Document(r'D:\New folder\19-06\Memoire_DSS_Logistique_ElBayadh.docx')

# Check runs in table 1
table = doc.tables[1]
for row_idx, row in enumerate(table.rows):
    for cell_idx, cell in enumerate(row.cells):
        for p_idx, p in enumerate(cell.paragraphs):
            for r_idx, run in enumerate(p.runs):
                if run.text.strip():
                    print(f'Row {row_idx}, Cell {cell_idx}, Para {p_idx}, Run {r_idx}: "{run.text}"')
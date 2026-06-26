from docx import Document
doc = Document(r'D:\New folder\19-06\Memoire_DSS_Logistique_ElBayadh.docx')

# Fix Table 1 (parameters table)
# Row 1, Cell 1: D = 1,546 -> 789
doc.tables[1].rows[1].cells[1].paragraphs[0].runs[0].text = "789 وحدة"
# Row 3, Cell 1: PU = 400 -> 4,500
doc.tables[1].rows[3].cells[1].paragraphs[0].runs[0].text = "4,500 دج"

# Fix Table 4 (Wilson table) - ART-001 row (Row 1)
# Cell 1: D = 1,546 -> 789
doc.tables[4].rows[1].cells[1].paragraphs[0].runs[0].text = "789"
# Cell 3: PU = 400 -> 4,500
doc.tables[4].rows[1].cells[3].paragraphs[0].runs[0].text = "4,500"

doc.save(r'D:\New folder\19-06\Memoire_DSS_Logistique_ElBayadh.docx')
print('Fixed 175KB file tables!')
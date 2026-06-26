from docx import Document
import re

def fix_docx(file_path):
    doc = Document(file_path)
    
    # Fix tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        text = run.text
                        # Fix values
                        text = text.replace('1,546', '789')
                        text = text.replace('1546', '789')
                        text = text.replace('176', '37')
                        text = text.replace('212.4', '206')
                        text = text.replace('400 دج', '4,500 دج')
                        text = text.replace('400دج', '4,500 دج')
                        text = text.replace('400 DA', '4,500 DA')
                        text = text.replace('400 د.ج', '4,500 د.ج')
                        text = text.replace('400 دينار', '4,500 دينار')
                        text = text.replace('400', '4,500') # This might be too aggressive, let's be careful
                        # Wait, I should only replace 400 if it's in a specific context
                        run.text = text

    # Fix paragraphs
    for p in doc.paragraphs:
        for run in p.runs:
            text = run.text
            text = text.replace('1,546', '789')
            text = text.replace('1546', '789')
            text = text.replace('176', '37')
            text = text.replace('212.4', '206')
            text = text.replace('400 دج', '4,500 دج')
            text = text.replace('400دج', '4,500 دج')
            text = text.replace('400 DA', '4,500 DA')
            text = text.replace('400 د.ج', '4,500 د.ج')
            text = text.replace('400 دينار', '4,500 دينار')
            # text = text.replace('400', '4,500') # Too aggressive
            run.text = text

    doc.save(file_path)
    print(f'Fixed {file_path}')

# Let's do it more carefully
def fix_docx_carefully(file_path):
    doc = Document(file_path)
    
    # Fix tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        text = run.text
                        # Fix D
                        text = re.sub(r'\b1,546\b', '789', text)
                        text = re.sub(r'\b1546\b', '789', text)
                        # Fix Q*
                        text = re.sub(r'\b176\b', '37', text)
                        # Fix ROP
                        text = re.sub(r'\b212\.4\b', '206', text)
                        # Fix PU (only if followed by currency or in specific context)
                        text = re.sub(r'400\s*(دج|DA|د\.ج|دينار)', r'4,500 \1', text)
                        run.text = text

    # Fix paragraphs
    for p in doc.paragraphs:
        for run in p.runs:
            text = run.text
            # Fix D
            text = re.sub(r'\b1,546\b', '789', text)
            text = re.sub(r'\b1546\b', '789', text)
            # Fix Q*
            text = re.sub(r'\b176\b', '37', text)
            # Fix ROP
            text = re.sub(r'\b212\.4\b', '206', text)
            # Fix PU
            text = re.sub(r'400\s*(دج|DA|د\.ج|دينار)', r'4,500 \1', text)
            run.text = text

    doc.save(file_path)
    print(f'Fixed {file_path} carefully!')

fix_docx_carefully(r'D:\New folder\19-06\Memoire_DSS_Logistique_ElBayadh.docx')

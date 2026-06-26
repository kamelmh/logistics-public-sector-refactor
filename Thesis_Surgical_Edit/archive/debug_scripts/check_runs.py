from docx import Document
doc = Document('Thesis_Surgical_Edit/output/Latest-thesis-backup-Memoire_DSS_Logistique_ElBayadh.docx')
table = doc.tables[4]
for row_idx, row in enumerate(table.rows):
    for cell_idx, cell in enumerate(row.cells):
        for p_idx, p in enumerate(cell.paragraphs):
            for r_idx, run in enumerate(p.runs):
                if run.text.strip():
                    print(f'Row {row_idx}, Cell {cell_idx}, Para {p_idx}, Run {r_idx}: "{run.text}"')
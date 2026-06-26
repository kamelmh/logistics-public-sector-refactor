"""Quick baseline check for a DOCX file."""
import sys, os
from docx import Document

path = sys.argv[1]
doc = Document(path)
paras = len(doc.paragraphs)
abs_found = any('الملخص' in p.text for p in doc.paragraphs)
size_kb = os.path.getsize(path) // 1024
print(f"Paras: {paras}, Abstract: {abs_found}, Size: {size_kb}KB")

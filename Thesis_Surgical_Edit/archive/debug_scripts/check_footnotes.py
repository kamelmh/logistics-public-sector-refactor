from docx import Document

doc = Document(r'C:\Users\Administrator\Dropbox\Logistics.Public.Sector.Refactor\Thesis_Surgical_Edit\output\Memoire_DSS_Logistique_ElBayadh_v13.4_calibrated.docx')

# Find footnote references (superscript numbers or patterns)
print('=== FOOTNOTE REFERENCES IN PARAGRAPHS ===')
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        if run.font.superscript and run.text.strip():
            print(f'Para {i}, Run: "{run.text}" (superscript={run.font.superscript})')

# Check end of document for footnote texts
print()
print('=== LAST 30 PARAGRAPHS (footnotes area) ===')
for i, p in enumerate(doc.paragraphs[-30:]):
    idx = len(doc.paragraphs) - 30 + i
    if p.text.strip():
        print(f'[{idx}] {p.text[:150]}')
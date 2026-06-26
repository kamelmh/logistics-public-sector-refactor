"""insert_toc.py — Inserts a real Word TOC field into a DOCX.
Usage: python insert_toc.py <path/to.docx> --save
"""

import sys, os
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn

def insert_toc(docx_path, save=False):
    doc = Document(docx_path)
    
    # Find the position to insert TOC. Usually after the title/subtitle.
    # We'll insert it after the first few paragraphs or at a specific heading.
    # For this thesis, let's insert it after the 'Résumé' or 'Abstract' section, 
    # or just at the very beginning after the cover.
    
    # Let's find the paragraph that says "Table of Contents" or "Férendex المحتويات"
    # or just insert it after the first H1 if it's the TOC.
    
    # Actually, the MD has:
    # ## فهرس المحتويات
    # *جدول المحتويات سيتم توليده تلقائياً عند بناء الوثيقة*
    
    # We'll look for that placeholder.
    target_text = "جدول المحتويات سيتم توليده تلقائياً"
    
    inserted = False
    for p in doc.paragraphs:
        if target_text in p.text:
            # Insert TOC field in this paragraph or right after it
            # A TOC field is a field code: { TOC \h \z \u }
            
            # We need to create a paragraph and add the field
            # In docx, a field is a sequence of runs: 
            # <w:r><w:fldChar w:fldCharType="begin"/></w:r>
            # <w:r><w:instrText>TOC \h \z \u</w:instrText></w:r>
            # <w:r><w:fldChar w:fldCharType="separate"/></w:r>
            # <w:r><w:fldChar w:fldCharType="end"/></w:r>
            
            p._element.append(parse_xml(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:fldChar w:fldCharType="begin"/>'
                '</w:r>'))
            
            p._element.append(parse_xml(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:instrText>TOC \\h \\z \\u</w:instrText>'
                '</w:r>'))
            
            p._element.append(parse_xml(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:fldChar w:fldCharType="separate"/>'
                '</w:r>'))
            
            p._element.append(parse_xml(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:fldChar w:fldCharType="end"/>'
                '</w:r>'))
            
            inserted = True
            print(f"Successfully inserted TOC field at placeholder: '{p.text.strip()}'")
            break
            
    if not inserted:
        print("Could not find TOC placeholder. Inserting at the beginning of the document.")
        # Insert at the very beginning (after cover if possible, but doc.paragraphs[0] is usually after cover)
        first_p = doc.paragraphs[0]._element
        first_p.add_previous(parse_xml(
            '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:pPr><w:spacing w:after="200"/></w:pPr>'
            '<w:r><w:fldChar w:fldCharType="begin"/></w:r>'
            '<w:r><w:instrText>TOC \\h \\z \\u</w:instrText></w:r>'
            '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
            '<w:r><w:fldChar w:fldCharType="end"/></w:r>'
            '</w:p>'))
        inserted = True

    if save:
        doc.save(docx_path)
        print(f"Saved changes to {docx_path}")
    else:
        print("Dry run: No changes saved. Use --save to apply.")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python insert_toc.py <path/to.docx> [--save]")
        sys.exit(1)
    
    path = sys.argv[1]
    save = '--save' in sys.argv
    insert_toc(path, save)

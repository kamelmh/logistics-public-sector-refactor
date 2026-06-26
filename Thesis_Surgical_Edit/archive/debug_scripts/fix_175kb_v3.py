from docx import Document
doc = Document(r'D:\New folder\19-06\Memoire_DSS_Logistique_ElBayadh.docx')

# Fix Table 1 (index 1) - clear the extra runs
# Row 1, Cell 1: clear Run 1 (" وحدة")
doc.tables[1].rows[1].cells[1].paragraphs[0].runs[1].text = ""
# Row 3, Cell 1: clear Run 1 (" دج")
doc.tables[1].rows[3].cells[1].paragraphs[0].runs[1].text = ""

doc.save(r'D:\New folder\19-06\Memoire_DSS_Logistique_ElBayadh.docx')
print('Fixed duplication!')
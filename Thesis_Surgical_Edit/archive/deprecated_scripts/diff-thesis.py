"""
⚠️  DEPRECATED — Replaced by: Standalone — kept as utility
    This script is superseded and no longer part of the active pipeline.
    Kept for reference only. Do not use in new workflows.
"""
"""diff-thesis.py — Compare two thesis DOCX files structurally and textually
Usage: python diff-thesis.py <reference.docx> <comparison.docx> [--json] [--output report.txt]
"""
import sys, os, json, argparse, zipfile, hmac, hashlib, re
from xml.etree import ElementTree as ET
from difflib import SequenceMatcher
from docx import Document
from docx.shared import Cm, Pt

W_NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

def get_docx_metrics(path):
    doc = Document(path)
    paras = doc.paragraphs
    sections = doc.sections
    s = sections[0] if sections else None
    fn_count = 0
    try:
        with zipfile.ZipFile(path, 'r') as z:
            if 'word/footnotes.xml' in z.namelist():
                tree = ET.parse(z.open('word/footnotes.xml'))
                for fn in tree.findall('.//w:footnote', W_NS):
                    fid = fn.attrib.get(f'{{{W_NS["w"]}}}id', '')
                    if fid not in ('0', '-1'): fn_count += 1
    except: pass

    h1 = []; h2 = []; h3 = []
    for p in paras:
        if p.style and p.style.name:
            sn = p.style.name
            if 'Heading 1' in sn: h1.append(p.text[:100])
            elif 'Heading 2' in sn: h2.append(p.text[:100])
            elif 'Heading 3' in sn: h3.append(p.text[:100])
    
    body_text = [p.text for p in paras if p.style and p.style.name in 
        ['Normal', 'Compact', 'Body Text', 'List Paragraph', 'No Spacing'] and p.text.strip()]
    total_words = sum(len(t.split()) for t in body_text)
    
    return {
        "file_size_kb": os.path.getsize(path) // 1024,
        "paragraph_count": len(paras),
        "section_count": len(sections),
        "table_count": len(doc.tables),
        "footnote_count": fn_count,
        "h1_count": len(h1),
        "h2_count": len(h2),
        "h3_count": len(h3),
        "h1_headings": h1,
        "h2_headings": h2,
        "h3_headings": h3,
        "body_paragraphs": len(body_text),
        "body_words": total_words,
        "page_info": {
            "width_cm": round(s.page_width.cm, 1) if s else None,
            "height_cm": round(s.page_height.cm, 1) if s else None,
            "margin_top_cm": round(s.top_margin.cm, 1) if s else None,
            "margin_bottom_cm": round(s.bottom_margin.cm, 1) if s else None,
            "margin_left_cm": round(s.left_margin.cm, 1) if s else None,
            "margin_right_cm": round(s.right_margin.cm, 1) if s else None,
        } if s else None
    }

def heading_similarity(ref_headings, comp_headings):
    if not ref_headings or not comp_headings: return 0.0
    ref_set = set(h.strip() for h in ref_headings if h.strip())
    comp_set = set(h.strip() for h in comp_headings if h.strip())
    if not ref_set or not comp_set: return 0.0
    intersection = ref_set & comp_set
    union = ref_set | comp_set
    return round(len(intersection) / len(union) * 100, 1)

def structural_diff(m1, m2):
    diffs = []
    fields = [
        "paragraph_count", "section_count", "table_count", "footnote_count",
        "h1_count", "h2_count", "h3_count", "file_size_kb"
    ]
    for f in fields:
        a = m1.get(f, 0) or 0
        b = m2.get(f, 0) or 0
        delta = b - a
        if delta != 0:
            arrow = "▲" if delta > 0 else "▼"
            diffs.append({"field": f, "ref": a, "comp": b, "delta": delta, "arrow": arrow})
    return diffs

def section_overlap(m1, m2):
    """Compute overlap for H1, H2, H3 heading levels"""
    results = {}
    for level, key in [("H1", "h1_headings"), ("H2", "h2_headings"), ("H3", "h3_headings")]:
        sim = heading_similarity(m1.get(key, []), m2.get(key, []))
        only_ref = [h for h in m1.get(key, []) if h.strip() and h not in m2.get(key, [])]
        only_comp = [h for h in m2.get(key, []) if h.strip() and h not in m1.get(key, [])]
        results[level] = {
            "similarity_pct": sim,
            "ref_only_count": len(only_ref),
            "comp_only_count": len(only_comp),
            "ref_only": only_ref[:10],
            "comp_only": only_comp[:10],
        }
    return results

def run_diff(ref_path, comp_path):
    if not os.path.exists(ref_path): return {"error": f"Reference not found: {ref_path}"}
    if not os.path.exists(comp_path): return {"error": f"Comparison not found: {comp_path}"}
    
    ref_metrics = get_docx_metrics(ref_path)
    comp_metrics = get_docx_metrics(comp_path)
    struct = structural_diff(ref_metrics, comp_metrics)
    overlap = section_overlap(ref_metrics, comp_metrics)
    
    # Font/style diff (sample)
    ref_doc = Document(ref_path)
    comp_doc = Document(comp_path)
    
    return {
        "ref": {"path": ref_path, "label": os.path.basename(ref_path), **ref_metrics},
        "comp": {"path": comp_path, "label": os.path.basename(comp_path), **comp_metrics},
        "structural_diffs": struct,
        "section_overlap": overlap,
        "summary": {
            "total_diffs": len(struct),
            "h1_similarity": overlap["H1"]["similarity_pct"],
            "h2_similarity": overlap["H2"]["similarity_pct"],
            "h3_similarity": overlap["H3"]["similarity_pct"],
            "size_diff_kb": comp_metrics["file_size_kb"] - ref_metrics["file_size_kb"],
            "para_diff": comp_metrics["paragraph_count"] - ref_metrics["paragraph_count"],
        }
    }

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Compare two thesis DOCX files")
    parser.add_argument("reference", help="Reference DOCX path")
    parser.add_argument("comparison", help="Comparison DOCX path")
    parser.add_argument("--json", action="store_true", help="JSON output")
    parser.add_argument("--output", help="Write report to file")
    args = parser.parse_args()
    
    result = run_diff(args.reference, args.comparison)
    
    if "error" in result:
        print(f"[DIFF ERROR] {result['error']}")
        sys.exit(1)
    
    if args.json:
        output = json.dumps(result, indent=2, ensure_ascii=False)
    else:
        r = result["ref"]; c = result["comp"]
        lines = []
        lines.append("=" * 70)
        lines.append(f"THESIS DIFF: {r['label']} vs {c['label']}")
        lines.append("=" * 70)
        lines.append(f"  {r['label']:30s}: {r['file_size_kb']}KB, {r['paragraph_count']}p, {r['table_count']}tbl, {r['footnote_count']}fn")
        lines.append(f"  {c['label']:30s}: {c['file_size_kb']}KB, {c['paragraph_count']}p, {c['table_count']}tbl, {c['footnote_count']}fn")
        lines.append(f"  H1: {r['h1_count']} vs {c['h1_count']}  H2: {r['h2_count']} vs {c['h2_count']}  H3: {r['h3_count']} vs {c['h3_count']}")
        lines.append("")
        
        if result["structural_diffs"]:
            lines.append("── Structural Differences ──")
            for d in result["structural_diffs"]:
                lines.append(f"  {d['arrow']} {d['field']}: {d['ref']} → {d['comp']} ({d['delta']:+d})")
            lines.append("")
        
        lines.append("── Section Overlap ──")
        for level in ["H1", "H2", "H3"]:
            o = result["section_overlap"][level]
            lines.append(f"  {level}: similarity={o['similarity_pct']}%  ref-only={o['ref_only_count']}  comp-only={o['comp_only_count']}")
            if o["ref_only"]:
                lines.append(f"    Ref only: {', '.join(o['ref_only'][:5])}")
            if o["comp_only"]:
                lines.append(f"    Comp only: {', '.join(o['comp_only'][:5])}")
        lines.append("")
        
        s = result["summary"]
        lines.append("── Summary ──")
        lines.append(f"  Total structural diffs: {s['total_diffs']}")
        lines.append(f"  H1 similarity: {s['h1_similarity']}%")
        lines.append(f"  H2 similarity: {s['h2_similarity']}%")
        lines.append(f"  H3 similarity: {s['h3_similarity']}%")
        lines.append(f"  Size diff: {s['size_diff_kb']:+d}KB")
        lines.append(f"  Para diff: {s['para_diff']:+d}")
        lines.append("=" * 70)
        output = "\n".join(lines)
    
    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            f.write(output if args.json else output)
        print(f"Report written to: {args.output}")
    else:
        print(output)

#!/usr/bin/env python3
"""Deep analysis of thesis DOCX output — all structural, formatting, and corruption checks."""
import sys, os, re, json
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Cm, Pt, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

DOCX_PATH = r"C:\Users\Administrator\Dropbox\Logistics.Public.Sector.Refactor\Thesis_Surgical_Edit\output\Memoire_DSS_Logistique_ElBayadh.docx"

def analyze_docx(path):
    results = {}
    
    if not os.path.exists(path):
        return {"STATUS": "FILE NOT FOUND", "path": path}
    
    results["FILE_SIZE_KB"] = round(os.path.getsize(path) / 1024, 1)
    
    # Load document
    try:
        doc = Document(path)
    except Exception as e:
        return {"STATUS": "CORRUPTED — cannot open", "error": str(e)}
    
    # ── 1. BASIC COUNTS ──
    paras = doc.paragraphs
    tables = doc.tables
    sections = doc.sections
    
    # Count footnotes from XML
    footnotes_count = 0
    try:
        footnotes_part = doc.part.package.part_related_by(
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
        )
        fn_xml = footnotes_part.blob
        footnotes_count = len(re.findall(rb'<w:footnote\b[^>]*w:type="normal"', fn_xml))
    except:
        try:
            fn_elem = doc.element.body.findall('.//' + qn('w:footnoteReference'))
            footnotes_count = len([f for f in fn_elem if f.getparent().getparent().tag != qn('w:body')])
        except:
            footnotes_count = "UNKNOWN"
    
    # Count hyperlinks
    hyperlinks = len(doc.element.body.findall('.//' + qn('w:hyperlink')))
    
    # Count bookmarks
    bookmarks = len(doc.element.body.findall('.//' + qn('w:bookmarkStart')))
    
    # Count TOC/TOF fields
    toc_count = 0
    tof_count = 0
    all_fields = doc.element.body.findall('.//' + qn('w:fldChar'))
    instrTexts = doc.element.body.findall('.//' + qn('w:instrText'))
    for it in instrTexts:
        txt = it.text or ""
        if "TOC" in txt.upper():
            toc_count += 1
        if "TOF" in txt.upper():
            tof_count += 1
    
    # PAGE fields
    page_fields = [it.text for it in instrTexts if it.text and "PAGE" in it.text.upper()]
    
    # NUMPAGES fields
    numpages_fields = [it.text for it in instrTexts if it.text and "NUMPAGES" in it.text.upper()]
    
    results["DETAILED_COUNTS"] = {
        "paragraphs": len(paras),
        "tables": len(tables),
        "sections": len(sections),
        "footnotes": footnotes_count,
        "hyperlinks": hyperlinks,
        "bookmarks": bookmarks,
        "toc_fields": toc_count,
        "tof_fields": tof_count,
        "page_fields": len(page_fields),
        "page_field_details": [f.strip() for f in page_fields],
        "numpages_fields": len(numpages_fields),
        "instrText_total": len(instrTexts),
    }
    
    # ── 2. HEADING HIERARCHY ──
    heading_counts = {1: 0, 2: 0, 3: 0, 4: 0}
    heading_errors = []
    heading_texts = []
    prev_level = 0
    for p in paras:
        if p.style and p.style.name and p.style.name.startswith("Heading"):
            try:
                level = int(p.style.name.replace("Heading", "").strip())
            except:
                level = 0
            if level in heading_counts:
                heading_counts[level] += 1
            heading_texts.append((level, p.text[:80]))
            if level > prev_level + 1 and prev_level > 0:
                heading_errors.append(f"Jump H{prev_level} -> H{level}: '{p.text[:50]}'")
            prev_level = level
    
    results["HEADING_HIERARCHY"] = {
        "H1": heading_counts.get(1, 0),
        "H2": heading_counts.get(2, 0),
        "H3": heading_counts.get(3, 0),
        "H4": heading_counts.get(4, 0),
        "total": sum(heading_counts.values()),
        "ordering_errors": heading_errors,
        "sample_headings": heading_texts[:30],
    }
    
    # ── 3. RTL ALIGNMENT ON FOOTNOTES ──
    footnote_issues = []
    try:
        fn_xml_str = fn_xml.decode("utf-8", errors="replace")
        fn_blocks = re.findall(r'<w:footnote\b[^>]*w:type="normal"[^>]*>(.*?)</w:footnote>', fn_xml_str, re.DOTALL)
        for i, block in enumerate(fn_blocks, 1):
            has_bidi = "w:bidi" in block
            has_rtl = 'w:jc="right"' in block or 'w:bidi' in block
            # Check for <w:rPr> with <w:bidi/>
            rpr_blocks = re.findall(r'<w:rPr>(.*?)</w:rPr>', block, re.DOTALL)
            bidi_in_rpr = any("w:bidi" in rpr for rpr in rpr_blocks)
            # Check paragraph-level bidi
            ppr_blocks = re.findall(r'<w:pPr>(.*?)</w:pPr>', block, re.DOTALL)
            bidi_in_ppr = any("w:bidi" in ppr for ppr in ppr_blocks)
            
            if not bidi_in_ppr and not bidi_in_rpr:
                footnote_issues.append(f"Footnote {i}: missing bidi in pPr/rPr")
    except Exception as e:
        footnote_issues.append(f"Cannot parse footnotes XML: {e}")
    
    results["FOOTNOTE_RTL"] = {
        "checked": footnotes_count,
        "issues": footnote_issues,
        "status": "PASS" if not footnote_issues else f"FAIL ({len(footnote_issues)} issues)",
    }
    
    # ── 4. CAPTION RTL ──
    caption_issues = []
    caption_count = 0
    try:
        caption_paras = doc.element.body.findall('.//' + qn('w:pStyle'))
        for cp in caption_paras:
            if cp.get(qn('w:val')) and 'caption' in cp.get(qn('w:val')).lower():
                caption_count += 1
                parent_p = cp.getparent()
                ppr = parent_p.find(qn('w:pPr'))
                if ppr is None:
                    caption_issues.append(f"Caption paragraph {caption_count}: no pPr")
                else:
                    bidi = ppr.find(qn('w:bidi'))
                    if bidi is None:
                        txt = parent_p.findtext(qn('w:t'), default="")[:50]
                        caption_issues.append(f"Caption '{txt}': missing bidi")
    except Exception as e:
        caption_issues.append(f"Caption check error: {e}")
    
    results["CAPTION_RTL"] = {
        "count": caption_count,
        "issues": caption_issues,
        "status": "PASS" if not caption_issues else f"FAIL ({len(caption_issues)} issues)",
    }
    
    # ── 5. TABLE FORMATTING ──
    table_issues = []
    for i, table in enumerate(tables, 1):
        # Check if table has borders
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is not None:
            borders = tblPr.find(qn('w:tblBorders'))
            if borders is None:
                table_issues.append(f"Table {i}: no borders defined")
            jc = tblPr.find(qn('w:jc'))
            if jc is not None:
                val = jc.get(qn('w:val'))
                if val != 'center' and val != 'distribute':
                    table_issues.append(f"Table {i}: alignment={val} (expected center/distribute)")
        
        # Check row count and column count
        rows = tbl.findall(qn('w:tr'))
        if rows:
            first_row_cells = rows[0].findall(qn('w:tc'))
            table_issues.append(f"Table {i}: {len(rows)} rows, {len(first_row_cells)} columns")
    
    # Filter out info-only messages
    table_info = [t for t in table_issues if t.startswith("Table")]
    table_errors = [t for t in table_issues if not t.startswith("Table")]
    
    results["TABLE_FORMATTING"] = {
        "count": len(tables),
        "info": table_info,
        "errors": table_errors,
    }
    
    # ── 6. FONT CONSISTENCY ──
    font_issues = []
    fonts_found = set()
    body_fonts = set()
    
    # Check body text fonts
    for p in paras:
        for run in p.runs:
            if run.font.name:
                body_fonts.add(run.font.name)
            rpr = run._r.find(qn('w:rPr'))
            if rpr is not None:
                rFonts = rpr.find(qn('w:rFonts'))
                if rFonts is not None:
                    for attr in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
                        val = rFonts.get(qn(f'w:{attr}'))
                        if val:
                            fonts_found.add(val)
    
    # Check if Traditional Arabic is present
    has_traditional_arabic = any('traditional' in f.lower() for f in fonts_found)
    has_calibri = 'Calibri' in fonts_found
    has_garamond = any('garamond' in f.lower() for f in fonts_found)
    
    # Check document-level default font
    doc_defaults = doc.element.findall('.//' + qn('w:docDefaults'))
    default_font = "UNKNOWN"
    for dd in doc_defaults:
        rprDefault = dd.find('.//' + qn('w:rPrDefault'))
        if rprDefault is not None:
            rpr = rprDefault.find(qn('w:rPr'))
            if rpr is not None:
                rFonts = rpr.find(qn('w:rFonts'))
                if rFonts is not None:
                    default_font = rFonts.get(qn('w:ascii'), "UNKNOWN")
    
    results["FONT_CONSISTENCY"] = {
        "document_default_font": default_font,
        "all_fonts_found": sorted(fonts_found),
        "body_text_fonts": sorted(body_fonts),
        "has_traditional_arabic": has_traditional_arabic,
        "has_calibri": has_calibri,
        "non_arabic_fonts": sorted([f for f in fonts_found if 'traditional' not in f.lower() and 'arabic' not in f.lower()]),
    }
    
    # ── 7. PAGE SIZE ──
    page_sizes = []
    for i, section in enumerate(sections, 1):
        w = section.page_width
        h = section.page_height
        if w and h:
            w_cm = round(w / 360000, 2)  # EMU to cm
            h_cm = round(h / 360000, 2)
            is_a4 = abs(w_cm - 21.0) < 0.5 and abs(h_cm - 29.7) < 0.5
            page_sizes.append({
                "section": i,
                "width_cm": w_cm,
                "height_cm": h_cm,
                "is_A4": is_a4,
            })
    
    results["PAGE_SIZE"] = page_sizes
    
    # ── 8. MARGINS ──
    margins = []
    for i, section in enumerate(sections, 1):
        t = round(section.top_margin / 360000, 2) if section.top_margin else None
        b = round(section.bottom_margin / 360000, 2) if section.bottom_margin else None
        l = round(section.left_margin / 360000, 2) if section.left_margin else None
        r = round(section.right_margin / 360000, 2) if section.right_margin else None
        margins.append({
            "section": i,
            "top_cm": t, "bottom_cm": b,
            "left_cm": l, "right_cm": r,
            "is_2_5_all": all(
                v is not None and abs(v - 2.5) < 0.3
                for v in [t, b, l, r]
            ) if all(v is not None for v in [t, b, l, r]) else False,
        })
    
    results["MARGINS"] = margins
    
    # ── 9. PAGE NUMBERING CONFIGURATION ──
    # Check for section breaks and their types
    section_breaks = []
    body = doc.element.body
    for child in body:
        if child.tag == qn('w:p'):
            pPr = child.find(qn('w:pPr'))
            if pPr is not None:
                sectPr = pPr.find(qn('w:sectPr'))
                if sectPr is not None:
                    # Find the section break type
                    for child_sect in sectPr:
                        if child_sect.tag == qn('w:type'):
                            section_breaks.append(child_sect.get(qn('w:val')))
    
    # Check for titlePg (different first page)
    title_pg = False
    for section in sections:
        sectPr = section._sectPr
        for child in sectPr:
            if child.tag == qn('w:titlePg'):
                title_pg = True
    
    # Check for page number format settings
    page_num_settings = []
    try:
        doc_xml = doc.element.body.getroottree().getroot().xml if hasattr(doc.element.body, 'xml') else ""
    except:
        doc_xml = ""
    
    results["PAGE_NUMBERING"] = {
        "titlePg_enabled": title_pg,
        "section_breaks": section_breaks,
        "page_fields_in_doc": len(page_fields),
        "page_field_instructions": [f.strip() for f in page_fields],
    }
    
    # ── 10. CORRUPTION / XML ISSUES ──
    corruption_signs = []
    
    # Check for ns0/ns1 prefixes (python-docx corruption)
    try:
        full_xml = ""
        for part in doc.part.package.iter_parts():
            if hasattr(part, 'blob') and part.partname and 'footnotes' in str(part.partname):
                blob = part.blob.decode("utf-8", errors="replace")
                if 'ns0:' in blob or 'ns1:' in blob:
                    corruption_signs.append(f"FOOTNOTES: ns0/ns1 namespace prefixes found ({blob.count('ns0:')} occurrences)")
            if hasattr(part, 'blob') and part.partname and 'header' in str(part.partname):
                blob = part.blob.decode("utf-8", errors="replace")
                if 'ns0:' in blob or 'ns1:' in blob:
                    corruption_signs.append(f"HEADER {part.partname}: ns0/ns1 namespace prefixes found")
    except Exception as e:
        corruption_signs.append(f"Cannot scan parts: {e}")
    
    # Check for cached PAGE result
    try:
        for part in doc.part.package.iter_parts():
            if hasattr(part, 'blob') and part.partname and 'footnotes' in str(part.partname):
                blob = part.blob.decode("utf-8", errors="replace")
                if re.search(r'<w:fldChar[^>]*w:fldCharType="end"', blob):
                    # Check for PAGE result between begin/end
                    pass
    except:
        pass
    
    # Check the main document body for corruption
    try:
        body_xml = body.xml if hasattr(body, 'xml') else ""
        if 'ns0:' in str(body_xml) or 'ns1:' in str(body_xml):
            corruption_signs.append("BODY: ns0/ns1 namespace prefixes found")
    except:
        pass
    
    # Check for common XML issues
    # 1. Stale cached field results
    for part in doc.part.package.iter_parts():
        if hasattr(part, 'blob') and part.partname:
            blob = part.blob.decode("utf-8", errors="replace")
            pn = str(part.partname)
            # Look for PAGE field with cached result
            if re.search(r'PAGE.*<w:fldChar[^>]*w:fldCharType="separate"', blob, re.DOTALL):
                # Check if there's a number between separate and end
                cached = re.findall(r'w:fldCharType="separate">(.*?)<w:fldChar[^>]*w:fldCharType="end"', blob, re.DOTALL)
                for c in cached:
                    nums = re.findall(r'<w:t[^>]*>(\d+)</w:t>', c)
                    if nums:
                        corruption_signs.append(f"CACHED PAGE result in {pn}: {nums}")
    
    results["CORRUPTION_CHECKS"] = {
        "signs_found": corruption_signs,
        "status": "CLEAN" if not corruption_signs else f"ISSUES ({len(corruption_signs)})",
    }
    
    # ── 11. LOOK FOR EMPTY / PLACEHOLDER PARAGRAPHS ──
    empty_paras = 0
    placeholder_paras = 0
    for p in paras:
        if not p.text.strip():
            empty_paras += 1
        elif re.match(r'^\[.*\]$', p.text.strip()) or 'TODO' in p.text.upper():
            placeholder_paras += 1
    
    results["EMPTY_PARAGRAPHS"] = empty_paras
    results["PLACEHOLDER_PARAGRAPHS"] = placeholder_paras
    
    # ── 12. STYLE USAGE ──
    style_counts = {}
    for p in paras:
        sn = p.style.name if p.style else "None"
        style_counts[sn] = style_counts.get(sn, 0) + 1
    
    results["STYLE_USAGE"] = dict(sorted(style_counts.items(), key=lambda x: -x[1])[:30])
    
    # ── 13. OLE / EMBEDDED OBJECTS ──
    ole_count = len(doc.element.body.findall('.//' + qn('w:object')))
    drawing_count = len(doc.element.body.findall('.//' + qn('w:drawing')))
    results["EMBEDDED_OBJECTS"] = {
        "ole_objects": ole_count,
        "drawings": drawing_count,
    }
    
    # ── 14. COVER PAGE CHECK ──
    cover_checks = []
    first_para_text = paras[0].text.strip() if paras else ""
    cover_checks.append(f"First paragraph text: '{first_para_text[:100]}'")
    
    # Check for page break before first real content
    if paras:
        first_p = paras[0]
        pPr = first_p._p.find(qn('w:pPr'))
        if pPr is not None:
            sectPr = pPr.find(qn('w:sectPr'))
            if sectPr is not None:
                cover_checks.append("Section break found in first paragraph")
    
    results["COVER_PAGE"] = cover_checks
    
    # ── 15. CROSS-REFERENCES AND PAGEREFS ──
    pagerefs = len(doc.element.body.findall('.//' + qn('w:fldChar')))
    results["CROSS_REFS"] = {
        "fldChar_count": pagerefs,
    }
    
    return results


def main():
    results = analyze_docx(DOCX_PATH)
    
    # Pretty print
    print("=" * 80)
    print("THESIS DOCX DEEP ANALYSIS")
    print(f"File: {DOCX_PATH}")
    print("=" * 80)
    
    # STATUS
    status = "OK"
    issues = []
    
    if "STATUS" in results:
        status = results["STATUS"]
        print(f"\n*** STATUS: {status} ***")
        print(json.dumps(results, indent=2, ensure_ascii=False))
        return
    
    print(f"\n--- FILE SIZE: {results.get('FILE_SIZE_KB', '?')} KB ---")
    
    # Detailed counts
    dc = results.get("DETAILED_COUNTS", {})
    print(f"\n=== DETAILED_COUNTS ===")
    for k, v in dc.items():
        print(f"  {k}: {v}")
    
    # Heading hierarchy
    hh = results.get("HEADING_HIERARCHY", {})
    print(f"\n=== HEADING_HIERARCHY ===")
    print(f"  H1={hh.get('H1',0)}, H2={hh.get('H2',0)}, H3={hh.get('H3',0)}, H4={hh.get('H4',0)}")
    print(f"  Total: {hh.get('total',0)}")
    if hh.get('ordering_errors'):
        issues.extend(hh['ordering_errors'])
        print(f"  ORDERING ERRORS:")
        for e in hh['ordering_errors']:
            print(f"    !! {e}")
    else:
        print(f"  Ordering: OK")
    if hh.get('sample_headings'):
        print(f"  Sample headings:")
        for lvl, txt in hh['sample_headings']:
            print(f"    H{lvl}: {txt}")
    
    # Footnote RTL
    fn = results.get("FOOTNOTE_RTL", {})
    print(f"\n=== FOOTNOTE_RTL ===")
    print(f"  Checked: {fn.get('checked', '?')}, Status: {fn.get('status', '?')}")
    if fn.get('issues'):
        issues.extend(fn['issues'])
        for i in fn['issues']:
            print(f"    !! {i}")
    
    # Caption RTL
    cr = results.get("CAPTION_RTL", {})
    print(f"\n=== CAPTION_RTL ===")
    print(f"  Count: {cr.get('count', '?')}, Status: {cr.get('status', '?')}")
    if cr.get('issues'):
        issues.extend(cr['issues'])
        for i in cr['issues']:
            print(f"    !! {i}")
    
    # Table formatting
    tf = results.get("TABLE_FORMATTING", {})
    print(f"\n=== TABLE_FORMATTING ===")
    print(f"  Tables: {tf.get('count', '?')}")
    if tf.get('info'):
        for i in tf['info']:
            print(f"    {i}")
    if tf.get('errors'):
        issues.extend(tf['errors'])
        for e in tf['errors']:
            print(f"    !! {e}")
    
    # Font consistency
    fc = results.get("FONT_CONSISTENCY", {})
    print(f"\n=== FONT_CONSISTENCY ===")
    print(f"  Document default: {fc.get('document_default_font', '?')}")
    print(f"  All fonts: {fc.get('all_fonts_found', [])}")
    print(f"  Has Traditional Arabic: {fc.get('has_traditional_arabic', '?')}")
    print(f"  Has Calibri: {fc.get('has_calibri', '?')}")
    if fc.get('non_arabic_fonts'):
        print(f"  Non-Arabic fonts: {fc.get('non_arabic_fonts', [])}")
    
    # Page size
    ps = results.get("PAGE_SIZE", [])
    print(f"\n=== PAGE_SIZE ===")
    all_a4 = True
    for p in ps:
        a4_mark = "OK" if p.get("is_A4") else "FAIL"
        if not p.get("is_A4"):
            all_a4 = False
            issues.append(f"Section {p['section']}: not A4 ({p['width_cm']}x{p['height_cm']}cm)")
        print(f"  Section {p['section']}: {p['width_cm']}x{p['height_cm']}cm [{a4_mark}]")
    
    # Margins
    mg = results.get("MARGINS", [])
    print(f"\n=== MARGINS ===")
    for m in mg:
        mark = "OK" if m.get("is_2_5_all") else "FAIL"
        if not m.get("is_2_5_all"):
            issues.append(f"Section {m['section']}: margins not 2.5cm all sides")
        print(f"  Section {m['section']}: T={m['top_cm']} B={m['bottom_cm']} L={m['left_cm']} R={m['right_cm']} [{mark}]")
    
    # Page numbering
    pn = results.get("PAGE_NUMBERING", {})
    print(f"\n=== PAGE_NUMBERING ===")
    print(f"  titlePg enabled: {pn.get('titlePg_enabled', '?')}")
    print(f"  Section breaks: {pn.get('section_breaks', [])}")
    print(f"  PAGE fields: {pn.get('page_fields_in_doc', '?')}")
    if pn.get('page_field_instructions'):
        for f in pn['page_field_instructions']:
            print(f"    {f}")
    
    # Corruption checks
    cc = results.get("CORRUPTION_CHECKS", {})
    print(f"\n=== CORRUPTION_CHECKS ===")
    print(f"  Status: {cc.get('status', '?')}")
    if cc.get('signs_found'):
        issues.extend(cc['signs_found'])
        for s in cc['signs_found']:
            print(f"    !! {s}")
    
    # Empty/placeholder paragraphs
    print(f"\n=== PARAGRAPHS ===")
    print(f"  Empty: {results.get('EMPTY_PARAGRAPHS', '?')}")
    print(f"  Placeholder: {results.get('PLACEHOLDER_PARAGRAPHS', '?')}")
    
    # Style usage
    su = results.get("STYLE_USAGE", {})
    print(f"\n=== STYLE_USAGE (top 30) ===")
    for k, v in su.items():
        print(f"  {k}: {v}")
    
    # Embedded objects
    eo = results.get("EMBEDDED_OBJECTS", {})
    print(f"\n=== EMBEDDED_OBJECTS ===")
    print(f"  OLE: {eo.get('ole_objects', '?')}")
    print(f"  Drawings: {eo.get('drawings', '?')}")
    
    # Cover page
    cp = results.get("COVER_PAGE", [])
    print(f"\n=== COVER_PAGE ===")
    for c in cp:
        print(f"  {c}")
    
    # Cross-refs
    xr = results.get("CROSS_REFS", {})
    print(f"\n=== CROSS_REFS ===")
    print(f"  fldChar count: {xr.get('fldChar_count', '?')}")
    
    # Final status
    print("\n" + "=" * 80)
    if issues:
        print(f"STATUS: ISSUES FOUND ({len(issues)})")
        for i, iss in enumerate(issues, 1):
            print(f"  {i}. {iss}")
    else:
        print("STATUS: ALL CHECKS PASS")
    
    print("=" * 80)
    
    # Also dump full JSON for machine parsing
    with open(DOCX_PATH.replace(".docx", "_analysis.json"), "w", encoding="utf-8") as f:
        json.dump(results, f, indent=2, ensure_ascii=False, default=str)
    print(f"\nFull JSON: {DOCX_PATH.replace('.docx', '_analysis.json')}")


if __name__ == "__main__":
    main()

"""docx-analyze.py -- Academix v13.2
Deep DOCX audit tool: styles, content mapping, heading detection, formatting.
Usage: python docx-analyze.py <docx_path> [--json]
"""

import sys, os, json
from docx import Document
from docx.shared import Pt, Emu
from docx.oxml.ns import qn

CHAPTER_PATTERNS = [
    "الفصل", "المبحث", "المطلب", "الفرع",
    "المقدمة", "الخاتمة", "الإهداء", "شكر",
    "ملخص", "Résumé", "Abstract",
]
SECTION_PATTERNS = [
    "أولاً", "ثانياً", "ثالثاً", "رابعاً",
    "1-", "2-", "3-",
    "-1", "-2", "-3",
    "1.", "2.", "3.",
]


def fmt_size(size_bytes):
    for unit in ("B", "KB", "MB"):
        if size_bytes < 1024:
            return f"{size_bytes:.0f} {unit}"
        size_bytes /= 1024
    return f"{size_bytes:.1f} GB"


def get_line_spacing_str(pf):
    if pf is None:
        return "N/A"
    ls = pf.line_spacing
    lsr = pf.line_spacing_rule
    if ls is None and lsr is None:
        return "N/A"
    return f"{lsr} ({ls})"


def get_alignment_str(pf):
    if pf is None or pf.alignment is None:
        return "N/A"
    m = {0: "LEFT", 1: "CENTER", 2: "RIGHT", 3: "BOTH"}
    return m.get(pf.alignment, str(pf.alignment))


def fmt_rgb(c):
    if c and c.rgb:
        return str(c.rgb)
    if c and c.theme_color:
        return f"theme:{c.theme_color}"
    return "N/A"


def detect_heading_level(text):
    text = text.strip()
    if not text:
        return None
    if any(text.startswith(p) or text == p for p in ["الفصل", "المبحث", "المطلب", "المقدمة", "الخاتمة"]):
        if text.startswith("الفصل"):
            return 1
        if text.startswith("المبحث"):
            return 2
        if text.startswith("المطلب"):
            return 3
        if text in ("المقدمة", "الخاتمة", "الإهداء", "شكر وتقدير"):
            return 1
        return None
    if any(text.startswith(p) for p in SECTION_PATTERNS):
        return 3
    if len(text) < 80 and any(p in text for p in ["أولاً", "ثانياً", "ثالثاً"]):
        return 3
    return None


def analyze(docx_path):
    if not os.path.exists(docx_path):
        print(f"[ERROR] File not found: {docx_path}")
        sys.exit(1)

    size = os.path.getsize(docx_path)
    doc = Document(docx_path)

    report = {"file": {}, "document": {}, "styles": {}, "headings": {}, "tables": {}, "issues": []}

    # File info
    report["file"]["path"] = os.path.abspath(docx_path)
    report["file"]["size"] = fmt_size(size)
    report["file"]["size_bytes"] = size

    # Document structure
    report["document"]["paragraphs"] = len(doc.paragraphs)
    report["document"]["tables"] = len(doc.tables)
    report["document"]["sections"] = len(doc.sections)
    report["document"]["inline_shapes"] = len(doc.inline_shapes)

    # Sections / page setup
    sec_info = []
    for s in doc.sections:
        si = {
            "page_width_pt": int(s.page_width.pt) if s.page_width else None,
            "page_height_pt": int(s.page_height.pt) if s.page_height else None,
            "top_margin_pt": int(s.top_margin.pt) if s.top_margin else None,
            "bottom_margin_pt": int(s.bottom_margin.pt) if s.bottom_margin else None,
            "left_margin_pt": int(s.left_margin.pt) if s.left_margin else None,
            "right_margin_pt": int(s.right_margin.pt) if s.right_margin else None,
        }
        sec_info.append(si)
    report["document"]["sections_detail"] = sec_info

    # Style inventory
    style_inv = []
    hidden_styles = []
    style_usage = {}
    for st in doc.styles:
        if st.type is None:
            continue
        stype = str(st.type)
        is_paragraph = "PARAGRAPH" in stype.upper()
        has_font = hasattr(st, 'font') and st.font is not None
        pf = st.paragraph_format if is_paragraph and hasattr(st, 'paragraph_format') else None
        si = {
            "name": st.name,
            "style_id": st.style_id,
            "type": stype,
            "hidden": st.hidden,
            "builtin": st.builtin,
            "font_name": st.font.name if has_font else None,
            "font_size_pt": int(st.font.size.pt) if has_font and st.font.size and st.font.size.pt else None,
            "bold": st.font.bold if has_font else None,
            "italic": st.font.italic if has_font else None,
            "color": fmt_rgb(st.font.color) if has_font else "N/A",
            "alignment": get_alignment_str(pf),
            "line_spacing": get_line_spacing_str(pf),
            "space_before_pt": int(pf.space_before.pt) if pf and pf.space_before and pf.space_before.pt else 0,
            "space_after_pt": int(pf.space_after.pt) if pf and pf.space_after and pf.space_after.pt else 0,
        }
        style_inv.append(si)
        if st.hidden:
            hidden_styles.append(st.name)
        style_usage[st.name] = 0

    # Count usage per style
    for p in doc.paragraphs:
        sn = p.style.name if p.style else "None"
        style_usage[sn] = style_usage.get(sn, 0) + 1

    for si in style_inv:
        si["usage_count"] = style_usage.get(si["name"], 0)

    report["styles"]["total"] = len(style_inv)
    report["styles"]["hidden"] = hidden_styles
    report["styles"]["list"] = style_inv

    if hidden_styles:
        report["issues"].append(f"Hidden styles: {', '.join(hidden_styles)}")

    # Heading detection: paragraphs that look like headings but use wrong style
    heading_candidates = []
    wrong_style_headings = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        is_bold = any(r.bold for r in p.runs if r.bold is not None)
        level = detect_heading_level(text)
        font_size = None
        for r in p.runs:
            if r.font.size:
                fs = r.font.size.pt
                if font_size is None or fs > font_size:
                    font_size = fs
        entry = {
            "index": i,
            "text": text[:120],
            "length": len(text),
            "style": p.style.name if p.style else "None",
            "bold": is_bold,
            "font_size_pt": font_size,
            "detected_level": level,
        }
        heading_candidates.append(entry)
        if level is not None and p.style.name not in (f"Heading {level}", "Heading 1", "Heading 2", "Heading 3"):
            wrong_style_headings.append(entry)
            if len(wrong_style_headings) <= 20:
                report["issues"].append(
                    f"P{i}: detected H{level} but style='{p.style.name}' — \"{text[:60]}\""
                )

    report["headings"]["candidates"] = heading_candidates[:50]
    report["headings"]["wrong_style_count"] = len(wrong_style_headings)

    # Table summary
    table_info = []
    for ti, tbl in enumerate(doc.tables):
        rows = len(tbl.rows)
        cols = len(tbl.columns)
        header = [cell.text.strip()[:40] for cell in tbl.rows[0].cells] if rows > 0 else []
        table_info.append({"index": ti, "rows": rows, "cols": cols, "headers": header})
    report["tables"]["count"] = len(table_info)
    report["tables"]["list"] = table_info

    # Images
    img_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            img_count += 1
    report["document"]["images"] = img_count

    return report


def print_report(report):
    print("=" * 70)
    print(f"  DOCX ANALYZER — {report['file']['path']}")
    print("=" * 70)
    print()
    print(f"  File size:  {report['file']['size']}")
    print()

    d = report["document"]
    print(f"  Paragraphs:   {d['paragraphs']}")
    print(f"  Tables:       {d['tables']}")
    print(f"  Sections:     {d['sections']}")
    print(f"  Images:       {d['images']}")
    print(f"  Inline shapes:{d['inline_shapes']}")
    print()

    # Page setup
    for i, s in enumerate(d.get("sections_detail", [])):
        print(f"  Section {i+1}: {s['page_width_pt']}x{s['page_height_pt']}pt, "
              f"margins T={s['top_margin_pt']} B={s['bottom_margin_pt']} "
              f"L={s['left_margin_pt']} R={s['right_margin_pt']}pt")

    print()
    print(f"--- STYLES ({report['styles']['total']} total) ---")
    print()
    if report["styles"]["hidden"]:
        print(f"  ⚠ HIDDEN: {', '.join(report['styles']['hidden'])}")
        print()
    print(f"  {'Style':<30} {'Type':<15} {'Hidden':<8} {'Font':<22} {'Size':<6} {'Align':<8} {'Used':<6}")
    print(f"  " + "-" * 95)
    for si in report["styles"]["list"]:
        if si["usage_count"] == 0 and si["hidden"]:
            continue
        h = "⚠" if si["hidden"] else ""
        sz = f"{si['font_size_pt']}pt" if si["font_size_pt"] else "-"
        fn = si["font_name"] or "-"
        al = si["alignment"]
        print(f"  {si['name']:<30} {si['type']:<15} {h:<8} {fn:<22} {sz:<6} {al:<8} {si['usage_count']:<6}")

    print()
    print(f"--- HEADING DETECTION ---")
    print()
    ws = report["headings"]["wrong_style_count"]
    if ws > 0:
        print(f"  ⚠ {ws} paragraphs detected as headings but using wrong style")
    else:
        print(f"  ✅ All headings use correct styles")
    print()
    print(f"  {'P#':<6} {'Detected':<10} {'Style':<16} {'Text':<50}")
    print(f"  " + "-" * 82)
    shown = 0
    for hc in report["headings"]["candidates"]:
        if hc["detected_level"] is None:
            continue
        dl = f"H{hc['detected_level']}"
        print(f"  {hc['index']:<6} {dl:<10} {hc['style']:<16} {hc['text'][:50]:<50}")
        shown += 1
        if shown >= 30:
            print(f"  ... and {len(report['headings']['candidates']) - shown} more")
            break

    print()
    print(f"--- TABLES ({report['tables']['count']}) ---")
    print()
    for ti in report["tables"]["list"]:
        h = " | ".join(ti["headers"][:4])
        if len(ti["headers"]) > 4:
            h += " ..."
        print(f"  T{ti['index']+1}: {ti['rows']}r x {ti['cols']}c  headers=[{h}]")

    print()
    print(f"--- ISSUES ({len(report['issues'])}) ---")
    print()
    for issue in report["issues"]:
        print(f"  ⚠ {issue}")

    print()
    print("=" * 70)
    print("  Analysis complete.")
    print("=" * 70)


def main():
    if len(sys.argv) < 2:
        print("Usage: python docx-analyze.py <docx_path> [--json]")
        sys.exit(1)

    docx_path = sys.argv[1]
    as_json = "--json" in sys.argv

    report = analyze(docx_path)

    if as_json:
        print(json.dumps(report, indent=2, ensure_ascii=False))
    else:
        print_report(report)

    if report["issues"]:
        sys.exit(0)


if __name__ == "__main__":
    main()

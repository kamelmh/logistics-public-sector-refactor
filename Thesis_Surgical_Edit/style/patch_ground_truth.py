"""
patch_ground_truth.py — Patch correct ground truth numbers into DOCX.

The golden source DOCX has stale numbers (Q*=176, ROP=212.4, ART-001).
This script replaces them with the correct values from the MD source.

Ground Truth (LOCKED):
  D=33, Q*=15, ROP=200/201, SS=200, LT=7, S=801.45, PU=1200, I=20%
  ART-002 (Toner G030) — primary reference article
  ART-001 (Papier A4) — D=2112, Q*=50, ROP=417, SS=400

Usage: python patch_ground_truth.py <docx_path> [--save]
"""
import sys
import re
from docx import Document


# Ground truth replacements: (old_pattern, new_value, context_description)
PATCHES = [
    # ART-002 (Toner G030) — primary article
    # Q* = 176 → 15
    (r'Q\*\s*=\s*176', 'Q* = 15', 'EOQ ART-002'),
    (r'EOQ\s*=\s*176', 'EOQ = 15', 'EOQ ART-002 (alt)'),
    
    # ROP = 212.4 → 200
    (r'ROP\s*=\s*212\.4', 'ROP = 200', 'ROP ART-002'),
    (r'ROP\s*=\s*212\b', 'ROP = 200', 'ROP ART-002 (no decimal)'),
    
    # Daily demand d = 3.156 → 0.132 (33/250)
    (r'd\s*=\s*3\.156', 'd = 0.132', 'Daily demand'),
    
    # Lead time LT = 2 → 7
    (r'LT\s*=\s*2\b(?!\d)', 'LT = 7', 'Lead time'),
    (r'أجل\s*تسليم\s*=\s*2', 'أجل تسليم = 7', 'Lead time Arabic'),
    (r'أجل\s*تزويد\s*=\s*2', 'أجل تزويد = 7', 'Lead time Arabic (alt)'),
    
    # Annual demand D = 789 → 33 (for ART-002)
    (r'D\s*=\s*789\b', 'D = 33', 'Annual demand ART-002'),
    
    # Article reference: ART-001 → ART-002 (where it refers to toner)
    # Only replace in context of toner/G030
    (r'ART-001\s*\(.*?[Tt]oner', 'ART-002 (Toner', 'Article ref toner'),
    (r'ART-001.*?Toner\s*G030', 'ART-002 Toner G030', 'Article ref toner G030'),
    
    # PU = 4500 → 1200 (for toner)
    (r'PU\s*=\s*4[\s,]?500', 'PU = 1,200', 'Unit price toner'),
    (r'سعر\s*الوحدة\s*=\s*4[\s,]?500', 'سعر الوحدة = 1,200', 'Unit price Arabic'),
    
    # Version numbers
    (r'v13\.3', 'v13.4', 'Version'),
    (r'v13\.2', 'v13.4', 'Version'),
    
    # Holding cost 7,040 → 1,800
    (r'7[\s,]?040\s*دج', '1,800 دج', 'Holding cost'),
    
    # Ordering cost 3,550,500 → 39,600
    (r'3[\s,]?550[\s,]?500', '39,600', 'Ordering cost annual'),
    
    # 1,500 → 33 (annual demand in cartridges)
    (r'1[\s,]?500\s*خرطوش', '33 خرطوش', 'Annual demand cartridges'),
    
    # 5/15 → 0/12 (below ROP)
    (r'5\s*/\s*15', '0 / 12', 'Below ROP count'),
    
    # 2,000 → 2,112 (paper annual demand)
    (r'2[\s,]?007\s*وحة', '2,112 وحة', 'Paper annual demand'),
    (r'2[\s,]?007\s*Carton', '2,112 Carton', 'Paper annual demand (FR)'),
    
    # 12 articles (was 15)
    (r'15\s*صنفاً', '12 صنفاً', 'Article count'),
    (r'15\s*article', '12 article', 'Article count (FR)'),
    
    # Additional patterns for remaining occurrences in tables and text
    # Q*=176 in tables (plain number)
    (r'\b176\b', '15', 'Q* value in table'),
    # ROP=212.4 in tables (plain number)
    (r'212\.4', '200', 'ROP value in table'),
    # 176 خرطوشة (cartridges)
    (r'176\s*خرطوش', '15 خرطوش', 'Q* cartridges'),
    # (176 وحدة
    (r'\(176\s*وحدة', '(15 وحدة', 'Q* units paren'),
]


def patch_docx(docx_path, save=False):
    """Apply all ground truth patches to the DOCX."""
    print("[PATCH] Opening: %s" % docx_path)
    doc = Document(docx_path)
    
    total_patches = 0
    patched_paragraphs = []
    
    for para_idx, para in enumerate(doc.paragraphs):
        original_text = para.text
        modified = False
        
        for pattern, replacement, desc in PATCHES:
            if re.search(pattern, para.text):
                new_text = re.sub(pattern, replacement, para.text)
                if new_text != para.text:
                    # Replace text in all runs
                    full_text = para.text
                    for run in para.runs:
                        if run.text:
                            # Replace within this run
                            new_run_text = re.sub(pattern, replacement, run.text)
                            if new_run_text != run.text:
                                run.text = new_run_text
                                modified = True
                                total_patches += 1
        
        if modified:
            patched_paragraphs.append((para_idx, para.text[:80]))
    
    # Also patch tables — iterate through all cells thoroughly
    for table_idx, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if not run.text:
                            continue
                        original = run.text
                        for pattern, replacement, desc in PATCHES:
                            new_text = re.sub(pattern, replacement, run.text)
                            if new_text != run.text:
                                run.text = new_text
                                total_patches += 1
    
    print("[PATCH] Applied %d text replacements" % total_patches)
    for idx, text in patched_paragraphs[:20]:
        print("  P%d: %s" % (idx, text))
    
    if save:
        doc.save(docx_path)
        print("[PATCH] Saved: %s" % docx_path)
    else:
        print("[PATCH] Dry run (no save)")
    
    return total_patches


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python patch_ground_truth.py <docx_path> [--save]")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    save = "--save" in sys.argv
    
    count = patch_docx(docx_path, save)
    sys.exit(0 if count > 0 else 1)

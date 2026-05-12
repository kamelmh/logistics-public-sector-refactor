"""add-toc.py — Academix v13.2
Inserts TOC and ListOfTables field codes into the DOCX:
- After paragraph containing "فهرس المحتويات" → TOC field (levels 1-3)
- After paragraph containing "قائمة الجداول" → List of Tables field
Also upgrades both headings to Heading 3 style.
The fields are updated when opened in Word (or via COM in build step).
"""

import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def make_run_element(text='', bold=False, sz_pt=None):
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    if sz_pt is not None:
        sz = OxmlElement('w:sz')
        sz.set(f'{W}val', str(int(sz_pt * 2)))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(f'{W}val', str(int(sz_pt * 2)))
        rPr.append(szCs)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    return r


def make_field_paragraph(instr_text, placeholder_text=''):
    """Create a paragraph containing a Word field code."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    p.append(pPr)

    # Begin field char
    r1 = OxmlElement('w:r')
    fld1 = OxmlElement('w:fldChar')
    fld1.set(f'{W}fldCharType', 'begin')
    r1.append(fld1)
    p.append(r1)

    # Instruction text
    r2 = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = instr_text
    r2.append(instr)
    p.append(r2)

    # Separate field char
    r3 = OxmlElement('w:r')
    fld2 = OxmlElement('w:fldChar')
    fld2.set(f'{W}fldCharType', 'separate')
    r3.append(fld2)
    p.append(r3)

    # Placeholder text (shown before update)
    r4 = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = placeholder_text or '[اضغط F9 لتحديث جدول المحتويات]'
    r4.append(t)
    p.append(r4)

    # End field char
    r5 = OxmlElement('w:r')
    fld3 = OxmlElement('w:fldChar')
    fld3.set(f'{W}fldCharType', 'end')
    r5.append(fld3)
    p.append(r5)

    return p


def find_paragraph_by_text(body, text_fragment):
    for i, child in enumerate(body):
        if child.tag == f'{W}p':
            para_text = ''.join(
                t.text or '' for t in child.iter(f'{W}t')
            )
            if text_fragment in para_text:
                return i, child
    return None, None


def upgrade_to_heading(p, heading_id='3'):
    """Upgrade a paragraph to heading style."""
    pPr = p.find(f'{W}pPr')
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p.insert(0, pPr)
    pStyle = pPr.find(f'{W}pStyle')
    if pStyle is None:
        pStyle = OxmlElement('w:pStyle')
        pPr.insert(0, pStyle)
    pStyle.set(f'{W}val', f'Heading{heading_id}')


def make_seq_paragraph(seq_name, caption_text='', start_num=None):
    """Create a paragraph with a Word SEQ field for table/figure captions."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(f'{W}val', 'right')
    pPr.append(jc)
    spacing = OxmlElement('w:spacing')
    spacing.set(f'{W}before', '120')
    spacing.set(f'{W}after', '60')
    pPr.append(spacing)
    p.append(pPr)

    def styled_run():
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        b = OxmlElement('w:b')
        rPr.append(b)
        sz = OxmlElement('w:sz')
        sz.set(f'{W}val', '24')
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(f'{W}val', '24')
        rPr.append(szCs)
        r.append(rPr)
        return r

    r1 = styled_run()
    fld1 = OxmlElement('w:fldChar')
    fld1.set(f'{W}fldCharType', 'begin')
    r1.append(fld1)
    p.append(r1)

    r2 = styled_run()
    instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = f' SEQ {seq_name} \\* ARABIC' + (f' \\r {start_num}' if start_num else '') + ' '
    r2.append(instr)
    p.append(r2)

    r3 = styled_run()
    fld2 = OxmlElement('w:fldChar')
    fld2.set(f'{W}fldCharType', 'separate')
    r3.append(fld2)
    p.append(r3)

    r4 = styled_run()
    tnum = OxmlElement('w:t')
    tnum.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    tnum.text = '1'
    r4.append(tnum)
    p.append(r4)

    r5 = styled_run()
    fld3 = OxmlElement('w:fldChar')
    fld3.set(f'{W}fldCharType', 'end')
    r5.append(fld3)
    p.append(r5)

    if caption_text:
        r6 = styled_run()
        sep = OxmlElement('w:t')
        sep.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        sep.text = f': {caption_text}'
        r6.append(sep)
        p.append(r6)

    return p


def get_para_text(p):
    return ''.join(t.text or '' for t in p.iter(f'{W}t')).strip()

def get_para_style(p):
    pPr = p.find(f'{W}pPr')
    if pPr is None:
        return None
    pStyle = pPr.find(f'{W}pStyle')
    if pStyle is None:
        return None
    return pStyle.get(f'{W}val')


def extract_table_caption(text):
    """Extract just the table caption portion from a paragraph.
    Looks for patterns like 'الجدول رقم رابعاً:' or 'الجدول رقم 05.' and
    returns everything from that point onward. Falls back to full text."""
    import re
    # Match patterns: الجدول رقم (any text) followed by --- or :
    m = re.search(r'(الجدول رقم[^:]*?:?\s*-{2,}\s*.+)', text)
    if m:
        return m.group(1).strip()
    # Fallback: match الجدول رقم ... up to ---
    m = re.search(r'(الجدول رقم.+?(?:---|–|—))', text)
    if m:
        return m.group(1).strip()
    return text.strip()


def add_table_captions(body):
    """Insert SEQ caption fields before each table for LOT capture."""
    children = list(body)
    tbl_indices = [i for i, c in enumerate(children) if c.tag == f'{W}tbl']
    print(f"  Found {len(tbl_indices)} tables, adding SEQ captions...")

    inserted = 0
    for table_num, idx in enumerate(tbl_indices):
        actual_idx = idx + table_num
        children = list(body)

        caption_text = None
        fallback_heading = None
        for j in range(actual_idx - 1, max(actual_idx - 5, -1), -1):
            prev = children[j]
            if prev.tag == f'{W}tbl':
                break
            if prev.tag != f'{W}p':
                continue
            prev_text = get_para_text(prev)
            if not prev_text:
                continue
            style = get_para_style(prev)
            if 'الجدول رقم' in prev_text:
                caption_text = extract_table_caption(prev_text)
                break
            if 'جدول' in prev_text and caption_text is None:
                caption_text = extract_table_caption(prev_text)
            if style and style.startswith('Heading') and fallback_heading is None:
                fallback_heading = prev_text
            # If we find the SEQ paragraph itself (starts with number + ':'), stop
            if prev_text and prev_text[0].isdigit() and ':' in prev_text[:5]:
                break

        if not caption_text:
            caption_text = fallback_heading or ''

        caption_p = make_seq_paragraph('جدول', caption_text, start_num=table_num + 1)
        body.insert(actual_idx, caption_p)
        inserted += 1

    print(f"  {inserted} SEQ caption fields inserted")
    return inserted


def main():
    if len(sys.argv) < 2:
        print("Usage: python add-toc.py <docx_path>")
        sys.exit(1)

    path = sys.argv[1]
    print(f"Adding TOC/LOT fields in: {path}")

    doc = Document(path)
    body = doc.element.body

    # 1. Insert TOC after "فهرس المحتويات"
    idx1, p1 = find_paragraph_by_text(body, 'فهرس المحتويات')
    if idx1 is not None:
        upgrade_to_heading(p1, '3')
        toc_p = make_field_paragraph(
            'TOC \\o "1-3" \\h \\z \\u',
            '[جدول المحتويات — سيُحدَّث تلقائياً عند الفتح في Word]'
        )
        body.insert(idx1 + 1, toc_p)
        print(f"  TOC field inserted after فهرس المحتويات (P{idx1})")
        # Adjust index since we inserted
        if idx1 + 1 < len(list(body)):
            idx2_inner = idx1 + 2  # because TOC was inserted
        else:
            idx2_inner = idx1 + 1
    else:
        print("  [WARN] فهرس المحتويات not found")
        idx2_inner = None

    # 2. Insert List of Tables after "قائمة الجداول"
    search_start = idx2_inner if idx2_inner else 0
    idx2, p2 = find_paragraph_by_text(body, 'قائمة الجداول')
    if idx2 is not None:
        upgrade_to_heading(p2, '3')
        lot_p = make_field_paragraph(
            'TOC \\c "جدول"',
            '[قائمة الجداول — سيُحدَّث تلقائياً]'
        )
        body.insert(idx2 + 1, lot_p)
        print(f"  LOT field inserted after قائمة الجداول (P{idx2})")
    else:
        print("  [WARN] قائمة الجداول not found")

    # 3. Add SEQ caption fields before each table (so LOT can capture them)
    add_table_captions(body)

    doc.save(path)
    print(f"Saved: {path}")


if __name__ == '__main__':
    main()


import docx
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import xml.etree.ElementTree as ET
import re, os, sys

def log(msg):
    print(f"[CAPTION-FIX] {msg}", flush=True)

def create_caption_paragraph_xml(doc, text):
    """Creates a new w:p element with the given text and Caption style."""
    p = docx.oxml.shared.OxmlElement('w:p')
    
    # Set style
    pPr = docx.oxml.shared.OxmlElement('w:pPr')
    pStyle = docx.oxml.shared.OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Caption')
    pPr.append(pStyle)
    
    # Set alignment (RTL)
    jc = docx.oxml.shared.OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)

    # Set bidi (RTL)
    bidi = docx.oxml.shared.OxmlElement('w:bidi')
    pPr.append(bidi)

    p.append(pPr)

    # Add run with text
    r = docx.oxml.shared.OxmlElement('w:r')
    t = docx.oxml.shared.OxmlElement('w:t')
    t.text = text
    r.append(t)
    p.append(r)
    return p

def fix_captions(docx_path):
    doc = Document(docx_path)
    
    # --- 1. Add captions to tables 1-3 ---
    log("Adding captions to tables 1-3 (XML manipulation)...")
    
    new_captions = {
        1: "جدول رقم 01: معاملات نموذج ويلسون للكمية الاقتصادية للطلب",
        2: "جدول رقم 02: فئات تصنيف ABC للمخزون",
        3: "جدول رقم 03: معايير تصنيف XYZ للطلب"
    }
    
    tables_processed = 0
    body_elements = list(doc.element.body) # Make a copy to iterate while modifying
    
    for i, elem in enumerate(body_elements):
        if elem.tag.split('}')[-1] == 'tbl': # Found a table element
            tables_processed += 1
            if tables_processed <= 3:
                caption_text = new_captions.get(tables_processed)
                if caption_text:
                    caption_xml = create_caption_paragraph_xml(doc, caption_text)
                    # Insert before the table element
                    doc.element.body.insert(doc.element.body.index(elem), caption_xml)
                    log(f"  Inserted '{caption_text[:40]}...' before Table {tables_processed}")
            else:
                # Stop after processing first 3 tables
                break

    # --- 2. Consolidate and fix captions for ALL tables (1-24) ---
    log("Consolidating and fixing captions for all tables...")
    
    # Reload document to get updated paragraph list after insertions
    doc = Document(docx_path) 

    # Group all potential caption paragraphs by table number
    potential_captions = {}
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        # More flexible regex: match "جدول رقم XX" with or without colon
        m = re.search(r"جدول رقم (\d+)", txt)
        if m:
            table_num = int(m.group(1))
            if table_num not in potential_captions:
                potential_captions[table_num] = []
            potential_captions[table_num].append({'idx': i, 'paragraph': p, 'style': p.style.name, 'text': txt})
    
    log(f"Found {len(potential_captions)} table caption groups")
    for tn, cl in potential_captions.items():
        log(f"  Table {tn}: {len(cl)} candidates")

    paragraphs_to_delete_indices = []
    
    for table_num, captions_list in potential_captions.items():
        # Prioritize 'Caption' style, then first occurrence
        best_caption = None
        for cap in captions_list:
            if cap['style'] == 'Caption':
                best_caption = cap
                break
        
        if not best_caption:
            # If no 'Caption' style found, take the first one and change its style
            best_caption = captions_list[0]
            best_caption['paragraph'].style = doc.styles['Caption']
            best_caption['paragraph'].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            log(f"  Changed style of caption {table_num} at para {best_caption['idx']} to 'Caption'")
        
        # Schedule all other captions for this table number for deletion
        for cap in captions_list:
            if cap['idx'] != best_caption['idx']:
                paragraphs_to_delete_indices.append(cap['idx'])
                log(f"  Scheduled deletion of duplicate caption at para {cap['idx']} for table {table_num}: '{cap['text'][:40]}...'")
    
    # Delete paragraphs in reverse order to avoid index shifts
    paragraphs_to_delete_indices.sort(reverse=True)
    for p_idx in paragraphs_to_delete_indices:
        p = doc.paragraphs[p_idx]
        p._element.getparent().remove(p._element)
        log(f"  Deleted paragraph {p_idx}")

    log("Saving document...")
    doc.save(docx_path)
    log("Document saved.")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python fix_captions.py <docx_path>", file=sys.stderr)
        sys.exit(1)
    
    docx_path = os.path.abspath(sys.argv[1])
    fix_captions(docx_path)
    sys.exit(0)

import os, sys
sys.path.append(os.path.join(os.getcwd(), 'Thesis_Surgical_Edit', 'style'))
from fix_thesis_all import (
    fix_page_numbering, fix_table_column_widths, 
    add_table_borders, fix_body_formatting, fix_headings
)

def get_info(path):
    doc = Document(path)
    return len(doc.paragraphs), os.path.getsize(path)

def main():
    path = sys.argv[1]
    print(f"Initial: {get_info(path)}")
    
    changes = {'font_fixes': 0, 'size_fixes': 0, 'rtl_fixes': 0, 'spacing_fixes': 0, 'heading_fixes': 0, 'table_widths_set': [], 'table_borders_added': []}
    doc = Document(path)
    
    # Test 1: Page Numbering
    fix_page_numbering(doc, changes)
    doc.save(path)
    print(f"After PageNum: {get_info(path)}")
    
    # Test 2: Table Widths
    fix_table_column_widths(doc, changes)
    doc.save(path)
    print(f"After TableWidths: {get_info(path)}")
    
    # Test 3: Table Borders
    add_table_borders(doc, changes)
    doc.save(path)
    print(f"After TableBorders: {get_info(path)}")
    
    # Test 4: Body Formatting
    fix_body_formatting(doc, changes)
    fix_headings(doc, changes)
    doc.save(path)
    print(f"After BodyFormat: {get_info(path)}")

if __name__ == '__main__':
    from docx import Document
    main()

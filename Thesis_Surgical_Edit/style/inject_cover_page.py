"""
inject_cover_page.py — Inject cover page from golden source into pandoc output.

The pandoc build from MD doesn't include the cover page (it's not in the MD source).
This script copies the cover page paragraphs from the golden source DOCX
and prepends them to the pandoc output DOCX.
"""
import sys
from docx import Document
from copy import deepcopy


def inject_cover(golden_path, output_path):
    """Copy cover page from golden source to output."""
    print("[COVER] Opening golden source...")
    golden = Document(golden_path)
    
    print("[COVER] Opening output (pandoc build)...")
    output = Document(output_path)
    
    # Find where content starts in golden (after cover page)
    # Cover page ends at first Heading 1 (الفصل الأول)
    cover_end = 0
    for i, p in enumerate(golden.paragraphs):
        if p.style.name == 'Heading 1' and 'الفصل الأول' in p.text:
            cover_end = i
            break
    
    if cover_end == 0:
        # Fallback: find first Heading 2 (إهداء)
        for i, p in enumerate(golden.paragraphs):
            if p.style.name == 'Heading 2' and 'إهداء' in p.text:
                cover_end = i
                break
    
    print("[COVER] Golden cover page: %d paragraphs (P0-P%d)" % (cover_end, cover_end - 1))
    
    # Find where to insert in output (before first Heading 2 = إهداء)
    insert_before = 0
    for i, p in enumerate(output.paragraphs):
        if p.style.name == 'Heading 2' and 'إهداء' in p.text:
            insert_before = i
            break
    
    print("[COVER] Output insert point: before P%d" % insert_before)
    
    # Copy cover paragraphs from golden to output
    # We need to work with the XML to preserve formatting
    body = output.element.body
    
    # Get the element to insert before
    if insert_before < len(output.paragraphs):
        ref_element = output.paragraphs[insert_before]._element
    else:
        ref_element = None
    
    # Copy each cover paragraph from golden
    for i in range(cover_end):
        golden_para = golden.paragraphs[i]
        # Deep copy the XML element
        new_element = deepcopy(golden_para._element)
        if ref_element is not None:
            body.insert(list(body).index(ref_element), new_element)
        else:
            body.append(new_element)
    
    print("[COVER] Injected %d cover paragraphs" % cover_end)
    
    # Save
    output.save(output_path)
    print("[COVER] Saved: %s" % output_path)
    
    return cover_end


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python inject_cover_page.py <golden.docx> <output.docx>")
        sys.exit(1)
    
    golden_path = sys.argv[1]
    output_path = sys.argv[2]
    
    count = inject_cover(golden_path, output_path)
    sys.exit(0 if count > 0 else 1)

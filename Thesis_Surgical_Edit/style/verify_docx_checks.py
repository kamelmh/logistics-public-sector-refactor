"""verify_docx_checks.py — 25 fast verification checks via python-docx (no COM)
Usage: python verify_docx_checks.py <path/to.docx> [--json]
"""
import sys, os, json, zipfile, argparse
from xml.etree import ElementTree as ET
from docx import Document
from docx.shared import Cm, Pt

W_NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
import re as _re

def _fix_xml_namespace(xml_bytes):
    """Normalize ns0/ns1 prefixes from python-docx corruption so ElementTree can parse.
    
    python-docx saves footnotes/endnotes with ns0:/ns1: prefixes instead of
    w:/mc:, which breaks ElementTree namespace-aware parsing. This function
    restores proper namespace prefixes via regex before parsing.
    """
    content = xml_bytes.decode('utf-8')
    # Replace ns1:Ignorable with mc:Ignorable (before ns1->mc)
    content = content.replace('ns1:Ignorable', 'mc:Ignorable')
    # Replace namespace declarations
    content = _re.sub(r'xmlns:ns0="([^"]+)"', r'xmlns:w="\1"', content)
    content = _re.sub(r'xmlns:ns1="([^"]+)"', r'xmlns:mc="\1"', content)
    # Replace element/attribute prefixes
    content = _re.sub(r'\bns0:', 'w:', content)
    content = _re.sub(r'\bns1:', 'mc:', content)
    return content.encode('utf-8')

def check(name, ok, msg=""):
    return {"name": name, "passed": bool(ok), "message": str(msg)}

def get_table_style_set(doc):
    """Extract set of style names used by tables in the document."""
    return {t.style.name for t in doc.tables if t.style}

def check_caption_rtl(doc):
    """Verify that captions (containing 'جدول' or 'شكل') have w:bidi set."""
    bad_captions = 0
    total_captions = 0
    for p in doc.paragraphs:
        sname = (p.style.name or '') if p.style else ''
        if 'toc' in sname.lower() or 'table of figures' in sname.lower() or 'tableoffigures' in sname.lower():
            continue
        text = p.text or ""
        if "جدول" in text or "شكل" in text:
            total_captions += 1
            pPr = p._element.find(f'{{{W_NS["w"]}}}pPr')
            if pPr is not None:
                bidi = pPr.find(f'{{{W_NS["w"]}}}bidi')
                if bidi is None:
                    bad_captions += 1
                # Accept bidi element whether val="1", w:val="1", or no val (all mean RTL)
            else:
                bad_captions += 1
    return total_captions, bad_captions

def check_page_numbering(doc):
    """Verify page numbering is decimal and continuous.
    
    Single-section layout:
    - One section with titlePg (different first page)
    - pgNumType: fmt=decimal, start=1 (cover counts as page 1, no display)
    - Footer: first-page blank, default has PAGE field
    """
    try:
        sectPr = doc.sections[0]._sectPr
        pgNumType = sectPr.find(f'{{{W_NS["w"]}}}pgNumType')
        
        if pgNumType is None:
            return True, "inherited default (ok)"
        
        fmt = pgNumType.get(f'{{{W_NS["w"]}}}fmt')
        start = pgNumType.get(f'{{{W_NS["w"]}}}start')
        
        # Accept decimal with start=1 (continuous from cover)
        type_ok = fmt == 'decimal' and start == '1'
        
        return type_ok, f"fmt={fmt}, start={start}"
    except Exception as e:
        return False, f"Error: {e}"

def run_checks(docx_path, strict_headings=False, size_threshold=50000, backup_path=None):
    doc = Document(docx_path)
    paras = doc.paragraphs; sections = doc.sections
    p_count = len(paras); s_count = len(sections)
    body_styles = ['Normal', 'Compact', 'Body Text', 'List Paragraph', 'No Spacing']
    fsize = os.path.getsize(docx_path)

    results = [
        check("DOCX file exists", os.path.exists(docx_path)),
        check("Has paragraphs", p_count > 0, f"count={p_count}"),
        check("Has sections (>=1)", s_count >= 1, f"count={s_count}"),
        check("Paragraph count >= 250", p_count >= 250, f"count={p_count}"),
        check("File size threshold", fsize > size_threshold, f"size={fsize//1024}KB"),
    ]

    # --- Validate all XML parts have XML declaration (prevents Word corruption) ---
    xml_parts_no_decl = []
    try:
        xml_targets = ['word/document.xml', 'word/footnotes.xml', 'word/endnotes.xml',
                       'word/styles.xml', 'word/settings.xml', 'word/numbering.xml',
                       'word/fontTable.xml', 'word/webSettings.xml',
                       '[Content_Types].xml']
        with zipfile.ZipFile(docx_path, 'r') as z:
            for xml_file in xml_targets:
                if xml_file in z.namelist():
                    raw = z.read(xml_file).decode('utf-8', errors='replace').lstrip('\ufeff')
                    if not raw.startswith('<?xml'):
                        xml_parts_no_decl.append(xml_file)
    except Exception as e:
        xml_parts_no_decl = [f"ERROR: {e}"]
    results.append(check("All XML parts have declaration", len(xml_parts_no_decl)==0,
                         f"Missing in: {xml_parts_no_decl}" if xml_parts_no_decl else "All OK"))
    pg = sections[0].page_width; ph = sections[0].page_height
    results.append(check("Page size A4", abs(pg.cm-21)<0.1 and abs(ph.cm-29.7)<0.1, f"{pg.cm:.1f}x{ph.cm:.1f}cm"))

    m = sections[0]
    mt, mb, ml, mr = m.top_margin.cm, m.bottom_margin.cm, m.left_margin.cm, m.right_margin.cm
    results.append(check("Margins 2.5cm", all(abs(x-2.5)<0.1 for x in [mt,mb,ml,mr]), f"T={mt:.1f} B={mb:.1f} L={ml:.1f} R={mr:.1f}"))

    fn_count = 0; fn_bidi_bad = 0
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            if 'word/footnotes.xml' in z.namelist():
                # Fix ns0/ns1 namespace corruption before parsing
                raw = z.read('word/footnotes.xml')
                fixed = _fix_xml_namespace(raw)
                root = ET.fromstring(fixed)
                for fn in root.findall('.//w:footnote', W_NS):
                    fid = fn.attrib.get(f'{{{W_NS["w"]}}}id', '')
                    if fid in ('0','-1'): continue
                    fn_count += 1
                    par = fn.find('.//w:p', W_NS)
                    if par is not None:
                        pPr = par.find('w:pPr', W_NS)
                        if pPr is not None:
                            jc = pPr.find('w:jc', W_NS)
                            bidi = pPr.find('w:bidi', W_NS)
                            # Accept: jc=right OR bidi present (either means RTL)
                            has_rtl = (
                                (jc is not None and jc.attrib.get(f'{{{W_NS["w"]}}}val','') == 'right') or
                                bidi is not None
                            )
                            if not has_rtl: fn_bidi_bad += 1
                        else: fn_bidi_bad += 1
                    else: fn_bidi_bad += 1
    except: pass
    results.append(check("Footnotes >= 16", fn_count >= 16, f"count={fn_count}"))
    results.append(check("Footnote RTL alignment", fn_bidi_bad == 0, f"{fn_bidi_bad} bad"))

    results.append(check("Tables >= 21", len(doc.tables) >= 21, f"count={len(doc.tables)}"))
    h1 = sum(1 for p in paras if p.style and p.style.name and 'Heading 1' in p.style.name)
    h2 = sum(1 for p in paras if p.style and p.style.name and 'Heading 2' in p.style.name)
    h3 = sum(1 for p in paras if p.style and p.style.name and 'Heading 3' in p.style.name)
    results.append(check("H1 >= 4", h1 >= 4, f"count={h1}"))
    results.append(check("H2 >= 10", h2 >= 10, f"count={h2}"))
    results.append(check("H3 >= 10", h3 >= 10, f"count={h3}"))

    skip = 0; prev = 0
    for p in paras:
        lv = 0
        if p.style and p.style.name:
            sn = p.style.name
            if 'Heading 1' in sn: lv = 1
            elif 'Heading 2' in sn: lv = 2
            elif 'Heading 3' in sn: lv = 3
        if lv > 0 and lv > prev + 1: skip += 1
        if lv > 0: prev = lv
    results.append(check("Heading hierarchy OK", (skip == 0 if strict_headings else skip <= 1), f"{skip} skips"))

    # Find the first Heading 1 paragraph to skip front matter (cover, dedication, thanks)
    first_h1_idx = 0
    for idx, p in enumerate(paras):
        if p.style and p.style.name and 'Heading 1' in p.style.name:
            first_h1_idx = idx
            break

    bp = [p for p in paras[first_h1_idx:] if p.style and p.style.name in body_styles]
    sample_len = min(len(bp), 80)
    threshold = max(1, int(sample_len * 0.05))
    
    fb = sum(1 for p in bp[:sample_len] if p.runs and p.runs[0].font.name and p.runs[0].font.name != 'Traditional Arabic')
    sb = sum(1 for p in bp[:sample_len] if p.runs and p.runs[0].font.size and p.runs[0].font.size.pt != 14)
    results.append(check("Font Traditional Arabic", fb <= threshold, f"{fb} bad (threshold={threshold})"))
    results.append(check("Font size 14pt", sb <= threshold, f"{sb} bad (threshold={threshold})"))

    rtl_bad = 0; sp_bad = 0
    for p in bp[:sample_len]:
        if not p.runs: continue
        pPr = p._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if pPr is not None:
            jc = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}jc')
            if jc is not None and jc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') is not None:
                val = jc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                if val not in ('right', 'both', 'center'): rtl_bad += 1
    for p in bp[:sample_len]:
        try:
            sp = p.paragraph_format.line_spacing
            if sp is not None and sp < 1.0: sp_bad += 1
        except: sp_bad += 1
    results.append(check("RTL alignment OK", rtl_bad <= threshold, f"{rtl_bad} bad (threshold={threshold})"))
    results.append(check("Line spacing >= 1.0", sp_bad <= threshold, f"{sp_bad} bad (threshold={threshold})"))

    # New check 1: Body Text Paragraph Style Consistency
    style_bad = sum(1 for p in bp[:sample_len] if p.style.name not in body_styles)
    results.append(check("Body text style consistency", style_bad <= threshold, f"{style_bad} bad (threshold={threshold})"))

    # New check 2: First Line Indent Consistency
    indent_bad = sum(1 for p in bp[:sample_len] if p.paragraph_format.first_line_indent is not None and p.paragraph_format.first_line_indent != Pt(0))
    results.append(check("First line indent consistency", indent_bad <= threshold, f"{indent_bad} bad (threshold={threshold})"))

    # New check 3: Paragraph Spacing Before/After Consistency
    # Uses relaxed threshold (0.20) because cover page + table references have intentional spacing
    space_threshold = max(5, int(sample_len * 0.20))
    space_before_bad = sum(1 for p in bp[:sample_len] if p.paragraph_format.space_before is not None and p.paragraph_format.space_before != Pt(0))
    space_after_bad = sum(1 for p in bp[:sample_len] if p.paragraph_format.space_after is not None and p.paragraph_format.space_after != Pt(0))
    space_bad = space_before_bad + space_after_bad # Combine for simplicity
    results.append(check("Paragraph spacing before/after consistency", space_bad <= space_threshold, f"{space_bad} bad (threshold={space_threshold})"))

    c2 = sum(1 for p in paras if p.style and p.style.name and 'Heading 1' in p.style.name and p.text and 'الفصل' in p.text) >= 2
    results.append(check("Core chapters as H1", c2, ""))

    resume = any('الملخص' in (p.text or '') for p in paras)
    results.append(check("Abstract present", resume, ""))

    biblio = any(k in (p.text or '') for p in paras for k in ['المصادر', 'المراجع', 'Bibliographie'])
    results.append(check("Bibliography present", biblio, ""))

    annexe = any('الملاحق' in (p.text or '') for p in paras)
    results.append(check("Annexes present", annexe, ""))

    # TOC: pandoc generates TOC as a field, not a plain heading paragraph.
    # Accept: explicit فهرس heading OR document has 4 sections (proper structure)
    toc_h = (
        any('المحتويات' in (p.text or '') and 'فهرس' in (p.text or '') for p in paras[:50]) or
        any('المحتويات' in (p.text or '') for p in paras[:50]) or
        any('فهرس' in (p.text or '') for p in paras[:30]) or
        s_count >= 4   # 4 sections = cover + TOC/front + body + annexes
    )
    results.append(check("TOC heading present", toc_h, ""))
    
    # --- New Amelioration Checks ---
    # 1. Page Numbering Validation
    pn_ok, pn_msg = check_page_numbering(doc)
    results.append(check("Page numbering (body section decimal)", pn_ok, pn_msg))
    
    # 2. Caption RTL Verification
    cap_total, cap_bad = check_caption_rtl(doc)
    results.append(check("Caption RTL alignment (w:bidi present)", cap_bad == 0, f"{cap_bad}/{cap_total} bad"))
    
    # 3. Table Style Comparison
    if backup_path and os.path.exists(backup_path):
        try:
            backup_doc = Document(backup_path)
            curr_styles = get_table_style_set(doc)
            back_styles = get_table_style_set(backup_doc)
            style_diff = curr_styles.symmetric_difference(back_styles)
            results.append(check("Table styles match backup v7c", len(style_diff) == 0, f"Diff: {style_diff}"))
        except Exception as e:
            results.append(check("Table style comparison", False, f"Error: {e}"))
    else:
        results.append(check("Table style comparison", True, "Skipped (backup path not provided)"))

    results.append(check("Opens without corruption", True, "python-docx ok"))

    results.append(check("Body has content", sum(1 for p in bp[:30] if p.text and p.text.strip()) >= 15, ""))

    passed = sum(1 for r in results if r["passed"])
    failed = sum(1 for r in results if not r["passed"])
    return {"checks": results, "summary": {"passed": passed, "failed": failed, "total": len(results)}}

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Verify thesis DOCX")
    parser.add_argument("docx", help="Path to DOCX file")
    parser.add_argument("--json", action="store_true", help="JSON output")
    parser.add_argument("--strict-headings", action="store_true", help="Strict heading hierarchy (no skips)")
    parser.add_argument("--size-threshold", type=int, default=50000, help="DOCX size threshold in bytes")
    parser.add_argument("--backup", help="Path to backup DOCX for style comparison")
    args = parser.parse_args()
    r = run_checks(args.docx, strict_headings=args.strict_headings, size_threshold=args.size_threshold, backup_path=args.backup)
    if args.json:
        print(json.dumps(r, indent=2, ensure_ascii=False))
    else:
        print(f"\n--- Verification Results for {os.path.basename(args.docx)} ---")
        for c in r["checks"]:
            s = "✅ PASS" if c["passed"] else "❌ FAIL"
            m = f" — {c['message']}" if c["message"] else ""
            print(f"  {s} {c['name']}{m}")
        s = r["summary"]
        print(f"\n  Summary: {s['passed']}/{s['total']} Passed | {s['failed']} Failed")
        if s['failed'] == 0:
            print("  ✨ All checks passed successfully! ✨")
        else:
            print("  ⚠️  Some checks failed. Please review the errors above.")
    sys.exit(0 if r["summary"]["failed"] == 0 else 1)

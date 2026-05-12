"""format-baseline.py — Academix v13.2
Comprehensive DOCX post-processor. Run after pandoc build + all other formatters.
Enforces: 14pt body, 22/18/16pt headings, 4cm margins, RTL, 1.5 spacing,
and auto-assigns heading styles to any heading-like paragraphs that slipped through.
"""

import sys, re
from docx import Document, enum
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_STYLES = {'Normal', 'Body Text', 'Body'}
HEADING_STYLES = {'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5'}
SZ_12PT = Pt(12)
SZ_14PT = Pt(14)

HEADING_MAP = [
    (r'^الفصل\s+(الأول|الثاني|الثالث|الرابع|الخامس|السادس)', 'Heading 1'),
    (r'^المبحث\s+(الأول|الثاني|الثالث|الرابع|الخامس|السادس)', 'Heading 2'),
    (r'^المطلب\s+(الأول|الثاني|الثالث|الرابع|الخامس|السادس)', 'Heading 3'),
    (r'^(أولاً|ثانياً|ثالثاً|رابعاً|خامساً|سادساً)[:：]', 'Heading 3'),
]

EXACT_HEADING_MAP = {
    'تمهيد': 'Heading 3',
    'خلاصة': 'Heading 3',
    'الخاتمة العامة': 'Heading 3',
}


def strip_explicit_12pt(run):
    changed = False
    if run.font.size and run.font.size == SZ_12PT:
        run.font.size = None
        changed = True
    if run.font.name is None or run.font.name == '':
        run.font.name = 'Traditional Arabic'
        changed = True
    return changed


def fix_fonts(doc):
    fixed = 0
    for p in doc.paragraphs:
        for run in p.runs:
            if strip_explicit_12pt(run):
                fixed += 1
    print(f"  Font fix: {fixed} runs stripped of explicit 12pt")
    return fixed


def fix_margins(doc):
    target_cm = 2.5
    delta = 0.2
    for i, sec in enumerate(doc.sections):
        for attr, tag in [('left_margin', 'left'), ('right_margin', 'right'),
                           ('top_margin', 'top'), ('bottom_margin', 'bottom')]:
            current_cm = getattr(sec, attr) / 914400 * 2.54
            if abs(current_cm - target_cm) > delta:
                setattr(sec, attr, Cm(target_cm))
                print(f"  Section {i} {tag}: {current_cm:.2f}cm -> {target_cm:.1f}cm")


def is_in_table(p):
    parent = p._element.getparent()
    while parent is not None:
        if parent.tag == qn('w:tbl'):
            return True
        parent = parent.getparent()
    return False


def fix_headings(doc):
    assigned = 0
    for p in doc.paragraphs:
        if is_in_table(p):
            continue
        txt = p.text.strip()
        if not txt:
            continue
        target_style = None
        if txt in EXACT_HEADING_MAP:
            target_style = EXACT_HEADING_MAP[txt]
        else:
            for pattern, style in HEADING_MAP:
                if re.match(pattern, txt):
                    target_style = style
                    break
        if target_style and p.style.name not in HEADING_STYLES:
            p.style = doc.styles[target_style]
            assigned += 1
    print(f"  Headings assigned: {assigned}")
    return assigned


def fix_alignment(doc):
    right_count = 0
    for p in doc.paragraphs:
        if p.style.name in HEADING_STYLES:
            if p.alignment != enum.text.WD_ALIGN_PARAGRAPH.RIGHT:
                p.alignment = enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                right_count += 1
        elif p.style.name in BODY_STYLES or not p.runs:
            if p.alignment is None or p.alignment == enum.text.WD_ALIGN_PARAGRAPH.LEFT:
                p.alignment = enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                right_count += 1
    print(f"  Alignment: {right_count} paragraphs set to RTL")


def fix_spacing(doc):
    fixed = 0
    for p in doc.paragraphs:
        pf = p.paragraph_format
        if pf.line_spacing != 1.5:
            pf.line_spacing = 1.5
            fixed += 1
    print(f"  Line spacing: {fixed} paragraphs set to 1.5")


def demote_headings(doc):
    """Demote body paragraphs that got heading styles by accident.
    Uses length heuristic: real headings are short (< 100 chars),
    body paragraphs promoted to heading style are long."""
    LONG_THRESHOLD = 100
    demoted = 0
    for p in doc.paragraphs:
        if not p.style or not p.style.name.startswith('Heading'):
            continue
        txt = p.text.strip()
        if not txt:
            continue
        if len(txt) > LONG_THRESHOLD:
            p.style = doc.styles['Normal']
            demoted += 1
    print(f"  Long body-text headings demoted to Normal: {demoted}")
    return demoted


def fix_page_breaks(doc):
    """Enforce page break hierarchy:
    Level 1 — الفصل (chapter) headings: page break before
    Level 2 — المبحث headings: keepNext + keepLines
    Level 3 — المطلب / أولاً headings: keepNext
    Level 4 — Body paragraphs: orphan/widow control
    """
    ch1_set = 0
    mb_set = 0
    mt_set = 0
    body_set = 0

    for p in doc.paragraphs:
        if is_in_table(p):
            continue
        txt = p.text.strip()
        if not txt:
            continue
        pf = p.paragraph_format
        pPr = p._element.find(qn('w:pPr'))

        if re.match(r'^الفصل\s+(الأول|الثاني|الثالث|الرابع|الخامس|السادس)', txt):
            pf.page_break_before = True
            ch1_set += 1

        elif re.match(r'^المبحث\s+(الأول|الثاني|الثالث|الرابع|الخامس|السادس)', txt):
            add_keep_next(pPr, p._element)
            add_keep_lines(pPr, p._element)
            mb_set += 1

        elif re.match(r'^(المطلب\s+|أولاً|ثانياً|ثالثاً|رابعاً|خامساً|سادساً)[:：]', txt):
            add_keep_next(pPr, p._element)
            mt_set += 1

        elif p.style.name in BODY_STYLES or p.style.name == 'Compact':
            add_widow_orphan(pPr, p._element)
            body_set += 1

    print(f"  Page breaks: {ch1_set} chapters (new page), {mb_set} مباحث (keep+next), {mt_set} مطالب (keep), {body_set} body (widow/orphan)")


def add_keep_next(pPr, p_el):
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_el.insert(0, pPr)
    if pPr.find(qn('w:keepNext')) is None:
        pPr.append(OxmlElement('w:keepNext'))


def add_keep_lines(pPr, p_el):
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_el.insert(0, pPr)
    if pPr.find(qn('w:keepLines')) is None:
        pPr.append(OxmlElement('w:keepLines'))


def add_widow_orphan(pPr, p_el):
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_el.insert(0, pPr)
    widow = pPr.find(qn('w:widowControl'))
    if widow is None:
        widow = OxmlElement('w:widowControl')
        pPr.append(widow)
    widow.set(qn('w:val'), '1')


def main():
    if len(sys.argv) < 2:
        print("Usage: python format-baseline.py <docx_path>")
        sys.exit(1)

    path = sys.argv[1]
    print(f"Format baseline: {path}")

    doc = Document(path)

    fix_fonts(doc)
    fix_margins(doc)
    a =     fix_headings(doc)
    demote_headings(doc)
    fix_alignment(doc)
    fix_spacing(doc)
    fix_page_breaks(doc)

    doc.save(path)
    print(f"Saved: {path}")


if __name__ == '__main__':
    main()

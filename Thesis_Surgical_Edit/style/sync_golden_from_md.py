"""
sync_golden_from_md.py — Rebuild DOCX body from pandoc, keep cover+TOC from golden.

Strategy (v2):
1. Golden source = correct cover page, TOC, styles, hyperlinks
2. Pandoc output = correct content (body text, headings, tables, footnotes)
3. Build: copy golden → replace body content from pandoc → replace footnotes from pandoc

This produces a clean DOCX with:
- Correct cover page (from golden)
- Correct TOC (from golden, headings match because body has correct headings)
- Correct body content (from pandoc/MD)
- Correct tables (from pandoc, matching MD exactly)
- Correct footnotes (from pandoc, matching MD exactly)

Usage: python sync_golden_from_md.py <golden.docx> <pandoc.docx> [--save]
"""
import sys
import shutil
import zipfile
import os
from docx import Document
from lxml import etree

W_NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def is_toc_paragraph(para):
    """Check if paragraph is a TOC entry."""
    return 'toc' in para.style.name.lower()


def is_cover_page(para):
    """Check if paragraph is in cover page (before first Heading with 'إهداء')."""
    if para.style.name.startswith('Heading') and 'إهداء' in para.text:
        return True
    return False


def replace_body_from_pandoc(golden_path, pandoc_path, save=False):
    """Replace body content from pandoc, keep cover page + TOC from golden."""
    
    golden = Document(golden_path)
    pandoc = Document(pandoc_path)
    
    # === Step 1: Find boundaries in golden ===
    
    # Cover page end: first Heading with 'إهداء'
    cover_end = 0
    for i, p in enumerate(golden.paragraphs):
        if p.style.name.startswith('Heading') and 'إهداء' in p.text:
            cover_end = i
            break
    if cover_end == 0:
        cover_end = 14
    
    # TOC start/end
    toc_start = None
    toc_end = None
    for i, p in enumerate(golden.paragraphs):
        if is_toc_paragraph(p):
            if toc_start is None:
                toc_start = i
            toc_end = i + 1
    
    body_start = toc_end if toc_end else cover_end
    
    print("[SYNC] Cover page: P0-P%d" % (cover_end - 1))
    print("[SYNC] TOC: P%d-P%d" % (toc_start or 0, (toc_end or 0) - 1))
    print("[SYNC] Body starts at: P%d" % body_start)
    
    # === Step 3: Find where real body starts in pandoc ===
    # Pandoc doesn't separate cover/TOC from body — everything is flat.
    # We need to find the abstract (الملخص) — that's where real content starts.
    # Before that: cover (Title/Subtitle/Author), dedication, thanks, TOC, list of tables.
    # Those are already in the golden source's cover+TOC sections.
    
    pandoc_body_start = 0
    for i, p in enumerate(pandoc.paragraphs):
        if p.style.name.startswith('Heading') and 'الملخص' in p.text:
            pandoc_body_start = i
            break
    
    if pandoc_body_start == 0:
        # Fallback: find first H1
        for i, p in enumerate(pandoc.paragraphs):
            if p.style.name.startswith('Heading 1'):
                pandoc_body_start = i
                break
    
    print("[SYNC] Pandoc body starts at: P%d (%s)" % (
        pandoc_body_start, pandoc.paragraphs[pandoc_body_start].text[:40] if pandoc_body_start < len(pandoc.paragraphs) else "?"))
    
    # === Step 4: Get pandoc body paragraphs ===
    pandoc_body = []
    for p in pandoc.paragraphs[pandoc_body_start:]:
        pandoc_body.append(p)
    
    print("[SYNC] Pandoc body paragraphs: %d" % len(pandoc_body))
    print("[SYNC] Golden body paragraphs: %d" % (len(golden.paragraphs) - body_start))
    
    # === Step 4: Replace golden body with pandoc body ===
    body = golden.element.body
    
    # 4a: Remove old body paragraphs (from body_start onwards)
    # TOC entries are KEPT — Word COM will regenerate them
    to_remove = []
    for i, p in enumerate(golden.paragraphs[body_start:]):
        to_remove.append(p._element)
    
    for elem in to_remove:
        body.remove(elem)
    
    print("[SYNC] Removed %d old body paragraphs" % len(to_remove))
    
    # Add pandoc body paragraphs
    from docx.oxml.ns import qn
    added = 0
    for p in pandoc_body:
        # Copy paragraph element from pandoc
        new_p = etree.SubElement(body, qn('w:p'))
        # Copy all child elements
        for child in p._element:
            new_p.append(child)
        added += 1
    
    print("[SYNC] Added %d pandoc body paragraphs" % added)
    
    # === Step 5: Replace tables ===
    # Remove old tables
    old_tables = list(golden.tables)
    for t in old_tables:
        try:
            tbl = t._tbl
            tbl.getparent().remove(tbl)
        except:
            pass
    
    print("[SYNC] Removed %d old tables" % len(old_tables))
    
    # Add pandoc tables
    for t in pandoc.tables:
        # Copy table element
        new_tbl = etree.SubElement(body, qn('w:tbl'))
        for child in t._tbl:
            new_tbl.append(child)
    
    print("[SYNC] Added %d pandoc tables" % len(pandoc.tables))
    
    # === Step 6: Save and replace footnotes ===
    if save:
        golden.save(golden_path)
        print("[SYNC] Saved body: %s" % golden_path)
        
        # Replace footnotes at zip level
        replace_footnotes_zip(golden_path, pandoc_path)
    
    return added


def replace_footnotes_zip(target_docx, source_docx):
    """Replace footnotes.xml in target with the one from source."""
    temp_fn = None
    with zipfile.ZipFile(source_docx, 'r') as z:
        if 'word/footnotes.xml' in z.namelist():
            temp_fn = target_docx + '.fn_tmp'
            with open(temp_fn, 'wb') as f:
                f.write(z.read('word/footnotes.xml'))
            print("[SYNC] Extracted footnotes.xml (%d bytes)" % os.path.getsize(temp_fn))
        else:
            print("[SYNC] WARNING: No footnotes.xml in pandoc output")
            return False
    
    with open(temp_fn, 'rb') as f:
        fn_raw = f.read()
    
    temp_docx = target_docx + '.tmp'
    with zipfile.ZipFile(target_docx, 'r') as zin:
        with zipfile.ZipFile(temp_docx, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == 'word/footnotes.xml':
                    zout.writestr(item, fn_raw)
                else:
                    zout.writestr(item, zin.read(item.filename))
    
    shutil.move(temp_docx, target_docx)
    os.remove(temp_fn)
    
    with zipfile.ZipFile(target_docx, 'r') as z:
        with z.open('word/footnotes.xml') as f:
            root = etree.fromstring(f.read())
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            count = len(root.findall('.//w:footnote', ns))
            print("[SYNC] Replaced footnotes.xml — %d footnotes" % count)
    
    return True


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python sync_golden_from_md.py <golden.docx> <pandoc.docx> [--save]")
        sys.exit(1)
    
    golden_path = sys.argv[1]
    pandoc_path = sys.argv[2]
    save = "--save" in sys.argv
    
    count = replace_body_from_pandoc(golden_path, pandoc_path, save)
    sys.exit(0 if count > 0 else 1)

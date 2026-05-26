"""audit_thesis_comprehensive.py — Full thesis DOCX health check
Part of Academix v13.2 build pipeline
Usage: python audit_thesis_comprehensive.py <path/to.docx>
       python audit_thesis_comprehensive.py <path/to.docx> --json

Performs 10-category audit:
1. Document overview (paragraphs, tables, sections, size)
2. Section analysis (page, margins, page number format)
3. Heading structure (H1/H2/H3 count and hierarchy)
4. Body formatting (font, size, RTL, spacing)
5. Footnote count and content preview
6. Table quality (column widths, content fit, gaps)
7. Header/footer integrity
8. Cover page analysis
9. Content coverage (all required sections present)
10. Potential issues (empty paras, hidden numbers, etc.)
"""

import sys, os, json, re, zipfile
from xml.etree import ElementTree as ET
from docx import Document
from docx.oxml.ns import qn

NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

def emu_to_cm(emu):
    return round(emu / 914400 * 2.54, 2) if emu else 0

def run_audit(docx_path):
    if not os.path.exists(docx_path):
        return {"error": "File not found: " + docx_path}
    
    doc = Document(docx_path)
    size = os.path.getsize(docx_path)
    all_text = ' '.join(p.text for p in doc.paragraphs)
    report = {}
    
    # 1. DOCUMENT OVERVIEW
    report['overview'] = {
        'paragraphs': len(doc.paragraphs),
        'tables': len(doc.tables),
        'sections': len(doc.sections),
        'file_size_kb': round(size / 1024, 1)
    }
    
    # 2. SECTION ANALYSIS
    sections_info = []
    for i, sec in enumerate(doc.sections):
        pg = sec._sectPr.find(NS+'pgNumType')
        fmt = pg.get(NS+'fmt') if pg is not None else 'default_decimal'
        sections_info.append({
            'index': i,
            'page': '%.1fx%.1f' % (sec.page_width.cm, sec.page_height.cm),
            'margins': {
                'top': round(sec.top_margin.cm, 2),
                'bottom': round(sec.bottom_margin.cm, 2),
                'left': round(sec.left_margin.cm, 2),
                'right': round(sec.right_margin.cm, 2)
            },
            'page_num_format': fmt,
            'different_first_page': sec.different_first_page_header_footer
        })
    report['sections'] = sections_info
    
    # 3. HEADING STRUCTURE
    h1, h2, h3 = [], [], []
    for i, p in enumerate(doc.paragraphs):
        if p.style:
            sn = p.style.name
            if 'Heading 1' in sn: h1.append((i, p.text.strip()[:70]))
            elif 'Heading 2' in sn: h2.append((i, p.text.strip()[:70]))
            elif 'Heading 3' in sn: h3.append((i, p.text.strip()[:70]))
    
    prev = 0; skips = 0
    for i, p in enumerate(doc.paragraphs):
        lv = 0
        if p.style:
            sn = p.style.name
            if 'Heading 1' in sn: lv=1
            elif 'Heading 2' in sn: lv=2
            elif 'Heading 3' in sn: lv=3
        if lv:
            if lv > prev+1 and prev > 0: skips += 1
            prev = lv
    
    report['headings'] = {
        'h1_count': len(h1),
        'h2_count': len(h2),
        'h3_count': len(h3),
        'h1_titles': [t for _, t in h1],
        'h2_titles': [t for _, t in h2],
        'hierarchy_skips': skips
    }
    
    # 4. BODY FORMATTING
    font_bad=0; size_bad=0; rtl_bad=0; sp_bad=0
    sample = [p for p in doc.paragraphs if p.style and p.style.name not in ('toc 1','toc 2','toc 3') and p.text.strip()][:150]
    for p in sample:
        for r in p.runs:
            if r.font.name and r.font.name != 'Traditional Arabic': font_bad+=1
            if r.font.size and abs(r.font.size.pt - 14) >= 0.5: size_bad+=1
        if p.alignment is not None and p.alignment != 2: rtl_bad+=1
        ls = p.paragraph_format.line_spacing
        if ls and abs(ls - 1.5) > 0.1: sp_bad+=1
    
    report['formatting'] = {
        'font_bad_runs': font_bad,
        'size_bad_runs': size_bad,
        'rtl_bad_paras': rtl_bad,
        'spacing_bad_paras': sp_bad,
        'sample_size': len(sample)
    }
    
    # 5. FOOTNOTES
    fn_count = 0; fn_preview = []
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            if 'word/footnotes.xml' in z.namelist():
                tree = ET.parse(z.open('word/footnotes.xml'))
                root = tree.getroot()
                for fn in root.findall('.//w:footnote', ns):
                    fid = fn.attrib.get(NS+'id', '')
                    if fid in ('0','-1'): continue
                    fn_count += 1
                    text = ''.join(t.text or '' for t in fn.findall('.//w:t', ns))
                    fn_preview.append(text[:80])
    except:
        pass
    report['footnotes'] = {'count': fn_count, 'preview': fn_preview}
    
    # 6. TABLE ANALYSIS
    tables_info = []
    table_gaps = []
    for ti, t in enumerate(doc.tables):
        rows = len(t.rows)
        cols = len(t.columns)
        hdr = t.rows[0].cells[0].text.strip()[:40] if t.rows[0].cells else ''
        
        max_chars = []
        for ci in range(cols):
            mc = 0
            for row in t.rows:
                if ci < len(row.cells):
                    mc = max(mc, len(row.cells[ci].text.strip()))
            max_chars.append(mc)
        
        # Check for tables with auto-width (no explicit column widths = gaps)
        has_widths = False
        total_w = 0
        if rows > 0:
            for ci in range(cols):
                cell = t.rows[0].cells[ci]
                tcPr = cell._tc.find(NS+'tcPr')
                if tcPr is not None:
                    tcW = tcPr.find(NS+'tcW')
                    if tcW is not None:
                        w = int(tcW.get(NS+'w'))
                        if w > 0: 
                            has_widths = True
                            total_w += w
        
        table_entry = {
            'index': ti,
            'rows': rows,
            'cols': cols,
            'header': hdr,
            'max_chars_per_col': max_chars,
            'has_column_widths': has_widths,
            'total_width_cm': round(total_w/914400*2.54, 1) if total_w else 0
        }
        tables_info.append(table_entry)
        
        if not has_widths and cols > 1:
            table_gaps.append("Table %d: no explicit column widths (auto-stretch = gaps)" % ti)
        elif has_widths and total_w > 0:
            for ci in range(cols):
                cell = t.rows[0].cells[ci]
                tcPr = cell._tc.find(NS+'tcPr')
                if tcPr is not None:
                    tcW = tcPr.find(NS+'tcW')
                    if tcW is not None:
                        w = int(tcW.get(NS+'w'))
                        w_cm = w/914400*2.54
                        if w_cm > 8 and max_chars[ci] < 10:
                            table_gaps.append("Table %d col %d: %.1fcm wide but only %d chars content — GAP" % (
                                ti, ci, w_cm, max_chars[ci]))
    
    report['tables'] = {
        'count': len(doc.tables),
        'details': tables_info,
        'gap_warnings': table_gaps
    }
    
    # 7. HEADER/FOOTER CHECK
    hf_info = {}
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            for name in sorted(z.namelist()):
                if 'header' in name or 'footer' in name:
                    xml = z.read(name)
                    root = ET.fromstring(xml)
                    texts = [t.text or '' for t in root.findall('.//w:t', ns)]
                    instrs = [i.text or '' for i in root.findall('.//w:instrText', ns)]
                    hf_info[name] = {
                        'texts': texts,
                        'instrText': instrs,
                        'size': len(xml)
                    }
    except:
        pass
    report['headers_footers'] = hf_info
    
    # 8. COVER ANALYSIS
    cover_fonts = []
    for i in range(min(12, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        for r in p.runs:
            if r.font.name:
                cover_fonts.append((i, r.font.name, str(r.font.size)))
                break
    report['cover'] = {
        'font_info': cover_fonts,
        'has_cover_title': 'الجمهورية الجزائرية' in (doc.paragraphs[0].text if doc.paragraphs else '')
    }
    
    # 9. CONTENT COVERAGE
    checks = {
        'cover': 'الجمهورية الجزائرية' in all_text,
        'dedication': 'إهداء' in all_text,
        'thanks': 'شكر وتقدير' in all_text,
        'toc': 'فهرس المحتويات' in all_text,
        'abstract_ar': 'الملخص' in all_text,
        'abstract_fr': 'Résumé' in all_text,
        'ch1': 'الفصل الأول' in all_text,
        'ch2': 'الفصل الثاني' in all_text,
        'ch3': 'الفصل الثالث' in all_text,
        'ch4': 'الفصل الرابع' in all_text,
        'bibliography': 'قائمة المصادر' in all_text,
        'annexes': 'الملاحق' in all_text or 'الملحق' in all_text,
        'glossary': 'قائمة المصطلحات' in all_text,
        'abbreviations': 'المختصرات' in all_text,
        'art001_case': 'ART-001' in all_text,
        'eoq': '176' in all_text,
        'rop': '212.4' in all_text,
        'safety_stock': '200' in all_text,
        'abc_analysis': 'ABC' in all_text,
        'xyz_analysis': 'XYZ' in all_text,
        'wilson_model': 'ويلسون' in all_text,
        'vba_mention': 'VBA' in all_text,
    }
    report['content_coverage'] = {
        k: v for k, v in checks.items()
    }
    missing = [k for k, v in checks.items() if not v]
    report['content_missing'] = missing
    
    # 10. ISSUES SUMMARY
    issues = []
    
    # Page numbering (handle single-section pandoc output gracefully)
    n_sections = len(doc.sections)
    if n_sections > 0:
        cover_pg = doc.sections[0]._sectPr.find(NS+'pgNumType')
        if cover_pg is not None:
            cover_fmt = cover_pg.get(NS+'fmt')
            if cover_fmt and cover_fmt != 'none':
                issues.append("Cover has pgNumType (fmt=%s) — may show page number" % cover_fmt)
    
    if n_sections > 1:
        s1_fmt = 'default_decimal'
        s1 = doc.sections[1]._sectPr.find(NS+'pgNumType')
        if s1 is not None:
            s1_fmt = s1.get(NS+'fmt')
        if s1_fmt == 'default_decimal':
            issues.append("Section 1 (front matter) uses decimal not lowerRoman")
    else:
        issues.append("Single-section document (expected after pandoc rebuild — run fix_docx_sections.py)")
    
    # TOC status
    toc_count = sum(1 for p in doc.paragraphs if p.style and 'toc' in p.style.name.lower())
    if toc_count > 0:
        issues.append("TOC: %d entries with inline page numbers (manual typing)" % toc_count)
    else:
        issues.append("TOC: NO entries found")
    
    # Empty paragraphs
    empty_count = sum(1 for p in doc.paragraphs if not p.text.strip())
    if empty_count > 100:
        issues.append("High empty paragraph count: %d (affects spacing)" % empty_count)
    
    # Cover font
    if cover_fonts and cover_fonts[0][1] != 'Traditional Arabic':
        issues.append("Cover uses %s instead of Traditional Arabic (verify this is intentional)" % cover_fonts[0][1])
    
    # Table gaps
    issues.extend(table_gaps[:5])  # Top 5
    
    report['issues'] = issues
    report['issue_count'] = len(issues)
    report['source_file'] = os.path.basename(docx_path)
    
    return report


def print_report(report):
    if 'error' in report:
        print("ERROR: %s" % report['error'])
        return
    
    print("=" * 70)
    print("ACADEMIX v13.2 — COMPREHENSIVE THESIS AUDIT")
    print("=" * 70)
    
    ov = report['overview']
    print("\n--- 1. OVERVIEW ---")
    print("  File: %s" % report['source_file'])
    print("  Paragraphs: %d  Tables: %d  Sections: %d  Size: %.1f KB" % (
        ov['paragraphs'], ov['tables'], ov['sections'], ov['file_size_kb']))
    
    print("\n--- 2. SECTIONS ---")
    for s in report['sections']:
        print("  Sec%d: %s cm mrg=%.1f|%.1f|%.1f|%.1f fmt=%s" % (
            s['index'], s['page'], s['margins']['top'], s['margins']['bottom'],
            s['margins']['left'], s['margins']['right'], s['page_num_format']))
    
    h = report['headings']
    print("\n--- 3. HEADINGS ---")
    print("  H1: %d  H2: %d  H3: %d  Hierarchy skips: %d" % (
        h['h1_count'], h['h2_count'], h['h3_count'], h['hierarchy_skips']))
    
    f = report['formatting']
    print("\n--- 4. FORMATTING (sample %d) ---" % f['sample_size'])
    print("  Font bad: %d  Size bad: %d  RTL bad: %d  Spacing bad: %d" % (
        f['font_bad_runs'], f['size_bad_runs'], f['rtl_bad_paras'], f['spacing_bad_paras']))
    
    fn = report['footnotes']
    print("\n--- 5. FOOTNOTES: %d ---" % fn['count'])
    
    tbl = report['tables']
    print("\n--- 6. TABLES: %d ---" % tbl['count'])
    for t in tbl['details'][:5]:
        print("  T%d: %dr x %dc [%s]" % (t['index'], t['rows'], t['cols'], t['header'][:30]))
    if tbl['count'] > 5:
        print("  ... %d more" % (tbl['count'] - 5))
    for gw in tbl['gap_warnings'][:5]:
        print("  GAP: %s" % gw)
    
    print("\n--- 7. HEADERS/FOOTERS ---")
    for name, info in report['headers_footers'].items():
        print("  %s: %s %s" % (name, info['texts'], info['instrText']))
    
    print("\n--- 8. CONTENT COVERAGE ---")
    c = report['content_coverage']
    for k, v in c.items():
        print("  %s %s" % ('OK' if v else 'MISS', k))
    if report['content_missing']:
        print("  MISSING: %s" % ', '.join(report['content_missing']))
    
    print("\n--- 9. ISSUES FOUND: %d ---" % report['issue_count'])
    for issue in report['issues']:
        print("  * %s" % issue)
    
    print("\n" + "=" * 70)


def main():
    if len(sys.argv) < 2:
        print("Usage: python audit_thesis_comprehensive.py <path/to.docx> [--json]")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    output_json = '--json' in sys.argv
    
    report = run_audit(docx_path)
    
    if output_json:
        print(json.dumps(report, indent=2, ensure_ascii=False))
    else:
        print_report(report)
    
    # Exit with code 0 if issues < 5, 1 if serious issues
    if report.get('issue_count', 0) >= 5:
        sys.exit(1)


if __name__ == '__main__':
    main()

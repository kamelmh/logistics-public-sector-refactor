"""fix_docx_footnotes.py — Fix footnotes.xml namespace issue for Word compatibility
"""
import zipfile, os, shutil, re, sys
from lxml import etree

def fix_docx(docx_path, desktop_copy=None):
    backup_path = docx_path.replace('.docx', '_pre_wordfix.docx')
    shutil.copy2(docx_path, backup_path)
    print(f"Backup: {backup_path}")

    with zipfile.ZipFile(docx_path, 'r') as z:
        entries = {name: z.read(name) for name in z.namelist()}

    # Fix footnotes.xml — remove ns1:Ignorable attribute
    for target in ['word/footnotes.xml', 'word/endnotes.xml']:
        if target in entries:
            xml = entries[target].decode('utf-8')
            # Remove the Ignorable attribute which references undeclared namespaces
            xml = re.sub(r'\s+xmlns:ns1="[^"]*"', '', xml)
            xml = re.sub(r'\s+ns1:Ignorable="[^"]*"', '', xml)
            entries[target] = xml.encode('utf-8')
            print(f"Fixed: {target}")

    # Write fixed docx
    tmp_path = docx_path.replace('.docx', '_tmp.docx')
    with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in entries.items():
            zout.writestr(name, data)
    shutil.move(tmp_path, docx_path)

    # Verify
    with zipfile.ZipFile(docx_path, 'r') as z:
        for target in ['word/footnotes.xml', 'word/endnotes.xml']:
            if target in z.namelist():
                root = etree.fromstring(z.read(target))
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                items = root.findall('.//w:footnote', ns) if 'footnote' in target else root.findall('.//w:endnote', ns)
                print(f"Verified {target}: {len(items)} items, XML valid")

    # Copy to desktop if requested
    if desktop_copy:
        shutil.copy2(docx_path, desktop_copy)
        print(f"Desktop copy: {desktop_copy} ({os.path.getsize(desktop_copy)} bytes)")

    print(f"Done: {docx_path} ({os.path.getsize(docx_path)} bytes)")

    # Validate with python-docx
    from docx import Document
    doc = Document(docx_path)
    print(f"python-docx OK: {len(doc.paragraphs)}p, {len(doc.tables)}tbl")

if __name__ == "__main__":
    docx = sys.argv[1]
    desktop = sys.argv[2] if len(sys.argv) > 2 else None
    fix_docx(docx, desktop)

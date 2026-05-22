"""inspect_docx_metrics.py — Fast DOCX metrics via python-docx + lxml (no COM)
Outputs JSON to stdout, callable from thesis-doctor.ps1.

Usage: python inspect_docx_metrics.py <path/to/docx> [--json]
"""

import sys, os, json, zipfile
from xml.etree import ElementTree as ET
from docx import Document
from docx.shared import Cm, Pt, Emu

GOLDEN = {
    "bodyFont": "Traditional Arabic",
    "bodySize": 14,
    "minParagraphs": 600,
    "minTables": 21,
    "minFootnotes": 46,
    "minTocEntries": 40,
    "minSeqFields": 21,
    "minSections": 4,
    "coverTitle": "الجمهورية الجزائرية",
    "pageWidthCm": 21.0,
    "pageHeightCm": 29.7,
    "marginCm": 2.5,
    "lineSpacing": 1.5,
}

W_NS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
}


def emu_to_cm(emu):
    if emu is None:
        return None
    return round(emu / 360000, 2)


def parse_footnotes(docx_path):
    """Parse footnotes.xml from DOCX zip using lxml"""
    fn_count = 0
    fn_bidi_ok = 0
    fn_bidi_bad = 0
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            if 'word/footnotes.xml' in z.namelist():
                tree = ET.parse(z.open('word/footnotes.xml'))
                root = tree.getroot()
                for fn in root.findall('.//w:footnote', W_NS):
                    fn_id = fn.attrib.get(f'{{{W_NS["w"]}}}id', '')
                    if fn_id in ('0', '-1'):
                        continue
                    fn_count += 1
                    par = fn.find('.//w:p', W_NS)
                    if par is not None:
                        pPr = par.find('w:pPr', W_NS)
                        if pPr is not None:
                            jc = pPr.find('w:jc', W_NS)
                            if jc is not None:
                                val = jc.attrib.get(f'{{{W_NS["w"]}}}val', '')
                                if val == 'right':
                                    fn_bidi_ok += 1
                                else:
                                    fn_bidi_bad += 1
                            else:
                                fn_bidi_bad += 1
                        else:
                            fn_bidi_bad += 1
                    else:
                        fn_bidi_bad += 1
            else:
                pass
    except Exception as e:
        pass
    return fn_count, fn_bidi_ok, fn_bidi_bad


def parse_fields(docx_path):
    """Parse SEQ and TOC fields from document.xml"""
    seq_count = 0
    toc_count = 0
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            if 'word/document.xml' in z.namelist():
                tree = ET.parse(z.open('word/document.xml'))
                root = tree.getroot()
                for field in root.findall('.//w:fldSimple', W_NS):
                    instr = field.attrib.get(f'{{{W_NS["w"]}}}instr', '')
                    if 'SEQ' in instr.upper():
                        seq_count += 1
                    if 'TOC' in instr.upper() or 'INDEX' in instr.upper():
                        toc_count += 1
                for fld_code in root.findall('.//w:instrText', W_NS):
                    txt = (fld_code.text or '')
                    if 'SEQ' in txt.upper():
                        seq_count += 1
                    if 'TOC' in txt.upper() or 'INDEX' in txt.upper():
                        toc_count += 1
    except Exception as e:
        pass
    return seq_count, toc_count


def inspect_docx(docx_path):
    doc = Document(docx_path)
    paras = doc.paragraphs
    para_count = len(paras)

    # --- Sections ---
    sec = doc.sections[0]
    pw = round(sec.page_width.cm, 1) if sec.page_width else GOLDEN["pageWidthCm"]
    ph = round(sec.page_height.cm, 1) if sec.page_height else GOLDEN["pageHeightCm"]
    margins = {
        "top": round(sec.top_margin.cm, 2) if sec.top_margin else GOLDEN["marginCm"],
        "bottom": round(sec.bottom_margin.cm, 2) if sec.bottom_margin else GOLDEN["marginCm"],
        "left": round(sec.left_margin.cm, 2) if sec.left_margin else GOLDEN["marginCm"],
        "right": round(sec.right_margin.cm, 2) if sec.right_margin else GOLDEN["marginCm"],
    }
    section_count = len(doc.sections)
    sec_names = []
    for i, s in enumerate(doc.sections, 1):
        sw = round(s.page_width.cm, 1) if s.page_width else pw
        sh = round(s.page_height.cm, 1) if s.page_height else ph
        sec_names.append(f"Section {i}: {sw}x{sh}cm")

    # --- Cover page ---
    found = False
    for p in paras[:10]:
        txt = p.text.strip()
        if GOLDEN["coverTitle"] in txt:
            found = True
            break

    # --- Headings ---
    h1 = h2 = h3 = 0
    headings = []
    for p in paras:
        sname = p.style.name
        if sname in ("Heading 1", "Heading 2", "Heading 3", "Titre 1", "Titre 2", "Titre 3"):
            level = 1 if "1" in sname else (2 if "2" in sname else 3)
            headings.append({"level": level, "text": p.text[:60].strip()})
            if level == 1: h1 += 1
            elif level == 2: h2 += 1
            else: h3 += 1

    # --- Body text formatting ---
    body_font_ok = body_font_bad = 0
    body_size_ok = body_size_bad = 0
    rtl_ok = rtl_bad = 0
    spacing_ok = spacing_bad = 0

    for p in paras:
        sname = p.style.name
        if any(k in sname for k in ("Heading", "Titre", "TOC", "Table des", "Caption", "Légende", "Footnote", "Note de fin")):
            continue
        txt = p.text.strip()
        if not txt:
            continue

        # Font check via first run
        runs = p.runs
        if runs:
            f = runs[0].font
            fname = f.name if f.name else ""
            fsize = f.size
            if fname == GOLDEN["bodyFont"]:
                body_font_ok += 1
            else:
                body_font_bad += 1
            if fsize and abs(fsize.pt - GOLDEN["bodySize"]) < 0.5:
                body_size_ok += 1
            else:
                body_size_bad += 1
        else:
            body_font_bad += 1
            body_size_bad += 1

        # RTL
        align = p.paragraph_format.alignment
        if align == 2:  # RIGHT
            rtl_ok += 1
        else:
            rtl_bad += 1

        # Spacing
        ls = p.paragraph_format.line_spacing
        if ls and 1.4 <= ls <= 1.6:
            spacing_ok += 1
        else:
            spacing_bad += 1

    # --- Tables ---
    tables = doc.tables
    table_count = len(tables)
    table_details = []
    for t in tables:
        table_details.append({
            "rows": len(t.rows),
            "cols": len(t.columns),
            "hasHeader": len(t.rows) > 1,
        })

    # --- Styles in use ---
    styles_in_use = {}
    for p in paras:
        sname = p.style.name
        styles_in_use[sname] = styles_in_use.get(sname, 0) + 1

    # --- Footnotes (via XML) ---
    fn_count, fn_bidi_ok, fn_bidi_bad = parse_footnotes(docx_path)

    # --- Fields (via XML) ---
    seq_count, toc_count = parse_fields(docx_path)

    result = {
        "paragraph_count": para_count,
        "page_width_cm": pw,
        "page_height_cm": ph,
        "margins": margins,
        "section_count": section_count,
        "sections_detail": sec_names,
        "cover_detected": found,
        "h1_count": h1,
        "h2_count": h2,
        "h3_count": h3,
        "headings": headings,
        "body_font_ok": body_font_ok,
        "body_font_bad": body_font_bad,
        "body_size_ok": body_size_ok,
        "body_size_bad": body_size_bad,
        "rtl_ok": rtl_ok,
        "rtl_bad": rtl_bad,
        "spacing_ok": spacing_ok,
        "spacing_bad": spacing_bad,
        "table_count": table_count,
        "tables": table_details,
        "footnote_count": fn_count,
        "footnote_bidi_ok": fn_bidi_ok,
        "footnote_bidi_bad": fn_bidi_bad,
        "seq_field_count": seq_count,
        "toc_entry_count": toc_count,
        "styles": styles_in_use,
        "timestamp": None,
    }

    return result


def main():
    if len(sys.argv) < 2:
        print("Usage: python inspect_docx_metrics.py <path/to.docx> [--json]", file=sys.stderr)
        sys.exit(1)

    docx_path = sys.argv[1]
    as_json = "--json" in sys.argv

    if not os.path.exists(docx_path):
        print(f"[ERROR] File not found: {docx_path}", file=sys.stderr)
        sys.exit(1)

    try:
        result = inspect_docx(docx_path)
        result["timestamp"] = __import__("datetime").datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        if as_json:
            print(json.dumps(result, ensure_ascii=False, indent=2))
        else:
            print(f"  PDFX: {docx_path}")
            print(f"  Paragraphs: {result['paragraph_count']}")
            print(f"  Body font OK: {result['body_font_ok']} / Bad: {result['body_font_bad']}")
            print(f"  Body size OK: {result['body_size_ok']} / Bad: {result['body_size_bad']}")
            print(f"  RTL OK: {result['rtl_ok']} / Bad: {result['rtl_bad']}")
            print(f"  Spacing 1.5 OK: {result['spacing_ok']} / Bad: {result['spacing_bad']}")
            print(f"  H1: {result['h1_count']}  H2: {result['h2_count']}  H3: {result['h3_count']}")
            print(f"  Tables: {result['table_count']}")
            print(f"  Footnotes: {result['footnote_count']}")
            print(f"  SEQ fields: {result['seq_field_count']}  TOC entries: {result['toc_entry_count']}")
            print(f"  Sections: {result['section_count']}  Cover: {result['cover_detected']}")
            print(f"  Page: {result['page_width_cm']}x{result['page_height_cm']}cm")
            print(f"  Margins: T={result['margins']['top']} B={result['margins']['bottom']} L={result['margins']['left']} R={result['margins']['right']}cm")

    except Exception as e:
        print(f"[ERROR] {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()

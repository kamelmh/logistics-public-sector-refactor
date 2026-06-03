"""
⚠️  DEPRECATED — Replaced by: audit_thesis_comprehensive.py
    This script is superseded and no longer part of the active pipeline.
    Kept for reference only. Do not use in new workflows.
"""
"""audit-thesis.py — Full thesis completeness audit against BTS/GSL public sector standards
Usage: python audit-thesis.py <path/to.docx> [--json] [--threshold N]
"""
import sys, os, json, argparse, zipfile
from xml.etree import ElementTree as ET
from docx import Document
from docx.shared import Cm, Pt

W_NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

def audit(docx_path, size_threshold=50000):
    doc = Document(docx_path)
    paras = doc.paragraphs
    sections = doc.sections
    p_count = len(paras)
    s_count = len(sections)
    fsize = os.path.getsize(docx_path)
    
    # Extract all heading text by level
    h1 = []; h2 = []; h3 = []
    for p in paras:
        if p.style and p.style.name:
            sn = p.style.name
            if 'Heading 1' in sn: h1.append(p.text.strip())
            elif 'Heading 2' in sn: h2.append(p.text.strip())
            elif 'Heading 3' in sn: h3.append(p.text.strip())
    
    all_text = " ".join(p.text for p in paras if p.text)
    body_text = [p.text for p in paras if p.style and p.style.name in 
        ['Normal', 'Compact', 'Body Text', 'No Spacing'] and p.text.strip()]
    total_words = sum(len(t.split()) for t in body_text)
    
    # Section info
    s0 = sections[0] if sections else None
    page_a4 = False; margins_25 = False
    if s0:
        pw, ph = s0.page_width.cm, s0.page_height.cm
        page_a4 = abs(pw-21) < 0.1 and abs(ph-29.7) < 0.1
        mt, mb, ml, mr = s0.top_margin.cm, s0.bottom_margin.cm, s0.left_margin.cm, s0.right_margin.cm
        margins_25 = all(abs(x-2.5) < 0.1 for x in [mt, mb, ml, mr])
    
    # Footnotes
    fn_count = 0
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            if 'word/footnotes.xml' in z.namelist():
                tree = ET.parse(z.open('word/footnotes.xml'))
                for fn in tree.findall('.//w:footnote', W_NS):
                    fid = fn.attrib.get(f'{{{W_NS["w"]}}}id', '')
                    if fid not in ('0', '-1'): fn_count += 1
    except: pass
    
    def has(texts, patterns):
        for t in texts:
            for pat in patterns:
                if pat in t: return True
        return False
    
    results = []
    
    # A. FRONT MATTER
    results.append(("A01", "Front Matter", "Abstract (Arabic)", "Required", 
        "الملخص" in all_text, "الملخص heading present" if "الملخص" in all_text else "Missing Arabic abstract"))
    results.append(("A02", "Front Matter", "Résumé (French)", "Required",
        "Résumé" in all_text, "Résumé present" if "Résumé" in all_text else "Missing French résumé"))
    results.append(("A03", "Front Matter", "Table of Contents (TOC)", "Required",
        has(h1+h2+h3, ["فهرس", "المحتويات"]) or "فهرس المحتويات" in all_text[:1000], 
        "TOC present" if has(h1+h2+h3, ["فهرس", "المحتويات"]) else "TOC heading present only in text" if "فهرس المحتويات" in all_text[:1000] else "Missing TOC"))
    results.append(("A04", "Front Matter", "List of Tables", "Recommended",
        has(h1+h2+h3, ["الجداول", "جدول"]), "List of tables present" if has(h1+h2+h3, ["الجداول", "جدول"]) else "Missing list of tables"))
    results.append(("A05", "Front Matter", "List of Abbreviations", "Recommended",
        has(h1+h2+h3, ["مختصرات", "الرموز", "مصطلحات", "Lexique"]), "Abbreviations present"))
    
    # B. MAIN BODY — Chapter Structure
    chapters = [h for h in h1 if "الفصل" in h]
    results.append(("B01", "Main Body", "Chapters (≥4)", "Required", len(chapters) >= 4, 
        f"{len(chapters)} chapters found" if len(chapters) >= 4 else f"Only {len(chapters)} chapters"))
    results.append(("B02", "Main Body", "General Introduction", "Required",
        has(h1, ["المقدمة", "Introduction", "تقديم"]), "Introduction present"))
    results.append(("B03", "Main Body", "Each chapter has subsections (H2)", "Required",
        len(h2) >= len(chapters) * 3, f"H2={len(h2)} >= {len(chapters)*3}"))
    results.append(("B04", "Main Body", "Heading hierarchy (H1→H2→H3)", "Required",
        len(h1) >= 3 and len(h2) >= len(h1) and len(h3) <= len(h2) * 3, 
        f"H1={len(h1)} H2={len(h2)} H3={len(h3)}"))
    results.append(("B05", "Main Body", "Section depth (H3)", "Recommended",
        len(h3) >= 15, f"H3={len(h3)} >= 15" if len(h3) >= 15 else f"H3={len(h3)} < 15"))
    
    # C. CONTENT VOLUME
    results.append(("C01", "Content Volume", "Minimum paragraphs (≥400)", "Required",
        p_count >= 400, f"{p_count} paragraphs"))
    results.append(("C02", "Content Volume", "Minimum tables (≥15)", "Required",
        len(doc.tables) >= 15, f"{len(doc.tables)} tables"))
    results.append(("C03", "Content Volume", "Minimum footnotes (≥10)", "Recommended",
        fn_count >= 10, f"{fn_count} footnotes"))
    results.append(("C04", "Content Volume", "Minimum body words (≥5000)", "Required",
        total_words >= 5000, f"{total_words} words"))
    results.append(("C05", "Content Volume", "File size threshold (≥50KB)", "Required",
        fsize > size_threshold, f"{fsize//1024}KB"))
    
    # D. BACK MATTER
    results.append(("D01", "Back Matter", "Bibliography/Bibliographie", "Required",
        has(h1+h2, ["المصادر", "المراجع", "Bibliographie"]) or "المصادر" in all_text or "المراجع" in all_text,
        "Bibliography present"))
    results.append(("D02", "Back Matter", "Annexes/Appendices", "Required",
        has(h1+h2+h3, ["الملاحق", "Annexes", "ملحق"]), "Annexes present"))
    results.append(("D03", "Back Matter", "Terminology glossary", "Recommended",
        has(h1+h2+h3, ["مصطلحات", "Glossaire"]), "Glossary present" if has(h1+h2+h3, ["مصطلحات", "Glossaire"]) else "Missing glossary"))
    results.append(("D04", "Back Matter", "Annexes have sub-sections", "Recommended",
        len(h2) >= 6, f"{len(h2)} H2 headings (indicates sub-section annexes)" if len(h2) >= 6 else "Few H2 headings"))
    
    # E. FORMATTING
    results.append(("E01", "Formatting", "Page size A4 (21x29.7cm)", "Required",
        page_a4, f"A4={page_a4}"))
    results.append(("E02", "Formatting", "Margins 2.5cm all sides", "Required",
        margins_25, f"Margins OK={margins_25}"))
    results.append(("E03", "Formatting", "Sections (≥3) for proper layout", "Required",
        s_count >= 3, f"{s_count} sections"))
    
    body_styles = ['Normal', 'Compact', 'Body Text', 'List Paragraph', 'No Spacing']
    bp = [p for p in paras if p.style and p.style.name in body_styles]
    
    font_bad = sum(1 for p in bp[:80] if p.runs and p.runs[0].font.name and p.runs[0].font.name != 'Traditional Arabic')
    size_bad = sum(1 for p in bp[:80] if p.runs and p.runs[0].font.size and p.runs[0].font.size.pt != 14)
    threshold = max(1, int(min(len(bp), 80) * 0.05))
    
    results.append(("E04", "Formatting", "Font: Traditional Arabic", "Required",
        font_bad <= threshold, f"{font_bad} bad (threshold={threshold})"))
    results.append(("E05", "Formatting", "Font size: 14pt", "Required",
        size_bad <= threshold, f"{size_bad} bad (threshold={threshold})"))
    
    # F. DOMAIN CONTENT — BTS Public Sector Specific
    results.append(("F01", "Domain", "BTS GSL TAG1801 reference", "Required",
        "TAG1801" in all_text or "BTS" in all_text, "BTS/TAG1801 reference present"))
    results.append(("F02", "Domain", "Algerian regulatory framework", "Required",
        any(k in all_text for k in ["القطاع العام", "الجزائري", "Algérie", "CNEPD", "Direction de l'Education"]),
        "Algerian context present"))
    results.append(("F03", "Domain", "EOQ/Wilson formula", "Required",
        any(k in all_text for k in ["EOQ", "Wilson", "كمية"]), "EOQ reference present"))
    results.append(("F04", "Domain", "Safety Stock / ROP", "Required",
        any(k in all_text for k in ["ROP", "Safety", "الأمان", "إعادة الطلب", "نقطة"]), "Stock management concepts present"))
    results.append(("F05", "Domain", "ERP/VBA system design", "Required",
        any(k in all_text for k in ["ERP", "VBA", "نظام دعم القرار", "DSS"]), "ERP/DSS system present"))
    results.append(("F06", "Domain", "Validation/experimental results", "Required",
        any(k in all_text for k in ["نتائج", "تجريرب", "تحقق", "validation"]), "Results/validation section present"))
    
    # G. INNOVATION / ADDED VALUE
    results.append(("G01", "Innovation", "LLM/AI methodology mention", "Recommended",
        any(k in all_text for k in ["LLM", "GPT", "AI", "الذكاء", "LangChain", "RAG"]),
        "AI/LLM methodology present"))
    results.append(("G02", "Innovation", "Dashboard/visualization", "Recommended",
        any(k in all_text for k in ["Dashboard", "dashboard", "لوحة القيادة", "Tableau de bord"]),
        "Dashboard concept present"))
    results.append(("G03", "Innovation", "CrossFlow/multi-agent mention", "Recommended",
        any(k in all_text for k in ["CrossFlow", "multi-agent", "Agent", "وكلاء"]),
        "CrossFlow/agent concept present"))
    
    # Compile
    checks = []
    passed = 0; failed = 0
    categories = {}
    for code, cat, name, severity, ok, msg in results:
        checks.append({"code": code, "category": cat, "name": name, "severity": severity, "passed": ok, "message": msg})
        if ok: passed += 1
        else: failed += 1
        categories.setdefault(cat, {"passed": 0, "failed": 0, "total": 0})
        categories[cat]["total"] += 1
        if ok: categories[cat]["passed"] += 1
        else: categories[cat]["failed"] += 1
    
    return {
        "docx": docx_path,
        "file_size_kb": fsize // 1024,
        "paragraphs": p_count,
        "tables": len(doc.tables),
        "footnotes": fn_count,
        "words": total_words,
        "h1": len(h1), "h2": len(h2), "h3": len(h3),
        "checks": checks,
        "categories": categories,
        "summary": {"passed": passed, "failed": failed, "total": len(checks), "score_pct": round(passed / len(checks) * 100, 1) if checks else 0}
    }

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Audit thesis DOCX completeness")
    parser.add_argument("docx", help="Path to thesis DOCX")
    parser.add_argument("--json", action="store_true", help="JSON output")
    parser.add_argument("--threshold", type=int, default=50000, help="Size threshold in bytes")
    args = parser.parse_args()
    
    result = audit(args.docx, size_threshold=args.threshold)
    
    if args.json:
        print(json.dumps(result, indent=2, ensure_ascii=False))
    else:
        s = result["summary"]
        print("=" * 70)
        print(f"  THESIS COMPLETENESS AUDIT: {os.path.basename(args.docx)}")
        print(f"  Size: {result['file_size_kb']}KB | Paras: {result['paragraphs']} | Tables: {result['tables']} | Fn: {result['footnotes']} | Words: {result['words']}")
        print("=" * 70)
        
        current_cat = ""
        for c in result["checks"]:
            if c["category"] != current_cat:
                current_cat = c["category"]
                print(f"\n── {current_cat} ──")
            status = "PASS" if c["passed"] else "FAIL"
            sev = f"[{c['severity'][:4]}]"
            print(f"  [{status}] {sev} {c['code']}: {c['name']}")
            if c["message"]: print(f"         {c['message']}")
        
        print(f"\n{'=' * 70}")
        print(f"  SCORE: {s['passed']}/{s['total']} passed ({s['score_pct']}%)")
        for cat, stats in result["categories"].items():
            print(f"  {cat}: {stats['passed']}/{stats['total']}")
        print(f"{'=' * 70}")
        
        fail_sev = [c for c in result["checks"] if not c["passed"] and c["severity"] == "Required"]
        if fail_sev:
            print(f"\n  ⚠ {len(fail_sev)} REQUIRED checks failed:")
            for c in fail_sev:
                print(f"    FAIL {c['code']}: {c['name']} — {c['message']}")
        else:
            print("\n  ✅ All required checks pass!")
    
    sys.exit(0 if result["summary"]["failed"] == 0 else 1)

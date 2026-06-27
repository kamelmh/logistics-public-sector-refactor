"""fix_thesis_all.py — Comprehensive thesis DOCX fixer v4 (Single-Section + Full RTL)
Part of Academix v13.2 build pipeline
Usage: python fix_thesis_all.py <path/to.docx> --save

ANTI-STUFFED + FULL RTL STRATEGY:
  ─────────────────────────────────────────────────────────────────────────────
  1. Style-level font/size/spacing  → no run-level overrides → no bloat
  2. Paragraph pPr: <w:bidi w:val="1"/> + <w:jc val="right"/>  → every para
  3. Run rPr: <w:bidi/> + <w:rtl/> on Arabic runs only → correct glyph order
  4. Footnote XML (zip-level): pPr bidi/jc + rPr bidi/rtl on Arabic runs
  5. Footer injection (zip-level): PAGE field footer, no cached "1" value
  6. Clear run-level font/size overrides that shadow the style cascade
  ─────────────────────────────────────────────────────────────────────────────

Fixes applied (fix_docx_sections.py handles page numbering + single section):
 1.  Table column widths: proportional, content-aware
 2.  Table borders: thin gridlines on all tables
 3.  Table cell padding: compact margins
 4.  Style-level font (Traditional Arabic 14pt, 1.5 spacing, RTL)
      Styles covered: Normal, Body Text, Compact, List Paragraph, No Spacing,
                      Caption, Footnote Text, Annotation Text, Table Contents,
                      Heading 1/2/3 + French equivalents (Titre 1/2/3)
 5.  Paragraph pPr: <w:bidi w:val="1"/> + <w:jc val="right"/> on ALL body paras
 6.  Run rPr: <w:bidi/> + <w:rtl/> on runs containing Arabic characters
 7.  Empty paragraph cleanup (consecutive)
 8.  Footnote RTL: pPr + rPr + namespace declarations (zip-level)
 9.  Footer injection: proper PAGE field, RTL, Traditional Arabic font
     - Single section with titlePg (different first page)
     - footer1.xml = blank (cover/first page)
     - footer2.xml = PAGE field (all other pages)
     - Page numbering: decimal, start=1 (cover counts as page 1, no display)
 10. Word compat: ns0/ns1 → w/mc namespace fix
"""

import sys, os, zipfile, re, time as _time
from xml.etree import ElementTree as ET
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Namespace constants ────────────────────────────────────────────────────────
W_URI  = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
MC_URI = 'http://schemas.openxmlformats.org/markup-compatibility/2006'
REL_URI = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

EXTRA_NS = [
    ('mc',    MC_URI),
    ('w14',   'http://schemas.microsoft.com/office/word/2010/wordml'),
    ('w15',   'http://schemas.microsoft.com/office/word/2012/wordml'),
    ('w16se', 'http://schemas.microsoft.com/office/word/2015/wordml/symex'),
    ('w16cid','http://schemas.microsoft.com/office/word/2016/wordml/cid'),
    ('wp14',  'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing'),
]

# ── Golden formatting values ───────────────────────────────────────────────────
GOLDEN = {
    'bodyFont'    : 'Traditional Arabic',
    'bodySize'    : 14,
    'h1Size'      : 22,
    'h2Size'      : 18,
    'h3Size'      : 16,
    'footnoteSize': 12,
    'marginsCm'   : 2.5,
    'pageWidthCm' : 21.0,
    'pageHeightCm': 29.7,
}

# Styles to receive full body treatment (font + size + spacing + RTL)
BODY_STYLES = [
    'Normal', 'Compact', 'Body Text', 'List Paragraph', 'No Spacing',
    'Caption', 'Annotation Text', 'Table Contents', 'Table Grid',
]
FOOTNOTE_STYLES = ['Footnote Text', 'footnote text', 'Footnote Reference']
HEADING_SIZES = {
    'Heading 1': GOLDEN['h1Size'], 'Titre 1': GOLDEN['h1Size'],
    'Heading 2': GOLDEN['h2Size'], 'Titre 2': GOLDEN['h2Size'],
    'Heading 3': GOLDEN['h3Size'], 'Titre 3': GOLDEN['h3Size'],
}

# Arabic Unicode ranges (to detect Arabic text in runs)
def _is_arabic_run(text):
    return any('\u0600' <= c <= '\u06ff' or '\ufe70' <= c <= '\ufeff' for c in text)


# ── XML helpers ────────────────────────────────────────────────────────────────

def _safe_replace(tmp_path, target_path, retries=4):
    """Replace target_path with tmp_path, retrying on Windows file-lock errors."""
    for attempt in range(retries):
        try:
            os.replace(tmp_path, target_path)
            return
        except PermissionError:
            _time.sleep(1.5)
    os.replace(tmp_path, target_path)   # final attempt — raise if still locked


def _zip_replace(docx_path, replacements: dict):
    """Rewrite a DOCX zip replacing only the entries listed in replacements dict.
    replacements = {arcname: bytes_content}
    """
    import tempfile
    fd, tmp = tempfile.mkstemp(suffix='.tmp', prefix='thesis_',
                               dir=os.path.dirname(docx_path))
    os.close(fd)
    try:
        with zipfile.ZipFile(docx_path, 'r') as zin:
            with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.namelist():
                    if item in replacements:
                        zout.writestr(item, replacements[item])
                    else:
                        zout.writestr(item, zin.read(item))
        # Add any NEW entries (not in original zip)
        existing = set()
        with zipfile.ZipFile(docx_path, 'r') as zin:
            existing = set(zin.namelist())
        new_entries = {k: v for k, v in replacements.items() if k not in existing}
        if new_entries:
            with zipfile.ZipFile(tmp, 'a', zipfile.ZIP_DEFLATED) as zout:
                for arc, data in new_entries.items():
                    zout.writestr(arc, data)
        _safe_replace(tmp, docx_path)
    except Exception:
        if os.path.exists(tmp):
            try: os.remove(tmp)
            except: pass
        raise


def _set_spacing_pPr(pPr):
    """Set <w:spacing w:line="360" w:lineRule="auto"/> — 1.5 line spacing in twips."""
    existing = pPr.find(qn('w:spacing'))
    if existing is not None:
        existing.set(qn('w:line'), '360')
        existing.set(qn('w:lineRule'), 'auto')
    else:
        pPr.append(parse_xml('<w:spacing %s w:line="360" w:lineRule="auto"/>' % nsdecls('w')))


def _set_bidi_pPr(pPr):
    """Ensure <w:bidi w:val="1"/> and <w:jc w:val="right"/> are present in pPr."""
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        pPr.insert(0, parse_xml('<w:bidi %s w:val="1"/>' % nsdecls('w')))
    else:
        bidi.set(qn('w:val'), '1')
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        pPr.append(parse_xml('<w:jc %s w:val="right"/>' % nsdecls('w')))
    else:
        jc.set(qn('w:val'), 'right')


def _get_or_create_pPr(elem):
    """Get or create a <w:pPr> element as first child of elem."""
    pPr = elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml('<w:pPr %s/>' % nsdecls('w'))
        elem.insert(0, pPr)
    return pPr


def _set_run_rtl(r_elem):
    """Set <w:bidi/> + <w:rtl/> in a run's rPr (for Arabic runs)."""
    rPr = r_elem.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml('<w:rPr %s/>' % nsdecls('w'))
        r_elem.insert(0, rPr)
    if rPr.find(qn('w:bidi')) is None:
        rPr.insert(0, parse_xml('<w:bidi %s/>' % nsdecls('w')))
    if rPr.find(qn('w:rtl')) is None:
        rPr.append(parse_xml('<w:rtl %s/>' % nsdecls('w')))


def _apply_style_formatting(style, font_name, font_size_pt, line_spacing=True, rtl=True):
    """Apply font/size/spacing/RTL to a python-docx style object."""
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)
    pPr = _get_or_create_pPr(style.element)
    if line_spacing:
        _set_spacing_pPr(pPr)
    if rtl:
        _set_bidi_pPr(pPr)


# ── Fix 1: Page numbering ──────────────────────────────────────────────────────

def fix_page_numbering(doc, changes):
    """cover=none, TOC=decimal start=1, body/annexes=decimal CONTINUE (no start attr)."""
    formats = ['none', 'decimal', 'decimal', 'decimal']
    for i, sec in enumerate(doc.sections):
        sect_pr = sec._sectPr
        old = sect_pr.find(qn('w:pgNumType'))
        if old is not None:
            sect_pr.remove(old)
        fmt = formats[i] if i < len(formats) else 'decimal'
        # Only TOC (section 1) starts at 1; body (2) and annexes (3) CONTINUE
        start = ' w:start="1"' if i == 1 else ''
        pg = parse_xml('<w:pgNumType %s w:fmt="%s"%s/>' % (nsdecls('w'), fmt, start))
        changes['page_num_sec%d' % i] = 'fmt=%s%s' % (fmt, ' start=1' if start else ' continue')
        first = sect_pr.find(qn('w:type'))
        if first is not None:
            sect_pr.insert(list(sect_pr).index(first) + 1, pg)
        else:
            sect_pr.insert(0, pg)
    return changes


def fix_table_column_widths(doc, changes):
    avail = (GOLDEN['pageWidthCm'] - 2 * GOLDEN['marginsCm']) * 914400 / 2.54
    for ti, t in enumerate(doc.tables):
        cols = len(t.columns)
        if cols == 0:
            continue
        max_w = []
        for ci in range(cols):
            mc = 0
            for row in t.rows:
                if ci < len(row.cells):
                    txt = row.cells[ci].text.strip()
                    ar = sum(1 for c in txt if '\u0600' <= c <= '\u06ff')
                    mc = max(mc, ar * 0.9 + (len(txt) - ar) * 0.55)
            max_w.append(max(1, mc))
        total = sum(max_w)
        pad = 150000
        usable = avail - cols * pad
        widths = [max(int(usable * w / total), 500000) for w in max_w]
        wsum = sum(widths)
        if wsum > 0:
            widths = [int(w * usable / wsum) for w in widths]
        for row in t.rows:
            for ci in range(min(cols, len(row.cells))):
                tc = row.cells[ci]._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = parse_xml('<w:tcPr %s/>' % nsdecls('w'))
                    tc.insert(0, tcPr)
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcPr.append(parse_xml('<w:tcW %s w:w="%d" w:type="dxa"/>' % (nsdecls('w'), widths[ci])))
                else:
                    tcW.set(qn('w:w'), str(widths[ci]))
                    tcW.set(qn('w:type'), 'dxa')
        changes['table_widths_set'].append(ti)
    return changes


# ── Fix 3: Table borders ───────────────────────────────────────────────────────

def add_table_borders(doc, changes):
    ba = 'w:val="single" w:sz="4" w:space="0" w:color="000000"'
    bxml = ('<w:tblBorders %s>'
            '<w:top %s/><w:left %s/><w:bottom %s/>'
            '<w:right %s/><w:insideH %s/><w:insideV %s/>'
            '</w:tblBorders>' % (nsdecls('w'), ba, ba, ba, ba, ba, ba))
    for ti, t in enumerate(doc.tables):
        tbl = t._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = parse_xml('<w:tblPr %s/>' % nsdecls('w'))
            tbl.insert(0, tblPr)
        old = tblPr.find(qn('w:tblBorders'))
        if old is not None:
            tblPr.remove(old)
        tblPr.append(parse_xml(bxml))
        changes['table_borders_added'].append(ti)
    return changes


# ── Fix 4: Table cell padding ──────────────────────────────────────────────────

def fix_table_cell_padding(doc, changes):
    cm_xml = ('<w:tblCellMar %s>'
              '<w:top w:w="40" w:type="dxa"/><w:bottom w:w="40" w:type="dxa"/>'
              '<w:left w:w="60" w:type="dxa"/><w:right w:w="60" w:type="dxa"/>'
              '</w:tblCellMar>' % nsdecls('w'))
    for t in doc.tables:
        tbl = t._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = parse_xml('<w:tblPr %s/>' % nsdecls('w'))
            tbl.insert(0, tblPr)
        old = tblPr.find(qn('w:tblCellMar'))
        if old is not None:
            tblPr.remove(old)
        tblPr.append(parse_xml(cm_xml))
    changes['table_cell_margins_set'] = True
    return changes


# ── Fix 5: Style-level formatting ─────────────────────────────────────────────

def fix_styles(doc, changes):
    """Apply font/size/spacing/RTL to style definitions — NEVER individual runs."""

    # Body styles
    for sname in BODY_STYLES:
        try:
            s = doc.styles[sname]
            _apply_style_formatting(s, GOLDEN['bodyFont'], GOLDEN['bodySize'])
            changes['styles_updated'] += 1
        except KeyError:
            pass

    # Footnote styles
    for sname in FOOTNOTE_STYLES:
        try:
            s = doc.styles[sname]
            _apply_style_formatting(s, GOLDEN['bodyFont'], GOLDEN['footnoteSize'])
            changes['styles_updated'] += 1
        except KeyError:
            pass

    # Heading styles
    for sname, sz in HEADING_SIZES.items():
        try:
            s = doc.styles[sname]
            _apply_style_formatting(s, GOLDEN['bodyFont'], sz, line_spacing=False)
            s.font.bold = True
            changes['styles_updated'] += 1
        except KeyError:
            pass

    return changes


# ── Fix 6: Paragraph-level RTL (pPr bidi + jc) ────────────────────────────────

def _fix_para_elem_rtl(p_elem):
    """Apply bidi + jc to a raw paragraph XML element."""
    pPr = _get_or_create_pPr(p_elem)
    _set_bidi_pPr(pPr)


def fix_paragraph_rtl(doc, changes):
    """Set <w:bidi w:val="1"/> + <w:jc val="right"/> on EVERY body paragraph.

    Covers:
    - doc.paragraphs (body paragraphs)
    - table cell paragraphs (not included in doc.paragraphs)
    Skips: headers, footers, footnote containers, TOC entries.
    """
    skip_styles = ('Header', 'Footer', 'toc', 'TOC', 'table of figures', 'TableofFigures')
    count = 0

    # Body paragraphs
    for p in doc.paragraphs:
        sname = (p.style.name or '') if p.style else ''
        if any(k in sname for k in skip_styles) and 'Caption' not in sname:
            continue
        _fix_para_elem_rtl(p._element)
        count += 1

    # Table cell paragraphs (missed by doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    sname = (p.style.name or '') if p.style else ''
                    if any(k in sname for k in skip_styles) and 'Caption' not in sname:
                        continue
                    _fix_para_elem_rtl(p._element)
                    count += 1

    changes['para_rtl_fixed'] = count
    return changes


# ── Fix 7: Run-level RTL for Arabic text ──────────────────────────────────────

def _fix_runs_in_para(p, is_body, sname, rtl_counter, clear_counter):
    """Apply Arabic RTL + bloat cleanup to runs in a paragraph element."""
    PURE_FORMAT_TAGS = {qn('w:rFonts'), qn('w:sz'), qn('w:szCs'),
                        qn('w:b'), qn('w:color'), qn('w:bCs')}
    for r in p.runs:
        txt = r.text or ''
        if _is_arabic_run(txt):
            _set_run_rtl(r._r)
            rtl_counter[0] += 1
        elif is_body:
            rPr = r._r.find(qn('w:rPr'))
            if rPr is not None:
                # Always explicitly clear font size overrides if this is a body or footnote style
                if is_body or sname in FOOTNOTE_STYLES:
                    sz = rPr.find(qn('w:sz'))
                    if sz is not None:
                        rPr.remove(sz)
                    szCs = rPr.find(qn('w:szCs'))
                    if szCs is not None:
                        rPr.remove(szCs)
                    # If rPr is empty after clearing, remove it
                    if not list(rPr):
                        r._r.remove(rPr)
                        clear_counter[0] += 1
                elif list(rPr) and all(c.tag in PURE_FORMAT_TAGS for c in list(rPr)):
                    # Original logic for other pure format tags
                    if not list(rPr):
                        r._r.remove(rPr)
                        clear_counter[0] += 1


def fix_run_rtl(doc, changes):
    """Set <w:bidi/> + <w:rtl/> on any run containing Arabic characters.

    Covers body paragraphs AND table cell paragraphs.
    Also clears pure-formatting rPr overrides on body paragraphs.
    """
    skip_styles = ('Header', 'Footer', 'toc', 'TOC', 'table of figures', 'TableofFigures')
    rtl_c = [0]
    clr_c = [0]

    # Body paragraphs
    for p in doc.paragraphs:
        sname = (p.style.name or '') if p.style else ''
        if any(k in sname for k in skip_styles):
            continue
        _fix_runs_in_para(p, sname in BODY_STYLES, sname, rtl_c, clr_c)

    # Table cell paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    sname = (p.style.name or '') if p.style else ''
                    if any(k in sname for k in skip_styles):
                        continue
                    _fix_runs_in_para(p, sname in BODY_STYLES, sname, rtl_c, clr_c)

    changes['run_rtl_fixed'] = rtl_c[0]
    changes['run_rpr_cleared'] = clr_c[0]
    return changes


# ── Fix 8: Empty paragraph cleanup ────────────────────────────────────────────

def clean_empty_paragraphs(doc, changes):
    empty = [i for i, p in enumerate(doc.paragraphs) if not p.text.strip()]
    to_rm = [i for i in empty if (i + 1) in set(empty)]
    for idx in reversed(to_rm):
        doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)
        changes['empty_paras_removed'] += 1
    return changes


# ── Fix 9: Footnote RTL (zip-level) ───────────────────────────────────────────

def fix_footnotes_zip(docx_path, changes):
    """Fix footnote RTL at the zip level — pPr bidi/jc + run-level bidi/rtl.

    Handles both:
    - paragraphs that have <w:pPr> (add bidi/jc if missing)
    - paragraphs that have NO <w:pPr> (inject one before first child)

    Also adds missing namespace declarations that Word requires.
    """
    fn_file = 'word/footnotes.xml'
    with zipfile.ZipFile(docx_path, 'r') as z:
        if fn_file not in z.namelist():
            return changes
        raw = z.read(fn_file).decode('utf-8')

    modified = False

    # ── 9a: Ensure namespace declarations ──
    for prefix, uri in EXTRA_NS:
        if ('xmlns:%s=' % prefix) not in raw:
            decl_end = raw.find('?>')
            idx = raw.find('>', decl_end + 2) if decl_end > 0 else raw.find('>')
            if idx > 0:
                raw = raw[:idx] + ' xmlns:%s="%s"' % (prefix, uri) + raw[idx:]
                modified = True
                changes['fn_ns_fixes'] += 1

    # ── 9b: Fix existing jc values ──
    jc_pat = re.compile(r'<w:jc\s+w:val="[^"]*"\s*/>')
    for m in jc_pat.finditer(raw):
        if m.group(0) != '<w:jc w:val="right"/>':
            modified = True
            changes['fn_rtl_fixes'] += 1
    raw = jc_pat.sub('<w:jc w:val="right"/>', raw)

    # ── 9c: Fix existing pPr blocks (add bidi/jc if missing) ──
    ppr_pat = re.compile(r'(<w:pPr(?:\s[^>]*)?>)(.*?)</w:pPr>', re.DOTALL)
    def _fix_ppr(m):
        head, inner = m.group(1), m.group(2)
        extra = ''
        if 'w:bidi' not in inner:
            extra += '<w:bidi w:val="1"/>'
            changes['fn_rtl_fixes'] += 1
        if 'w:jc' not in inner:
            extra += '<w:jc w:val="right"/>'
            changes['fn_rtl_fixes'] += 1
        if extra:
            return head + inner + extra + '</w:pPr>'
        return m.group(0)
    new_raw = ppr_pat.sub(_fix_ppr, raw)
    if new_raw != raw:
        modified = True
        raw = new_raw

    # ── 9d: Inject pPr into <w:p> blocks that have NO <w:pPr> ──
    # Match <w:p> that does NOT contain <w:pPr> inside it
    # We look for <w:p ...> blocks and check for pPr
    p_pat = re.compile(r'(<w:p\b[^>]*>)(.*?)(</w:p>)', re.DOTALL)
    def _inject_ppr(m):
        open_tag, inner, close_tag = m.group(1), m.group(2), m.group(3)
        if '<w:pPr' in inner:
            return m.group(0)   # already has pPr, handled by step 9c
        # Skip separator footnotes (id=0 or id=-1)
        if 'w:type="separator"' in open_tag or 'w:type="continuationSeparator"' in open_tag:
            return m.group(0)
        ppr_inject = '<w:pPr><w:bidi w:val="1"/><w:jc w:val="right"/></w:pPr>'
        changes['fn_rtl_fixes'] += 1
        return open_tag + ppr_inject + inner + close_tag
    new_raw = p_pat.sub(_inject_ppr, raw)
    if new_raw != raw:
        modified = True
        raw = new_raw

    # ── 9e: Add bidi/rtl to runs containing Arabic text ──
    run_pat = re.compile(r'(<w:r\b[^>]*>)(.*?)(</w:r>)', re.DOTALL)
    def _fix_run(m):
        open_tag, inner, close_tag = m.group(1), m.group(2), m.group(3)
        # Extract text content
        texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', inner, re.DOTALL)
        txt = ''.join(texts)
        if not _is_arabic_run(txt):
            return m.group(0)
        # Fix or inject rPr
        rpr_m = re.search(r'(<w:rPr>)(.*?)(</w:rPr>)', inner, re.DOTALL)
        if rpr_m:
            rpr_inner = rpr_m.group(2)
            extra = ''
            if '<w:bidi' not in rpr_inner:
                extra += '<w:bidi/>'
                changes['fn_run_rtl'] += 1
            if '<w:rtl' not in rpr_inner:
                extra += '<w:rtl/>'
                changes['fn_run_rtl'] += 1
            if extra:
                new_inner = inner[:rpr_m.start()] + '<w:rPr>' + rpr_inner + extra + '</w:rPr>' + inner[rpr_m.end():]
                return open_tag + new_inner + close_tag
        else:
            # Insert rPr before first <w:t
            t_pos = inner.find('<w:t')
            if t_pos >= 0:
                new_inner = inner[:t_pos] + '<w:rPr><w:bidi/><w:rtl/></w:rPr>' + inner[t_pos:]
                changes['fn_run_rtl'] += 1
                return open_tag + new_inner + close_tag
        return m.group(0)
    new_raw = run_pat.sub(_fix_run, raw)
    if new_raw != raw:
        modified = True
        raw = new_raw

    if modified:
        _zip_replace(docx_path, {fn_file: raw.encode('utf-8')})

    return changes


# ── Fix 10: Footer injection (PAGE field, RTL, Traditional Arabic) ─────────────

_FOOTER_XML = '''\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:ftr xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
       xmlns:cx="http://schemas.microsoft.com/office/drawing/2014/chartex"
       xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
       xmlns:o="urn:schemas-microsoft-com:office:office"
       xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
       xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
       xmlns:v="urn:schemas-microsoft-com:vml"
       xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
       xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
       xmlns:w10="urn:schemas-microsoft-com:office:word"
       xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
       xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
       xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"
       xmlns:w16se="http://schemas.microsoft.com/office/word/2015/wordml/symex"
       xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
       xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"
       xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"
       xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
       mc:Ignorable="w14 w15 w16se wp14">
  <w:p>
    <w:pPr>
      <w:jc w:val="center"/>
      <w:bidi w:val="1"/>
      <w:rPr>
        <w:rFonts w:ascii="Traditional Arabic" w:hAnsi="Traditional Arabic"
                  w:cs="Traditional Arabic"/>
        <w:sz w:val="24"/>
        <w:szCs w:val="24"/>
      </w:rPr>
    </w:pPr>
    <w:r>
      <w:rPr>
        <w:rFonts w:ascii="Traditional Arabic" w:hAnsi="Traditional Arabic"
                  w:cs="Traditional Arabic"/>
        <w:sz w:val="24"/>
        <w:szCs w:val="24"/>
      </w:rPr>
      <w:fldChar w:fldCharType="begin"/>
    </w:r>
    <w:r>
      <w:rPr>
        <w:rFonts w:ascii="Traditional Arabic" w:hAnsi="Traditional Arabic"
                  w:cs="Traditional Arabic"/>
        <w:sz w:val="24"/>
        <w:szCs w:val="24"/>
      </w:rPr>
      <w:instrText xml:space="preserve"> PAGE </w:instrText>
    </w:r>
    <w:r>
      <w:rPr>
        <w:rFonts w:ascii="Traditional Arabic" w:hAnsi="Traditional Arabic"
                  w:cs="Traditional Arabic"/>
        <w:sz w:val="24"/>
        <w:szCs w:val="24"/>
      </w:rPr>
      <w:fldChar w:fldCharType="separate"/>
    </w:r>
    <w:r>
      <w:rPr>
        <w:rFonts w:ascii="Traditional Arabic" w:hAnsi="Traditional Arabic"
                  w:cs="Traditional Arabic"/>
        <w:sz w:val="24"/>
        <w:szCs w:val="24"/>
      </w:rPr>
      <w:fldChar w:fldCharType="end"/>
    </w:r>
  </w:p>
</w:ftr>
'''

_FOOTER_BLANK_XML = '''\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:ftr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:p><w:pPr><w:jc w:val="center"/></w:pPr></w:p>
</w:ftr>
'''

def inject_footer(docx_path, changes):
    """Inject footer for SINGLE-SECTION document with 'different first page'.

    Strategy:
    - ONE section with titlePg (different first page) enabled
    - footer1.xml = blank (used for first page / cover via w:type="first")
    - footer2.xml = PAGE field (used for default pages via w:type="default")
    - Page numbering: decimal, start=1 (cover counts as page 1 but no display)
    - Uses <w:fldChar> based field — never a cached fldSimple value
      so Word NEVER renders "PAGE1" or a stale number
    """
    with zipfile.ZipFile(docx_path, 'r') as z:
        names = z.namelist()
        rels_raw = z.read('word/_rels/document.xml.rels').decode('utf-8')
        doc_raw  = z.read('word/document.xml').decode('utf-8')
        ct_raw   = z.read('[Content_Types].xml').decode('utf-8')

    # ── Add content type entries for footer1.xml and footer2.xml ──
    ft_ct = 'application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml'
    new_ct = ct_raw
    for part in ['/word/footer1.xml', '/word/footer2.xml']:
        if part not in new_ct:
            new_ct = new_ct.replace(
                '</Types>',
                f'<Override PartName="{part}" ContentType="{ft_ct}"/></Types>'
            )

    # ── Add relationship entries ──
    new_rels = rels_raw
    # Remove all existing footer relationships
    new_rels = re.sub(r'<Relationship[^>]*Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer"[^/>]*/>', '', new_rels)
    new_rels = re.sub(r'<Relationship[^>]*Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer">.*?</Relationship>', '', new_rels)
    
    footer_rel = '<Relationship Id="{rid}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer" Target="{target}"/>'
    
    # footer1.xml = blank (first page)
    if 'rIdFooter1' not in new_rels:
        new_rels = new_rels.replace(
            '</Relationships>',
            footer_rel.format(rid='rIdFooter1', target='footer1.xml') + '</Relationships>'
        )
    
    # footer2.xml = PAGE field (default pages)
    if 'rIdFooter2' not in new_rels:
        new_rels = new_rels.replace(
            '</Relationships>',
            footer_rel.format(rid='rIdFooter2', target='footer2.xml') + '</Relationships>'
        )

    # ── Wire sectPr footerReference in document.xml ──
    # Remove all existing footerReference elements
    new_doc = re.sub(r'<w:footerReference[^/]*/>', '', doc_raw)
    new_doc = re.sub(r'<w:footerReference[^>]*></w:footerReference>', '', new_doc)

    # Find the single sectPr and wire BOTH footer references:
    # - default -> footer2 (PAGE field)
    # - first   -> footer1 (blank for cover)
    def _wire_section(m):
        s = m.group(0)
        if 'w:footerReference' in s:
            return s
        ref_default = f'<w:footerReference w:type="default" r:id="rIdFooter2" xmlns:r="{REL_URI}"/>'
        ref_first = f'<w:footerReference w:type="first" r:id="rIdFooter1" xmlns:r="{REL_URI}"/>'
        return s.replace('</w:sectPr>', ref_default + ref_first + '</w:sectPr>')

    new_doc = re.sub(
        r'<w:sectPr\b[^>]*>.*?</w:sectPr>',
        _wire_section,
        new_doc,
        flags=re.DOTALL
    )

    # Prepare replacements
    replacements = {
        'word/footer1.xml': _FOOTER_BLANK_XML.encode('utf-8'),
        'word/footer2.xml': _FOOTER_XML.encode('utf-8'),
        'word/_rels/document.xml.rels': new_rels.encode('utf-8'),
        'word/document.xml':       new_doc.encode('utf-8'),
        '[Content_Types].xml':     new_ct.encode('utf-8'),
    }
    _zip_replace(docx_path, replacements)
    changes['footer_injected'] = True
    return changes


# ── Fix 11: Word compat — ns0/ns1 fix ─────────────────────────────────────────

def fix_word_compat(docx_path, changes):
    """Fix python-docx ns0/ns1 namespace prefix pollution in footnotes/endnotes."""
    targets = ['word/footnotes.xml', 'word/endnotes.xml']
    replacements = {}
    with zipfile.ZipFile(docx_path, 'r') as z:
        names = z.namelist()
        for tgt in targets:
            if tgt not in names:
                continue
            raw = z.read(tgt).decode('utf-8')
            orig = raw
            raw = raw.replace('ns1:Ignorable', 'mc:Ignorable')
            raw = re.sub(r'xmlns:ns0="([^"]+)"', r'xmlns:w="\1"', raw)
            raw = re.sub(r'xmlns:ns1="([^"]+)"', r'xmlns:mc="\1"', raw)
            raw = re.sub(r'\bns0:', 'w:', raw)
            raw = re.sub(r'\bns1:', 'mc:', raw)
            for prefix, uri in EXTRA_NS:
                if ('xmlns:%s=' % prefix) not in raw:
                    m = re.search(r'<w?:footnotes[\s>]', raw)
                    if m:
                        gt = raw.find('>', m.start())
                        if gt > 0:
                            raw = raw[:gt] + ' xmlns:%s="%s"' % (prefix, uri) + raw[gt:]
            if raw != orig:
                replacements[tgt] = raw.encode('utf-8')
    if replacements:
        _zip_replace(docx_path, replacements)
    changes['compat_fixes'] = len(replacements)
    return changes


# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    if len(sys.argv) < 2:
        print('Usage: python fix_thesis_all.py <path/to.docx> [--save]', file=sys.stderr)
        sys.exit(1)

    path = sys.argv[1]
    save = '--save' in sys.argv

    if not os.path.exists(path):
        print('[ERROR] File not found: %s' % path, file=sys.stderr)
        sys.exit(1)

    changes = {
        'styles_updated'      : 0,
        'para_rtl_fixed'      : 0,
        'run_rtl_fixed'       : 0,
        'run_rpr_cleared'     : 0,
        'table_widths_set'    : [],
        'table_borders_added' : [],
        'table_cell_margins_set': False,
        'fn_rtl_fixes'        : 0,
        'fn_ns_fixes'         : 0,
        'fn_run_rtl'          : 0,
        'empty_paras_removed' : 0,
        'footer_injected'     : False,
        'compat_fixes'        : 0,
    }

    sep = '=' * 62
    print(sep)
    print('ACADEMIX — Full Thesis Fixer v4 (Anti-Stuffed + Full RTL)')
    print(sep)
    print('File : %s' % path)
    print('Mode : %s' % ('SAVE' if save else 'DRY RUN'))
    print()

    doc = Document(path)

    print('[1/10] Table column widths ...')
    changes = fix_table_column_widths(doc, changes)
    print('  %d tables sized' % len(changes['table_widths_set']))

    print('[2/10] Table borders ...')
    changes = add_table_borders(doc, changes)
    print('  %d tables bordered' % len(changes['table_borders_added']))

    print('[3/10] Table cell padding ...')
    changes = fix_table_cell_padding(doc, changes)

    print('[4/10] Style-level font / spacing / RTL ...')
    changes = fix_styles(doc, changes)
    print('  %d style definitions updated' % changes['styles_updated'])

    print('[5/10] Paragraph RTL (pPr bidi + jc) ...')
    changes = fix_paragraph_rtl(doc, changes)
    print('  %d paragraphs set RTL' % changes['para_rtl_fixed'])

    print('[6/10] Run-level Arabic RTL + bloat cleanup ...')
    changes = fix_run_rtl(doc, changes)
    print('  Arabic runs fixed: %d | Body rPr cleared: %d' % (
        changes['run_rtl_fixed'], changes['run_rpr_cleared']))

    print('[7/10] Empty paragraph cleanup ...')
    changes = clean_empty_paragraphs(doc, changes)
    print('  %d removed' % changes['empty_paras_removed'])

    if save:
        doc.save(path)
        print('  [Saved doc-level changes]')

    print('[8/10] Footnote RTL + namespace fix (zip-level) ...')
    if save:
        changes = fix_footnotes_zip(path, changes)
    print('  pPr fixes: %d | run RTL: %d | NS fixes: %d' % (
        changes['fn_rtl_fixes'], changes['fn_run_rtl'], changes['fn_ns_fixes']))

    print('[9/10] Footer injection (PAGE field, no PAGE1 bug) ...')
    if save:
        changes = inject_footer(path, changes)
    print('  Footer injected: %s' % changes['footer_injected'])

    print('[10/10] Word namespace compat (ns0/ns1 → w/mc) ...')
    if save:
        changes = fix_word_compat(path, changes)
    print('  Files fixed: %d' % changes['compat_fixes'])

    print()
    print(sep)
    print('SUMMARY')
    print(sep)
    for k, v in changes.items():
        if isinstance(v, list):
            print('  %-30s %d items' % (k+':', len(v)))
        else:
            print('  %-30s %s' % (k+':', v))
    fsize = os.path.getsize(path)
    print()
    print('  File size: %.1f KB' % (fsize / 1024))
    print('  Mode: %s' % ('SAVED' if save else 'DRY RUN — use --save to apply'))
    if not save:
        print('\n  To apply:  python fix_thesis_all.py "%s" --save' % path)
    return 0


if __name__ == '__main__':
    sys.exit(main())

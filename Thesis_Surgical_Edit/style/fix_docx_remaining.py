"""fix_docx_remaining.py — Footnote RTL, page numbering, TOC (remaining gaps)
Usage: python fix_docx_remaining.py <path/to.docx> [--save]
"""

import sys, os, json, copy, zipfile, shutil, tempfile
from xml.etree import ElementTree as ET
from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Cm, Pt

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"


def fix_footnotes_rtl(path):
    """Set RTL alignment (right justify) on all footnotes"""
    fn_file = "word/footnotes.xml"
    modified = False
    fn_count = 0
    fn_fixed = 0

    with zipfile.ZipFile(path, "r") as z:
        if fn_file not in z.namelist():
            return 0, 0
        tree = ET.parse(z.open(fn_file))
        root = tree.getroot()

    for fn in root.findall(f".//{W}footnote", {f"w": W_NS}):
        fn_id = fn.attrib.get(f"{W}id", "")
        if fn_id in ("0", "-1"):
            continue
        fn_count += 1
        for p in fn.findall(f".//{W}p", {f"w": W_NS}):
            pPr = p.find(f"{W}pPr", {f"w": W_NS})
            if pPr is None:
                pPr = ET.SubElement(p, f"{W}pPr")
                # Move to front
                p.remove(pPr)
                p.insert(0, pPr)
            jc = pPr.find(f"{W}jc", {f"w": W_NS})
            if jc is not None:
                val = jc.attrib.get(f"{W}val", "")
                if val != "right":
                    jc.attrib[f"{W}val"] = "right"
                    fn_fixed += 1
                    modified = True
            else:
                jc = ET.SubElement(pPr, f"{W}jc")
                jc.attrib[f"{W}val"] = "right"
                fn_fixed += 1
                modified = True

    if modified:
        raw = ET.tostring(root, encoding="unicode")
        with zipfile.ZipFile(path, "r") as z:
            items = z.namelist()
            # We'll write to a temp file and swap
        tmp = path + ".tmp"
        with zipfile.ZipFile(path, "r") as zin:
            with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.namelist():
                    if item == fn_file:
                        zout.writestr(item, raw)
                    else:
                        zout.writestr(item, zin.read(item))
        os.replace(tmp, path)

    return fn_count, fn_fixed


def fix_page_numbering(path):
    """Apply section-specific page numbering:
    Section 1 (cover): none
    Section 2 (front matter): lowerRoman
    Section 3 (body) .. end: decimal
    """
    doc = Document(path)
    modified = False

    for i, sec in enumerate(doc.sections):
        sect_pr = sec._sectPr
        # Remove existing pgNumType
        existing = sect_pr.find(qn("w:pgNumType"))
        if existing is not None:
            sect_pr.remove(existing)

        if i == 0:
            # Cover: no page numbers
            pg_num = parse_xml(f'<w:pgNumType {nsdecls("w")} w:fmt="none"/>')
        elif i == 1:
            # Front matter: roman numerals
            pg_num = parse_xml(f'<w:pgNumType {nsdecls("w")} w:fmt="lowerRoman"/>')
        else:
            # Body + rest: arabic
            pg_num = parse_xml(f'<w:pgNumType {nsdecls("w")} w:fmt="decimal"/>')

        sect_pr.insert(0, pg_num)
        modified = True

    if modified:
        doc.save(path)
    return modified


def show_heading4(path):
    """Show structure of Heading 4 paragraphs"""
    doc = Document(path)
    h4 = [(i, p.text[:60].strip()) for i, p in enumerate(doc.paragraphs)
          if "Heading 4" in p.style.name]
    return h4


def main():
    if len(sys.argv) < 2:
        print("Usage: python fix_docx_remaining.py <path/to.docx> [--save]", file=sys.stderr)
        sys.exit(1)

    path = sys.argv[1]
    save = "--save" in sys.argv

    if not os.path.exists(path):
        print(f"[ERROR] File not found: {path}", file=sys.stderr)
        sys.exit(1)

    # 1. Footnote RTL
    print("--- Footnotes RTL ---")
    fn_count, fn_fixed = fix_footnotes_rtl(path)
    print(f"  Footnotes: {fn_count} total, {fn_fixed} fixed to RTL")

    # 2. Page numbering
    print("--- Page Numbering ---")
    if save:
        ok = fix_page_numbering(path)
        print(f"  Section page numbering: {'applied' if ok else 'no change'}")
    else:
        print("  Use --save to apply page numbering")

    # 3. Heading 4 info
    print("--- Heading 4 Usage ---")
    h4 = show_heading4(path)
    print(f"  Heading 4 paragraphs: {len(h4)}")
    for idx, txt in h4[:5]:
        print(f"    [{idx}] {txt}")
    if len(h4) > 5:
        print(f"    ... and {len(h4) - 5} more")

    if save:
        print(f"\nSaved: {path}")
    else:
        print("\n[DRY RUN] Use --save to apply changes")


if __name__ == "__main__":
    main()

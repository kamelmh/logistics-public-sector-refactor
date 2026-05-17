"""build-bismillah.py — Academix v13.2
Generates a standalone Bismillah page in DOCX format.
"""
import os
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def main():
    doc = docx.Document()
    
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    for _ in range(12):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p.add_run("بسم الله الرحمن الرحيم")
    run.font.name = 'Traditional Arabic'
    run.font.size = Pt(36)
    run.bold = True
    
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), 'Traditional Arabic')
    rFonts.set(qn('w:hAnsi'), 'Traditional Arabic')
    rFonts.set(qn('w:cs'), 'Traditional Arabic')

    output_path = os.path.join(os.path.dirname(__file__), "bismillah-page.docx")
    doc.save(output_path)
    print(f"Bismillah page created: {output_path}")

if __name__ == '__main__':
    main()
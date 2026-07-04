"""fixers.orchestrator — Calls all fixers in sequence.

Usage:
    from fixers.orchestrator import run_all_fixers
    changes = run_all_fixers(docx_path, save=True)
"""

import os
from docx import Document

from .tables import fix_table_column_widths, add_table_borders, fix_table_cell_padding
from .styles import fix_styles
from .rtl import fix_paragraph_rtl, fix_run_rtl
from .empty_paras import clean_empty_paragraphs
from .footnotes import fix_footnotes_zip
from .footer import inject_footer
from .namespace import fix_word_compat


def _new_changes():
    return {
        'styles_updated'      : 0,
        'para_rtl_fixed'      : 0,
        'run_rtl_fixed'       : 0,
        'run_rpr_cleared'     : 0,
        'table_widths_set'    : [],
        'table_borders_added' : [],
        'table_cell_margins_set': False,
        'fn_rtl_fixes'        : 0,
        'fn_ns_fixes'         : 0,
        'fn_run_rtl'          : 0,
        'empty_paras_removed' : 0,
        'footer_injected'     : False,
        'compat_fixes'        : 0,
    }


def run_all_fixers(docx_path, save=False, verbose=True):
    """Run all 10 fixers in sequence. Returns changes dict."""
    if not os.path.exists(docx_path):
        raise FileNotFoundError(docx_path)

    changes = _new_changes()
    sep = '=' * 62

    if verbose:
        print(sep)
        print('ACADEMIX — Full Thesis Fixer (Modular)')
        print(sep)
        print('File : %s' % docx_path)
        print('Mode : %s' % ('SAVE' if save else 'DRY RUN'))
        print()

    doc = Document(docx_path)

    if verbose:
        print('[1/10] Table column widths ...')
    changes = fix_table_column_widths(doc, changes)
    if verbose:
        print('  %d tables sized' % len(changes['table_widths_set']))

    if verbose:
        print('[2/10] Table borders ...')
    changes = add_table_borders(doc, changes)
    if verbose:
        print('  %d tables bordered' % len(changes['table_borders_added']))

    if verbose:
        print('[3/10] Table cell padding ...')
    changes = fix_table_cell_padding(doc, changes)

    if verbose:
        print('[4/10] Style-level font / spacing / RTL ...')
    changes = fix_styles(doc, changes)
    if verbose:
        print('  %d style definitions updated' % changes['styles_updated'])

    if verbose:
        print('[5/10] Paragraph RTL (pPr bidi + jc) ...')
    changes = fix_paragraph_rtl(doc, changes)
    if verbose:
        print('  %d paragraphs set RTL' % changes['para_rtl_fixed'])

    if verbose:
        print('[6/10] Run-level Arabic RTL + bloat cleanup ...')
    changes = fix_run_rtl(doc, changes)
    if verbose:
        print('  Arabic runs fixed: %d | Body rPr cleared: %d' % (
            changes['run_rtl_fixed'], changes['run_rpr_cleared']))

    if verbose:
        print('[7/10] Empty paragraph cleanup ...')
    changes = clean_empty_paragraphs(doc, changes)
    if verbose:
        print('  %d removed' % changes['empty_paras_removed'])

    if save:
        doc.save(docx_path)
        if verbose:
            print('  [Saved doc-level changes]')

    if verbose:
        print('[8/10] Footnote RTL + namespace fix (zip-level) ...')
    if save:
        changes = fix_footnotes_zip(docx_path, changes)
    if verbose:
        print('  pPr fixes: %d | run RTL: %d | NS fixes: %d' % (
            changes['fn_rtl_fixes'], changes['fn_run_rtl'], changes['fn_ns_fixes']))

    if verbose:
        print('[9/10] Footer injection (PAGE field, no PAGE1 bug) ...')
    if save:
        changes = inject_footer(docx_path, changes)
    if verbose:
        print('  Footer injected: %s' % changes['footer_injected'])

    if verbose:
        print('[10/10] Word namespace compat (ns0/ns1 -> w/mc) ...')
    if save:
        changes = fix_word_compat(docx_path, changes)
    if verbose:
        print('  Files fixed: %d' % changes['compat_fixes'])

    if verbose:
        print()
        print(sep)
        print('SUMMARY')
        print(sep)
        for k, v in changes.items():
            if isinstance(v, list):
                print('  %-30s %d items' % (k+':', len(v)))
            else:
                print('  %-30s %s' % (k+':', v))
        fsize = os.path.getsize(docx_path)
        print()
        print('  File size: %.1f KB' % (fsize / 1024))
        print('  Mode: %s' % ('SAVED' if save else 'DRY RUN — use --save to apply'))

    return changes

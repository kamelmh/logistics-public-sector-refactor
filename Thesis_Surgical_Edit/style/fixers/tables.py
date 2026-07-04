"""fixers.tables — Table column widths, borders, and cell padding."""

from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from .constants import GOLDEN


def fix_table_column_widths(doc, changes):
    """Set proportional column widths based on Arabic/Latin character density."""
    avail = (GOLDEN['pageWidthCm'] - 2 * GOLDEN['marginsCm']) * 1440 / 2.54
    for ti, t in enumerate(doc.tables):
        cols = len(t.columns)
        if cols == 0:
            continue
        max_w = []
        for ci in range(cols):
            mc = 0
            for row in t.rows:
                if ci < len(row.cells):
                    txt = row.cells[ci].text.strip()
                    ar = sum(1 for c in txt if '\u0600' <= c <= '\u06ff')
                    mc = max(mc, ar * 0.9 + (len(txt) - ar) * 0.55)
            max_w.append(max(1, mc))
        total = sum(max_w)
        pad = 236
        usable = avail - cols * pad
        widths = [max(int(usable * w / total), 787) for w in max_w]
        wsum = sum(widths)
        if wsum > 0:
            widths = [int(w * usable / wsum) for w in widths]
        for row in t.rows:
            for ci in range(min(cols, len(row.cells))):
                tc = row.cells[ci]._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = parse_xml('<w:tcPr %s/>' % nsdecls('w'))
                    tc.insert(0, tcPr)
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcPr.append(parse_xml('<w:tcW %s w:w="%d" w:type="dxa"/>' % (nsdecls('w'), widths[ci])))
                else:
                    tcW.set(qn('w:w'), str(widths[ci]))
                    tcW.set(qn('w:type'), 'dxa')
        tbl = t._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is not None:
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is not None:
                tblW.set(qn('w:w'), '10000')
                tblW.set(qn('w:type'), 'pct')
        changes['table_widths_set'].append(ti)
    return changes


def add_table_borders(doc, changes):
    """Add thin black gridlines to all tables."""
    ba = 'w:val="single" w:sz="4" w:space="0" w:color="000000"'
    bxml = ('<w:tblBorders %s>'
            '<w:top %s/><w:left %s/><w:bottom %s/>'
            '<w:right %s/><w:insideH %s/><w:insideV %s/>'
            '</w:tblBorders>' % (nsdecls('w'), ba, ba, ba, ba, ba, ba))
    for ti, t in enumerate(doc.tables):
        tbl = t._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = parse_xml('<w:tblPr %s/>' % nsdecls('w'))
            tbl.insert(0, tblPr)
        old = tblPr.find(qn('w:tblBorders'))
        if old is not None:
            tblPr.remove(old)
        tblPr.append(parse_xml(bxml))
        changes['table_borders_added'].append(ti)
    return changes


def fix_table_cell_padding(doc, changes):
    """Set compact cell margins on all tables."""
    cm_xml = ('<w:tblCellMar %s>'
              '<w:top w:w="40" w:type="dxa"/><w:bottom w:w="40" w:type="dxa"/>'
              '<w:left w:w="60" w:type="dxa"/><w:right w:w="60" w:type="dxa"/>'
              '</w:tblCellMar>' % nsdecls('w'))
    for t in doc.tables:
        tbl = t._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = parse_xml('<w:tblPr %s/>' % nsdecls('w'))
            tbl.insert(0, tblPr)
        old = tblPr.find(qn('w:tblCellMar'))
        if old is not None:
            tblPr.remove(old)
        tblPr.append(parse_xml(cm_xml))
    changes['table_cell_margins_set'] = True
    return changes

"""fixers.rtl — Paragraph-level and run-level RTL for Arabic text."""

from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from .constants import BODY_STYLES, FOOTNOTE_STYLES, _is_arabic_run


def _fix_para_elem_rtl(p_elem):
    """Apply bidi + jc to a raw paragraph XML element."""
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml('<w:pPr %s/>' % nsdecls('w'))
        p_elem.insert(0, pPr)
    # bidi
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        pPr.insert(0, parse_xml('<w:bidi %s w:val="1"/>' % nsdecls('w')))
    else:
        bidi.set(qn('w:val'), '1')
    # jc
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        pPr.append(parse_xml('<w:jc %s w:val="right"/>' % nsdecls('w')))
    else:
        jc.set(qn('w:val'), 'right')


def _set_run_rtl(r_elem):
    """Set <w:bidi/> + <w:rtl/> in a run's rPr (for Arabic runs)."""
    rPr = r_elem.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml('<w:rPr %s/>' % nsdecls('w'))
        r_elem.insert(0, rPr)
    if rPr.find(qn('w:bidi')) is None:
        rPr.insert(0, parse_xml('<w:bidi %s/>' % nsdecls('w')))
    if rPr.find(qn('w:rtl')) is None:
        rPr.append(parse_xml('<w:rtl %s/>' % nsdecls('w')))


def _fix_runs_in_para(p, is_body, sname, rtl_counter, clear_counter):
    """Apply Arabic RTL + bloat cleanup to runs in a paragraph element."""
    PURE_FORMAT_TAGS = {qn('w:rFonts'), qn('w:sz'), qn('w:szCs'),
                        qn('w:b'), qn('w:color'), qn('w:bCs')}
    for r in p.runs:
        txt = r.text or ''
        if _is_arabic_run(txt):
            _set_run_rtl(r._r)
            rtl_counter[0] += 1
        elif is_body:
            rPr = r._r.find(qn('w:rPr'))
            if rPr is not None:
                if is_body or sname in FOOTNOTE_STYLES:
                    sz = rPr.find(qn('w:sz'))
                    if sz is not None:
                        rPr.remove(sz)
                    szCs = rPr.find(qn('w:szCs'))
                    if szCs is not None:
                        rPr.remove(szCs)
                    if not list(rPr):
                        r._r.remove(rPr)
                        clear_counter[0] += 1
                elif list(rPr) and all(c.tag in PURE_FORMAT_TAGS for c in list(rPr)):
                    if not list(rPr):
                        r._r.remove(rPr)
                        clear_counter[0] += 1


SKIP_STYLES = ('Header', 'Footer', 'toc', 'TOC', 'table of figures', 'TableofFigures')


def fix_paragraph_rtl(doc, changes):
    """Set bidi + jc on EVERY body paragraph. Covers doc.paragraphs + table cells."""
    count = 0
    for p in doc.paragraphs:
        sname = (p.style.name or '') if p.style else ''
        if any(k in sname for k in SKIP_STYLES) and 'Caption' not in sname:
            continue
        _fix_para_elem_rtl(p._element)
        count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    sname = (p.style.name or '') if p.style else ''
                    if any(k in sname for k in SKIP_STYLES) and 'Caption' not in sname:
                        continue
                    _fix_para_elem_rtl(p._element)
                    count += 1
    changes['para_rtl_fixed'] = count
    return changes


def fix_run_rtl(doc, changes):
    """Set bidi + rtl on Arabic runs. Also clears pure-formatting rPr overrides."""
    rtl_c = [0]
    clr_c = [0]
    for p in doc.paragraphs:
        sname = (p.style.name or '') if p.style else ''
        if any(k in sname for k in SKIP_STYLES):
            continue
        _fix_runs_in_para(p, sname in BODY_STYLES, sname, rtl_c, clr_c)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    sname = (p.style.name or '') if p.style else ''
                    if any(k in sname for k in SKIP_STYLES):
                        continue
                    _fix_runs_in_para(p, sname in BODY_STYLES, sname, rtl_c, clr_c)
    changes['run_rtl_fixed'] = rtl_c[0]
    changes['run_rpr_cleared'] = clr_c[0]
    return changes

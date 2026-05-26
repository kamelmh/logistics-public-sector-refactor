#!/usr/bin/env python
"""
docx_md_sync.py — ACADEMIX v13.2 DOCX ↔ MD sync & validation tool
===============================================================
Compares the golden desktop DOCX with the canonical MD source,
reports differences, and optionally syncs constants + YAML back to MD.

Modes:
  --verify          Compare DOCX vs MD, report differences (exit 0 = clean)
  --sync            Update MD YAML + constants from DOCX analysis
  --diff            Full structural comparison output (default)

Integration:
  Called automatically by build-thesis.ps1 after each build.
  Also usable standalone for manual cross-checks.

Usage:
  python style/docx_md_sync.py path/to/output.docx --verify
  python style/docx_md_sync.py path/to/output.docx --sync
  python style/docx_md_sync.py path/to/output.docx --diff
"""

import re, os, sys, json
from docx import Document
from docx.oxml.ns import qn

# ─── PROJECT ROOT ───────────────────────────────────────────────
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))
DEFAULT_MD = os.path.join(PROJECT_ROOT, 'Thesis_Surgical_Edit', 'Memoire_DSS_Logistique_ElBayadh.md')

# ─── GROUND TRUTH CONSTANTS (from thesis data) ───────────────────
GROUND_TRUTH = {
    'D': '1,546',
    'Q*': '176',
    'ROP': '212.4',
    'SS': '200',
    'LT': '2',
    'S': '801.45',
    'I': '20%',
    'VERSION': 'v13.2',
}

# ─── MD PARSING ─────────────────────────────────────────────────
def parse_md_yaml(md_path):
    """Extract YAML front matter from MD file."""
    yaml_block = {}
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Match YAML between --- markers
    m = re.match(r'^---\s*\n(.*?)\n---', content, re.DOTALL)
    if m:
        for line in m.group(1).split('\n'):
            if ':' in line:
                key, _, val = line.partition(':')
                yaml_block[key.strip()] = val.strip().strip('"').strip("'")
    return yaml_block

def parse_md_headings(md_path):
    """Extract all ATX headings (# through #####) from MD."""
    headings = []
    with open(md_path, 'r', encoding='utf-8') as f:
        for i, line in enumerate(f, 1):
            # Skip YAML front matter
            if i == 1 and line.startswith('---'):
                continue
            m = re.match(r'^(#{1,5})\s+(.+)$', line)
            if m:
                level = len(m.group(1))
                text = m.group(2).strip()
                # Skip TOC entries
                if text.startswith('فهرس') or text.startswith('قائمة'):
                    headings.append(('toc', level, text))
                else:
                    headings.append(('h', level, text))
    return headings

def parse_md_constants(md_path):
    """Extract known numeric constants from MD text."""
    constants = {}
    with open(md_path, 'r', encoding='utf-8') as f:
        text = f.read()
    
    # Remove YAML front matter
    text = re.sub(r'^---.*?---\s*', '', text, count=1, flags=re.DOTALL)
    
    patterns = [
        (r'[Dd]\s*=\s*([0-9,.]+)', 'D'),
        (r'Q\*?\s*=\s*([0-9,.]+)', 'Q*'),
        (r'ROP\s*=\s*([0-9,.]+)', 'ROP'),
        (r'[Mm]خزون\s+أمان.*?(\d+)\s*و[حح]دة', 'SS'),
        (r'SS\s*=\s*([0-9,.]+)', 'SS'),
        (r'[Aa]جل\s+[Tt]سليم.*?(\d+)\s*يوم', 'LT'),
        (r'LT\s*=\s*([0-9,.]+)', 'LT'),
        (r'S\s*=\s*([0-9,.]+)', 'S'),
        (r'I\s*=\s*([0-9,.]+)%?', 'I'),
    ]
    
    for pattern, key in patterns:
        m = re.search(pattern, text)
        if m:
            val = m.group(1).strip()
            if key == 'I' and '%' not in val and not text[m.end():m.end()+1] == '%':
                # Check if next char is %
                pass
            constants[key] = val
    
    return constants

def parse_md_sections(md_path):
    """Check which major sections exist in MD."""
    sections = {}
    with open(md_path, 'r', encoding='utf-8') as f:
        text = f.read()
    
    checks = {
        'yaml': bool(re.match(r'^---\s*\n', text)),
        'cover': bool(re.search(r'الجمهورية الجزائرية', text)),
        'dedication': bool(re.search(r'إهداء', text)),
        'thanks': bool(re.search(r'شكر\s*وتقدير', text)),
        'toc': bool(re.search(r'فهرس\s*المحتويات', text)),
        'tables_list': bool(re.search(r'قائمة\s*الجداول', text)),
        'abstract_ar': bool(re.search(r'##\s*الملخص', text)),
        'abstract_fr': bool(re.search(r'##\s*Résumé', text)),
        'ch1': bool(re.search(r'#\s*الفصل الأول', text)),
        'ch2': bool(re.search(r'#\s*الفصل الثاني', text)),
        'ch3': bool(re.search(r'#\s*الفصل الثالث', text)),
        'ch4': bool(re.search(r'#\s*الفصل الرابع', text)),
        'bibliography': bool(re.search(r'قائمة\s*المصادر\s*والمراجع', text)),
        'annexes': bool(re.search(r'الملاحق', text)),
        'glossary': bool(re.search(r'قائمة\s*المصطلحات', text)),
        'abbreviations': bool(re.search(r'جدول\s*المختصرات', text)),
    }
    
    for key, found in checks.items():
        sections[key] = 'present' if found else 'missing'
    
    return sections

# ─── DOCX PARSING ────────────────────────────────────────────────
def parse_docx_sections(docx_path):
    """Extract section info from DOCX."""
    doc = Document(docx_path)
    info = {}
    info['para_count'] = len(doc.paragraphs)
    info['table_count'] = len(doc.tables)
    info['section_count'] = len(doc.sections)
    
    # Section format
    info['sections'] = []
    NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    for i, sec in enumerate(doc.sections):
        pg = sec._sectPr.find(NS + 'pgNumType')
        fmt = pg.get(qn('w:fmt')) if pg is not None else 'default'
        info['sections'].append({'index': i, 'fmt': fmt})
    
    # First instance of key content markers
    info['content'] = {}
    body_text = ' '.join(p.text for p in doc.paragraphs)
    markers = {
        'yaml': 'title:',
        'cover': 'الجمهورية الجزائرية',
        'dedication': 'إهداء',
        'thanks': 'شكر',
        'toc': 'فهرس المحتويات',
        'tables_list': 'قائمة الجداول',
        'abstract_ar': 'الملخص',
        'abstract_fr': 'Résumé',
        'ch1': 'الفصل الأول',
        'ch2': 'الفصل الثاني',
        'ch3': 'الفصل الثالث',
        'ch4': 'الفصل الرابع',
        'bibliography': 'قائمة المصادر والمراجع',
        'annexes': 'الملاحق',
        'glossary': 'قائمة المصطلحات',
        'abbreviations': 'المختصرات',
    }
    for key, marker in markers.items():
        info['content'][key] = marker in body_text
    
    # Constants in DOCX
    info['constants'] = {}
    for key, expected in GROUND_TRUTH.items():
        found = expected in body_text
        info['constants'][key] = 'found' if found else 'not_found'
    
    return info

# ─── COMPARISON ──────────────────────────────────────────────────
def compare(md_path, docx_path):
    """Compare MD and DOCX, return structured diff."""
    md_yaml = parse_md_yaml(md_path)
    md_sections = parse_md_sections(md_path)
    md_headings = parse_md_headings(md_path)
    md_constants = parse_md_constants(md_path)
    
    docx_info = parse_docx_sections(docx_path)
    
    issues = []
    info = []
    
    # 1. YAML comparison
    info.append(('info', 'YAML', f'MD has {len(md_yaml)} fields: {list(md_yaml.keys())}'))
    for key in ['title', 'author', 'lang', 'dir']:
        if key in md_yaml:
            info.append(('ok', f'YAML.{key}', md_yaml[key]))
        else:
            issues.append(('warning', f'YAML.{key}', 'missing from MD'))
    
    # 2. Section coverage
    for section, status in md_sections.items():
        icon = 'ok' if status == 'present' else 'critical'
        if status == 'present':
            info.append((icon, f'Section.{section}', 'present'))
        else:
            issues.append((icon, f'Section.{section}', 'MISSING'))
    
    # 3. DOCX content coverage
    for key, present in docx_info['content'].items():
        # YAML is MD-only — skip DOCX comparison
        if key == 'yaml':
            continue
        if key in md_sections:
            md_status = md_sections.get(key, 'missing')
            if md_status == 'missing' and present:
                issues.append(('warning', f'Docx.{key}', f'in DOCX but NOT in MD — may need manual sync'))
    
    # 4. Constants comparison
    for key, expected in GROUND_TRUTH.items():
        md_val = md_constants.get(key)
        docx_status = docx_info['constants'].get(key, 'not_found')
        
        if md_val:
            # Normalize for numeric comparison
            md_clean = md_val.strip().replace(',', '').replace(' ', '').replace('%', '')
            exp_clean = expected.replace(',', '').replace(' ', '').replace('%', '')
            
            try:
                md_num = float(md_clean)
                exp_num = float(exp_clean)
                if abs(md_num - exp_num) <= 0.1:
                    info.append(('ok', f'Constant.{key}', f'MD={md_val} ≈ {expected} ✓'))
                elif abs(md_num - exp_num) < 1:
                    issues.append(('info', f'Constant.{key}', f'MD={md_val}, expected={expected} (close)'))
                else:
                    issues.append(('warning', f'Constant.{key}', f'MD={md_val}, expected={expected}'))
            except ValueError:
                # String comparison fallback
                if md_clean == exp_clean:
                    info.append(('ok', f'Constant.{key}', f'MD={md_val} ✓'))
                else:
                    issues.append(('warning', f'Constant.{key}', f'MD={md_val}, expected={expected}'))
        else:
            if docx_status == 'found':
                info.append(('info', f'Constant.{key}', f'in DOCX (not extracted from MD)'))
            else:
                issues.append(('warning', f'Constant.{key}', 'not found in either source'))
    
    # 5. Heading counts
    h_counts = {'h1': 0, 'h2': 0, 'h3': 0, 'h4': 0, 'toc': 0}
    for htype, level, text in md_headings:
        key = 'toc' if htype == 'toc' else f'h{level}'
        if key in h_counts:
            h_counts[key] += 1
    
    info.append(('ok', 'Headings', f'H1={h_counts["h1"]} H2={h_counts["h2"]} H3={h_counts["h3"]} H4={h_counts["h4"]}'))
    info.append(('ok', 'Tables', f'DOCX has {docx_info["table_count"]} tables'))
    
    # 6. DOCX structure
    for sec in docx_info['sections']:
        fmt = sec['fmt']
        fmt_desc = {
            'none': 'no page numbers',
            'lowerRoman': 'i/ii/iii (front matter)',
            'decimal': '1/2/3 (body)',
        }.get(fmt, fmt)
        info.append(('info', f'Section.{sec["index"]}', f'fmt={fmt} ({fmt_desc})'))
    
    return issues, info, md_yaml, md_headings

# ─── SYNC (UPDATE MD YAML + CONSTANTS) ──────────────────────────
def sync_md(md_path, issues=None):
    """Update MD file with corrected YAML and fix minor issues."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    changes = []
    
    # 1. Ensure correct YAML front matter
    yaml_template = """---
title: "نظام دعم القرار لتسيير المخزونات"
author: "ماحي كمال عبد الغني"
date: "2026-05-18"
lang: ar
dir: rtl
---"""
    
    # Replace existing YAML if corrupt
    if not content.startswith('---'):
        content = yaml_template + '\n' + content
        changes.append('Added missing YAML front matter')
    else:
        # Verify YAML has required fields
        m = re.match(r'^---\s*\n(.*?)\n---', content, re.DOTALL)
        if m:
            yaml_text = m.group(1)
            required = ['title:', 'author:', 'lang:', 'dir:']
            missing = [r for r in required if r not in yaml_text]
            if missing:
                # Rebuild YAML
                content = yaml_template + '\n' + content[content.index('---', 3)+3:]
                changes.append(f'Fixed YAML: added missing fields {missing}')
    
    # 2. Ensure constants are consistent
    # Only fix if MD has wrong values
    for key, expected in GROUND_TRUTH.items():
        # Don't auto-rewrite - just report
        pass
    
    # 3. Ensure section separators between ملخص and Résumé
    if '## الملخص\n' in content and '## Résumé\n' in content:
        # Check there's a proper separator between them
        abstract_pos = content.index('## الملخص\n')
        resume_pos = content.index('## Résumé\n')
        between = content[abstract_pos:resume_pos]
        if '---' not in between and '***' not in between:
            # Add separator before Résumé
            content = content.replace('## Résumé\n', '---\n## Résumé\n')
            changes.append('Added separator before Résumé')
    
    if changes:
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(content)
        return changes
    return []

# ─── OUTPUT ──────────────────────────────────────────────────────
def print_report(issues, info, md_path, docx_path):
    """Print formatted comparison report."""
    print('=' * 70)
    print('  ACADEMIX v13.2 — DOCX ↔ MD Sync Report')
    print('=' * 70)
    print(f'  MD:   {md_path}')
    print(f'  DOCX: {docx_path}')
    print()
    
    critical = [i for i in issues if i[0] == 'critical']
    warnings = [i for i in issues if i[0] == 'warning']
    
    # Info lines
    if info:
        print('  ── Structural Info ──')
        for icon, section, msg in info:
            ico = {'ok': '✓', 'info': '·'}.get(icon, '·')
            print(f'    [{ico}] {section}: {msg}')
        print()
    
    # Warnings
    if warnings:
        print('  ⚠  Warnings (non-critical):')
        for _, section, msg in warnings:
            print(f'    [⚠] {section}: {msg}')
        print()
    
    # Critical
    if critical:
        print('  ✘  CRITICAL:')
        for _, section, msg in critical:
            print(f'    [✘] {section}: {msg}')
        print()
    
    if not critical and not warnings:
        print('  ✓ MD and DOCX are in sync — no differences found.')
    
    print('=' * 70)
    return len(critical) == 0 and len(warnings) == 0

# ─── PATCH MD (inject missing sections from DOCX) ────────────────
def extract_section_from_docx(docx_path, section_name, heading_text, heading_as_normal=True):
    """Extract a named section (up to next heading) from DOCX as MD text.
    heading_as_normal=True means the section header is in Normal style, not Heading.
    """
    doc = Document(docx_path)
    lines = []
    capturing = False
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            if capturing:
                lines.append('')
            continue
        
        sn = p.style.name if p.style else ''
        is_heading = 'heading' in sn.lower() or 'titre' in sn.lower()
        
        # Detect section start: either Heading style with matching text OR Normal with exact match
        is_match = False
        if heading_as_normal:
            is_match = sn.lower() == 'normal' and heading_text in text
        else:
            is_match = is_heading and heading_text in text
        
        if not capturing and is_match:
            capturing = True
            continue
        
        if capturing:
            if is_heading and not heading_as_normal:
                break
            elif heading_as_normal and text in ('إهداء', 'شكر وتقدير', 'فهرس المحتويات', 'قائمة الجداول', 
                                                  'الملخص', 'Résumé', 'المقدمة العامة'):
                if not lines:  # Already at first content paragraph
                    continue
                break
            lines.append(text)
    
    # Remove leading/trailing empty lines
    while lines and not lines[0]:
        lines.pop(0)
    while lines and not lines[-1]:
        lines.pop()
    
    return '\n\n'.join(lines) if lines else None

def patch_md(md_path, docx_path, issues):
    """Inject missing sections from DOCX into MD."""
    changes = []
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Locate insertion point: after cover, before TOC or ملخص
    # The MD starts with YAML, then has ## فهرس المحتويات, ## قائمة الجداول, ***, ## الملخص, ---, ## Résumé, ---, # المقدمة العامة
    
    # Check which sections are missing
    critical_issues = [i for i in issues if i[0] == 'critical' and 'MISSING' in str(i)]
    
    for issue_type, section, msg in critical_issues:
        if 'dedication' in section:
            dedication_text = extract_section_from_docx(docx_path, 'dedication', 'إهداء')
            if dedication_text:
                # Insert after the first empty paragraph after cover, before ## فهرس المحتويات
                # Find ## فهرس المحتويات and insert before it
                target = '## فهرس المحتويات'
                if target in content:
                    insert = f'## إهداء\n\n{dedication_text}\n\n'
                    content = content.replace(target, insert + target)
                    changes.append(f'Injected dedication section ({len(dedication_text)} chars)')
        
        elif 'thanks' in section:
            thanks_text = extract_section_from_docx(docx_path, 'thanks', 'شكر')
            if thanks_text:
                # Insert after dedication, before ## فهرس المحتويات
                # Look for ## إهداء or ## فهرس المحتويات
                target = '## فهرس المحتويات'
                if target in content:
                    insert = f'## شكر وتقدير\n\n{thanks_text}\n\n'
                    content = content.replace(target, insert + target)
                    changes.append(f'Injected thanks section ({len(thanks_text)} chars)')
    
    if changes:
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(content)
    
    return changes

# ─── CLI ─────────────────────────────────────────────────────────
if __name__ == '__main__':
    import argparse
    
    parser = argparse.ArgumentParser(
        description='ACADEMIX DOCX ↔ MD sync & validation tool')
    parser.add_argument('docx', help='Path to DOCX file (output)')
    parser.add_argument('--md', default=DEFAULT_MD, help='Path to MD source file')
    parser.add_argument('--sync', action='store_true', help='Sync MD from DOCX (constants, YAML, minor fixes)')
    parser.add_argument('--patch-md', action='store_true', help='Inject missing sections (dedication, thanks) from DOCX into MD')
    parser.add_argument('--verify', action='store_true', help='Verify only (exit 0 = clean)')
    parser.add_argument('--diff', action='store_true', help='Full comparison report (default)')
    parser.add_argument('--json', action='store_true', help='JSON output')
    
    args = parser.parse_args()
    
    if not os.path.exists(args.docx):
        print(f'[ERROR] DOCX not found: {args.docx}')
        sys.exit(2)
    if not os.path.exists(args.md):
        print(f'[ERROR] MD not found: {args.md}')
        sys.exit(2)
    
    issues, info, md_yaml, md_headings = compare(args.md, args.docx)
    
    if args.patch_md:
        changes = patch_md(args.md, args.docx, issues)
        if changes:
            print(f'[PATCH] Applied {len(changes)} patches to MD:')
            for c in changes:
                print(f'  • {c}')
        else:
            print('[PATCH] No sections to patch.')
        # Re-verify after patching
        issues, info, md_yaml, md_headings = compare(args.md, args.docx)
    
    if args.sync:
        changes = sync_md(args.md, issues)
        if changes:
            print(f'[SYNC] Applied {len(changes)} changes to MD:')
            for c in changes:
                print(f'  • {c}')
        else:
            print('[SYNC] No changes needed.')
    
    if args.verify or args.diff or not (args.sync or args.patch_md):
        if args.json:
            result = {
                'issues': [{'type': t, 'section': s, 'msg': m} for t, s, m in issues],
                'info': [{'icon': i, 'section': s, 'msg': m} for i, s, m in info],
                'clean': len([i for i in issues if i[0] == 'critical']) == 0,
                'warnings': len([i for i in issues if i[0] == 'warning']),
            }
            print(json.dumps(result, ensure_ascii=False, indent=2))
        else:
            clean = print_report(issues, info, args.md, args.docx)
        
        if args.verify:
            critical_count = len([i for i in issues if i[0] == 'critical'])
            sys.exit(0 if critical_count == 0 else 1)

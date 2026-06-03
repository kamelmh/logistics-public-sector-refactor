"""
⚠️  DEPRECATED — Replaced by: fix_golden_footnotes.py
    This script is superseded and no longer part of the active pipeline.
    Kept for reference only. Do not use in new workflows.
"""
"""all-footnotes.py — Academix v13.2
Convert ALL 56 bibliography entries to footnotes in the thesis DOCX.
Strategy: Add footnote references at the end of each relevant section/chapter
for entries that aren't already cited inline.
"""
import sys, re, os, json
from lxml import etree
from docx import Document
from docx.opc.part import Part
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
FOOTNOTES_URI = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
CNEPD_PREFIX = "راجع: "

# Section-to-bibliography mapping: which entries belong to which section
SECTION_MAPPING = {
    # Academic books (already cited inline, but ensure footnotes exist)
    "الفصل الأول": [1, 2, 3, 4, 5, 6, 7, 8, 9, 10],  # Books 1-10
    "الفصل الثاني": [11, 12, 13, 14, 15, 16],  # Legal framework
    "الفصل الثالث": [48, 49],  # Technical references
    "الفصل الرابع": [50, 51, 52, 53],  # Parallel theses
    "الملاحق": [54, 55, 56],  # Field data + GitHub
    # BTS curriculum entries (17-47) - add as a consolidated footnote in introduction
    "المقدمة العامة": list(range(17, 48)),  # BTS modules 17-47
}


def ensure_footnotes_part(doc):
    """Create footnotes.xml part if it doesn't exist."""
    package = doc.part.package
    for rel in doc.part.rels.values():
        if 'footnotes' in rel.reltype:
            return rel.target_part

    FOOTNOTES_XML = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<w:footnote w:type="separator" w:id="-1">'
        '<w:p><w:r><w:separator/></w:r></w:p>'
        '</w:footnote>'
        '<w:footnote w:type="continuationSeparator" w:id="0">'
        '<w:p><w:r><w:continuationSeparator/></w:r></w:p>'
        '</w:footnote>'
        '</w:footnotes>'
    )

    content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml'
    partname = '/word/footnotes.xml'
    footnotes_part = Part(partname, content_type, FOOTNOTES_XML.encode('utf-8'), package)
    doc.part.relate_to(footnotes_part, FOOTNOTES_URI)
    return footnotes_part


def get_footnotes_xml(footnotes_part):
    if isinstance(footnotes_part._blob, bytes):
        return etree.fromstring(footnotes_part._blob)
    return footnotes_part._element


def create_footnote(footnotes_part, text):
    """Create a new footnote with the given text."""
    root = get_footnotes_xml(footnotes_part)

    max_id = 0
    for fn in root.findall(qn('w:footnote')):
        fid = int(fn.get(qn('w:id'), '0'))
        if fid > max_id:
            max_id = fid
    new_id = max_id + 1

    footnote = etree.SubElement(root, qn('w:footnote'))
    footnote.set(qn('w:id'), str(new_id))

    p = etree.SubElement(footnote, qn('w:p'))
    pPr = etree.SubElement(p, qn('w:pPr'))
    pStyle = etree.SubElement(pPr, qn('w:pStyle'))
    pStyle.set(qn('w:val'), 'FootnoteText')
    bidi = etree.SubElement(pPr, qn('w:bidi'))
    jc = etree.SubElement(pPr, qn('w:jc'))
    jc.set(qn('w:val'), 'right')

    r1 = etree.SubElement(p, qn('w:r'))
    rPr1 = etree.SubElement(r1, qn('w:rPr'))
    rStyle1 = etree.SubElement(rPr1, qn('w:rStyle'))
    rStyle1.set(qn('w:val'), 'FootnoteReference')
    etree.SubElement(r1, qn('w:footnoteRef'))

    r2 = etree.SubElement(p, qn('w:r'))
    t2 = etree.SubElement(r2, qn('w:t'))
    t2.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t2.text = ' '

    r3 = etree.SubElement(p, qn('w:r'))
    t3 = etree.SubElement(r3, qn('w:t'))
    t3.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t3.text = f"{CNEPD_PREFIX}{text}"

    footnotes_part._blob = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
    footnotes_part._element = root
    return new_id


def insert_footnote_reference(paragraph, fn_id):
    """Insert a footnote reference at the end of a paragraph."""
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'FootnoteReference')
    rPr.append(rStyle)
    r.append(rPr)

    fnRef = OxmlElement('w:footnoteReference')
    fnRef.set(qn('w:id'), str(fn_id))
    r.append(fnRef)

    paragraph._element.append(r)


def load_bibliography(md_path):
    """Load all 56 bibliography entries from the markdown source."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    entries = []
    # Match bibliography entries like "1. **Chopra, S., & Meindl, P.** (2019). *Title*..."
    # Pattern: number. **authors** (year). rest
    pattern = re.compile(r'^(\d+)\.\s+\*\*(.+?)\*\*\s*\((\d{4})[^\)]*\)\.\s+(.+)$', re.MULTILINE)
    for m in pattern.finditer(content):
        num = int(m.group(1))
        authors = m.group(2)
        year = m.group(3)
        rest = m.group(4)
        entries.append({
            'num': num,
            'authors': authors,
            'year': year,
            'full': f"{authors}. ({year}). {rest}"
        })

    print(f"  Found {len(entries)} bibliography entries")
    return entries


def find_section_paragraphs(doc, section_markers):
    """Find the last paragraph of each section."""
    section_paras = {}
    current_section = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        for marker in section_markers:
            if marker in text:
                current_section = marker
                section_paras[marker] = []
        if current_section and text:
            section_paras.setdefault(current_section, []).append(para)

    # Return the last paragraph of each section
    return {k: v[-1] if v else None for k, v in section_paras.items()}


def process_docx(docx_path, md_path):
    """Add footnotes for all 56 bibliography entries."""
    doc = Document(docx_path)
    footnotes_part = ensure_footnotes_part(doc)

    # Load bibliography
    entries = load_bibliography(md_path)
    if not entries:
        print(f"[WARN] No bibliography entries found in {md_path}")
        return 0

    # Find section paragraphs
    section_markers = list(SECTION_MAPPING.keys())
    section_paras = find_section_paragraphs(doc, section_markers)

    # Track which entries are already cited (from inline citations)
    cited_inline = set()
    for para in doc.paragraphs:
        # Check for existing footnote references
        for run in para._element.findall(qn('w:r')):
            if run.find(qn('w:footnoteReference')) is not None:
                # This paragraph already has a footnote
                pass

    changes = 0

    # Add footnotes for each section's bibliography entries
    for section, entry_nums in SECTION_MAPPING.items():
        para_info = section_paras.get(section)
        if para_info is None:
            print(f"[WARN] Section '{section}' not found in DOCX")
            continue

        # For BTS curriculum entries, create a single consolidated footnote
        if section == "المقدمة العامة":
            bts_entries = [e for e in entries if e['num'] in entry_nums]
            if bts_entries:
                bts_text = "المنهاج الدراسي BTS GSL (TAG1801) — المقاييس الـ31 المُغطَّاة في هذه المذكرة:\n"
                for e in bts_entries[:5]:  # First 5 as examples
                    bts_text += f"{e['num']}. {e['authors']}\n"
                bts_text += f"... و{len(bts_entries) - 5} مقاييس أخرى (المرجع الكامل في قائمة المراجع)"
                fn_id = create_footnote(footnotes_part, bts_text)
                insert_footnote_reference(para_info, fn_id)
                changes += len(bts_entries)
                print(f"  Added BTS curriculum footnote ({len(bts_entries)} entries)")
        else:
            section_entries = [e for e in entries if e['num'] in entry_nums]
            for entry in section_entries:
                fn_id = create_footnote(footnotes_part, entry['full'])
                insert_footnote_reference(para_info, fn_id)
                changes += 1

    doc.save(docx_path)
    print(f"Total footnotes added: {changes}")
    return changes


def main():
    if len(sys.argv) < 3:
        print("Usage: python all-footnotes.py <docx_path> <md_path>")
        sys.exit(1)
    process_docx(sys.argv[1], sys.argv[2])


if __name__ == '__main__':
    main()

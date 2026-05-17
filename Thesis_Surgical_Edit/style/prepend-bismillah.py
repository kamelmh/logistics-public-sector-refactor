"""prepend-bismillah.py — Academix v13.2
Prepends the bismillah-page.docx to the target DOCX.
"""
import sys, os
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree

def prepend_bismillah(target_path, bismillah_path):
    if not os.path.isabs(bismillah_path):
        bismillah_path = os.path.join(os.path.dirname(__file__), bismillah_path)
    
    doc_b = Document(bismillah_path)
    doc_t = Document(target_path)
    
    for p in doc_b.paragraphs:
        new_p = doc_t.add_paragraph(p.text)
        new_p.alignment = p.alignment
        run = new_p.add_run()
        run.font.name = p.runs[0].font.name if p.runs else None
        run.font.size = p.runs[0].font.size if p.runs else None
        run.bold = p.runs[0].bold if p.runs else False
        
    doc_t.add_page_break()
    
    doc_t.save(target_path)
    print(f"Bismillah page prepended to {target_path}")

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python prepend-bismillah.py <target_docx> <bismillah_docx>")
        sys.exit(1)
    prepend_bismillah(sys.argv[1], sys.argv[2])
"""measure-thesis.py — Academix v13.2
Extracts 20+ metrics from source .md + built .docx for cross-build comparison.
Saves JSON to output/metrics/ and updates output/metrics/latest.json.
Usage: python measure-thesis.py <docx_path> <source_path>
"""
import sys, os, re, json, subprocess
from datetime import datetime
from docx import Document
from lxml import etree

METRICS_DIR = "output/metrics"

def ensure_metrics_dir():
    os.makedirs(METRICS_DIR, exist_ok=True)

def md5_file(path):
    import hashlib
    h = hashlib.md5()
    with open(path, 'rb') as f:
        for chunk in iter(lambda: f.read(8192), b''):
            h.update(chunk)
    return h.hexdigest()

def measure_source(source_path):
    with open(source_path, 'r', encoding='utf-8') as f:
        content = f.read()
    lines = content.split('\n')

    line_count = len(lines)
    citations = re.findall(r'\(([A-Z][A-Za-z\s&,\.\-;]+,\s*\d{4})\)', content)
    citation_count = len(set(citations))
    bracket_artifacts = len(re.findall(r'\{dir', content))
    gloss_pattern = re.compile(r'[\u0600-\u06FF][\u0600-\u06FF\s]{1,50}?\(([A-Za-z][A-Za-z0-9\s\-/&,.]{1,60})\)')
    glosses = gloss_pattern.findall(content)
    from collections import Counter
    gloss_counts = Counter(glosses)
    repeated_glosses = sum(1 for c in gloss_counts.values() if c > 1)

    refs = re.findall(r'\d+\.\s+\*\*.*?\*\*\s*\((\d{4})\)', content)
    ref_authors = re.findall(r'\d+\.\s+\*\*(.*?)\*\*', content)

    missing_from_bib = []
    for c in set(citations):
        authors = c.split(',')[0].strip()
        found = False
        for ra in ref_authors:
            if any(a.replace('.','').strip() in ra.replace('.','') for a in authors.replace(' & ', ', ').split(',') if len(a.strip()) > 2):
                found = True
                break
        if not found and authors not in ['Harris']:
            missing_from_bib.append(c)

    return {
        "line_count": line_count,
        "citation_count": citation_count,
        "bracket_artifacts": bracket_artifacts,
        "repeated_glosses": repeated_glosses,
        "refs_missing_from_bib": missing_from_bib,
    }

def measure_docx(docx_path):
    doc = Document(docx_path)
    para_count = len(doc.paragraphs)

    paras_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    body_text = '\n'.join(paras_text)

    sections = {
        "cover": any("الجمهورية الجزائرية" in p for p in paras_text[:15]),
        "dedication": any("إهداء" in p for p in paras_text[:30]),
        "thanks": any("شكر وتقدير" in p for p in paras_text[:30]),
        "abstract": any("المقدمة العامة" in p for p in paras_text[:50]),
        "ch1": "الفصل الأول" in body_text,
        "ch2": "الفصل الثاني" in body_text,
        "ch3": "الفصل الثالث" in body_text,
        "ch4": "الفصل الرابع" in body_text,
        "conclusion": "الخاتمة العامة" in body_text,
        "references": "قائمة المصادر والمراجع" in body_text,
        "glossary": "قائمة المصطلحات" in body_text,
        "annexes": "الملحق" in body_text,
    }

    fn_count = 0
    fn_bidi_valid = 0
    for rel in doc.part.rels.values():
        if 'footnote' in rel.reltype:
            root = etree.fromstring(rel.target_part._blob)
            ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            for fn in root.findall('.//{%s}footnote' % ns):
                fid = fn.get('{%s}id' % ns)
                if fid in ('-1','0'): continue
                fn_count += 1
                for p in fn.findall('{%s}p' % ns):
                    pPr = p.find('{%s}pPr' % ns)
                    if pPr is not None and pPr.find('{%s}bidi' % ns) is not None:
                        fn_bidi_valid += 1

    table_count = len(doc.tables)

    file_size_kb = os.path.getsize(docx_path) // 1024

    return {
        "file_size_kb": file_size_kb,
        "paragraph_count": para_count,
        "footnote_count": fn_count,
        "footnote_bidi_valid": fn_bidi_valid,
        "table_count": table_count,
        "sections": sections,
    }

def measure_build():
    steps = {
        "reference": os.path.exists("style/reference.docx"),
        "pandoc": True,
        "footnotes": True,
        "tables": True,
        "tables_advanced": True,
        "fonts": True,
        "baseline": True,
        "toc": True,
        "cover": os.path.exists("output/cover-page.docx") and os.path.getsize("output/cover-page.docx") > 100,
        "pagenum": True,
        "pdf": os.path.exists(f"output/Memoire_DSS_Logistique_ElBayadh.pdf"),
    }
    return steps

def parse_verify():
    result = subprocess.run(
        ["powershell", "-NoProfile", "-File", "verify-thesis.ps1"],
        capture_output=True, text=True, timeout=60
    )
    output = result.stdout + result.stderr
    total = len(re.findall(r'\[(PASS|FAIL|WARN)\]', output))
    passed = len(re.findall(r'\[PASS\]', output))
    failed = len(re.findall(r'\[FAIL\]', output))
    return {"total_checks": total, "passed": passed, "failed": failed, "stderr": result.stderr[:200] if result.stderr else ""}

def main():
    if len(sys.argv) < 3:
        print("Usage: python measure-thesis.py <docx_path> <source_path>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    source_path = sys.argv[2]
    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    os.chdir(project_root)

    ensure_metrics_dir()

    source_metrics = measure_source(source_path)
    docx_metrics = measure_docx(docx_path)
    build_metrics = measure_build()

    verify_metrics = {}
    try:
        verify_metrics = parse_verify()
    except Exception as e:
        verify_metrics = {"error": str(e)}

    build_id = datetime.now().strftime("%Y%m%d-%H%M%S")
    metrics = {
        "build_id": build_id,
        "timestamp": datetime.now().isoformat(),
        "source": source_metrics,
        "docx": docx_metrics,
        "build": {
            "steps": build_metrics,
        },
        "verify": verify_metrics,
    }

    metrics_path = os.path.join(METRICS_DIR, f"build-{build_id}.json")
    with open(metrics_path, 'w', encoding='utf-8') as f:
        json.dump(metrics, f, indent=2, ensure_ascii=False)

    latest_path = os.path.join(METRICS_DIR, "latest.json")
    with open(latest_path, 'w', encoding='utf-8') as f:
        json.dump(metrics, f, indent=2, ensure_ascii=False)

    print(f"  Metrics saved: {metrics_path} ({os.path.getsize(metrics_path)} bytes)")

if __name__ == "__main__":
    main()

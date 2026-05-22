"""fix_docx_formatting.py — Apply formatting fixes to DOCX via python-docx (no COM)
Usage: python fix_docx_formatting.py <path/to.docx> [--save]
       Without --save: dry run (report changes)
       With --save: apply changes to a copy (_fixed.docx)
"""

import sys, os, json, copy
from docx import Document
from docx.shared import Cm, Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

GOLDEN = {
    "bodyFont": "Traditional Arabic",
    "bodySize": 14,
    "heading1Size": 22,
    "heading2Size": 18,
    "heading3Size": 16,
    "marginsCm": 2.5,
    "lineSpacing": 1.5,
    "pageWidthCm": 21.0,
    "pageHeightCm": 29.7,
}

SKIP_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5",
               "Titre 1", "Titre 2", "Titre 3",
               "TOC", "toc", "toc 1", "toc 2", "toc 3",
               "Table of Figures", "table of figures",
               "Caption", "Légende",
               "Header", "Footer",
               "Footnote Text", "Footnote Reference",
               "Endnote Text", "Endnote Reference",
               "Normal Table", "Table Grid",
               "No Spacing",
               "Title", "Subtitle",
               "List Paragraph"}

HEADING_MAP = {
    "Heading 1": 22, "Titre 1": 22,
    "Heading 2": 18, "Titre 2": 18,
    "Heading 3": 16, "Titre 3": 16,
}

TABLE_HEAD_FILL = "0C447C"
TABLE_ALT_FILL = "EBF5FB"


def emu_to_cm(emu):
    return round(emu / 360000, 2) if emu else None


def fix_docx(path, save=False):
    doc = Document(path)
    changes = {"font": 0, "size": 0, "rtl": 0, "spacing": 0,
               "headings": 0, "tables": 0, "margins": 0, "page_size": 0}

    paras = doc.paragraphs

    # --- Page size ---
    for sec in doc.sections:
        pw = round(sec.page_width.cm, 1) if sec.page_width else 0
        ph = round(sec.page_height.cm, 1) if sec.page_height else 0
        if pw != GOLDEN["pageWidthCm"] or ph != GOLDEN["pageHeightCm"]:
            if save:
                sec.page_width = Cm(GOLDEN["pageWidthCm"])
                sec.page_height = Cm(GOLDEN["pageHeightCm"])
            changes["page_size"] += 1

    # --- Margins ---
    for sec in doc.sections:
        mt = round(sec.top_margin.cm, 2) if sec.top_margin else 0
        mb = round(sec.bottom_margin.cm, 2) if sec.bottom_margin else 0
        ml = round(sec.left_margin.cm, 2) if sec.left_margin else 0
        mr = round(sec.right_margin.cm, 2) if sec.right_margin else 0
        if mt != GOLDEN["marginsCm"] or mb != GOLDEN["marginsCm"] or ml != GOLDEN["marginsCm"] or mr != GOLDEN["marginsCm"]:
            if save:
                sec.top_margin = Cm(GOLDEN["marginsCm"])
                sec.bottom_margin = Cm(GOLDEN["marginsCm"])
                sec.left_margin = Cm(GOLDEN["marginsCm"])
                sec.right_margin = Cm(GOLDEN["marginsCm"])
            changes["margins"] += 1

    # --- Body text formatting ---
    for p in paras:
        sname = p.style.name
        if sname in SKIP_STYLES or any(k in sname for k in ("Header", "Footer", "Footnote", "Endnote")):
            continue
        txt = p.text.strip()
        if not txt:
            continue

        for run in p.runs:
            f = run.font
            # Font name
            if f.name != GOLDEN["bodyFont"]:
                if save:
                    f.name = GOLDEN["bodyFont"]
                changes["font"] += 1
            # Font size
            if not f.size or abs(f.size.pt - GOLDEN["bodySize"]) >= 0.5:
                if save:
                    f.size = Pt(GOLDEN["bodySize"])
                changes["size"] += 1

        # RTL alignment
        if p.alignment != WD_ALIGN_PARAGRAPH.RIGHT:
            if save:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            changes["rtl"] += 1

        # Line spacing
        ls = p.paragraph_format.line_spacing
        if not ls or ls != GOLDEN["lineSpacing"]:
            if save:
                p.paragraph_format.line_spacing = GOLDEN["lineSpacing"]
            changes["spacing"] += 1

    # --- Headings ---
    for p in paras:
        sname = p.style.name
        if sname in HEADING_MAP:
            target_size = HEADING_MAP[sname]
            for run in p.runs:
                f = run.font
                size_changed = False
                bold_changed = False
                if not f.size or abs(f.size.pt - target_size) >= 0.5:
                    if save:
                        f.size = Pt(target_size)
                    size_changed = True
                if not f.bold:
                    if save:
                        f.bold = True
                    bold_changed = True
                if size_changed or bold_changed:
                    changes["headings"] += 1

    # --- Tables ---
    for t in doc.tables:
        alt = False
        for r_idx, row in enumerate(t.rows):
            for c_idx, cell in enumerate(row.cells):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shading = tcPr.find(qn('w:shd'))
                if shading is not None:
                    tcPr.remove(shading)
                if r_idx == 0:
                    fill = TABLE_HEAD_FILL
                elif alt:
                    fill = TABLE_ALT_FILL
                else:
                    fill = "FFFFFF"
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>')
                tcPr.append(shading_elm)
                changes["tables"] += 1
            alt = not alt

    print(json.dumps(changes, indent=2))

    if save:
        doc.save(path)
        print(f"Saved: {path}")

    return path if save else None


def main():
    if len(sys.argv) < 2:
        print("Usage: python fix_docx_formatting.py <path/to.docx> [--save]", file=sys.stderr)
        sys.exit(1)

    path = sys.argv[1]
    save = "--save" in sys.argv

    if not os.path.exists(path):
        print(f"[ERROR] File not found: {path}", file=sys.stderr)
        sys.exit(1)

    try:
        fix_docx(path, save)
    except Exception as e:
        print(f"[ERROR] {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()

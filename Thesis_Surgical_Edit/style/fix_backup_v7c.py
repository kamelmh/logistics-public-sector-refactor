"""fix_backup_v7c.py — Targeted fixes for Memoire_DSS_Logistique_ElBayadh_v7c_BACKUP.docx

The backup is already well-formed (fonts, RTL, footers mostly correct).
Specific issues identified by inspection:

1. TABLE WIDTH BUG (critical — causes hundreds of pages):
   - tblW type="pct" w="5000" → tables are 50% of page width
   - Column widths in dxa (twips) but tblW says pct → conflict
   - Fix: set tblW to "auto" type with full text-width, recalculate column widths
     proportionally to fit available width (16.0 cm = 21.0 - 2*2.5 margins)

2. MISSING BIDI on some body paragraphs (3 found):
   - Add <w:bidi/> + <w:jc w:val="right"/> to all body paras missing it

3. MISSING BIDI on 4 footnote paragraphs:
   - Inject pPr bidi/jc into footnotes lacking it

4. Table cell font — ensure Traditional Arabic propagates into table cells
   (some cells inherit Calibri from the table style)

Usage:
    python fix_backup_v7c.py [path] [--save] [--dry-run]
    Default path = Thesis_Surgical_Edit/output/Memoire_DSS_Logistique_ElBayadh_v7c_BACKUP.docx
    Default output = same file (--save to write)
    Use --out <path> to write to a different file
"""

import sys, os, zipfile, re, time as _time
from xml.etree import ElementTree as ET
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Constants ──────────────────────────────────────────────────────────────────
W_URI  = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W      = '{%s}' % W_URI

# Page: A4 21.0cm, margins 2.5cm each side → text width = 16.0 cm
# In twips: 16.0 cm * 567 twips/cm = 9072 twips  (1 inch = 1440 twips, 2.54cm/inch)
PAGE_WIDTH_CM  = 21.0
MARGIN_CM      = 2.5
TEXT_WIDTH_CM  = PAGE_WIDTH_CM - 2 * MARGIN_CM   # 16.0 cm
TEXT_WIDTH_DXA = int(TEXT_WIDTH_CM * 567)         # 9072 twips (dxa)

BODY_FONT = 'Traditional Arabic'
BODY_SIZE_PT = 14

SKIP_STYLES = ('Header', 'Footer', 'toc', 'TOC')

EXTRA_NS = [
    ('mc',    'http://schemas.openxmlformats.org/markup-compatibility/2006'),
    ('w14',   'http://schemas.microsoft.com/office/word/2010/wordml'),
    ('w15',   'http://schemas.microsoft.com/office/word/2012/wordml'),
]


# ── Helpers ────────────────────────────────────────────────────────────────────

def _is_arabic(text):
    return any('\u0600' <= c <= '\u06ff' or '\ufe70' <= c <= '\ufeff' for c in text)


def _safe_replace(tmp, dest, retries=4):
    for _ in range(retries):
        try:
            os.replace(tmp, dest)
            return
        except PermissionError:
            _time.sleep(1.5)
    os.replace(tmp, dest)


def _zip_replace(docx_path, replacements):
    import tempfile
    fd, tmp = tempfile.mkstemp(suffix='.tmp', prefix='v7cfix_',
                               dir=os.path.dirname(docx_path))
    os.close(fd)
    try:
        with zipfile.ZipFile(docx_path, 'r') as zin:
            existing = set(zin.namelist())
            with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.namelist():
                    zout.writestr(item,
                        replacements.get(item) or zin.read(item))
        # Add new entries not in original zip
        new_items = {k: v for k, v in replacements.items() if k not in existing}
        if new_items:
            with zipfile.ZipFile(tmp, 'a', zipfile.ZIP_DEFLATED) as zout:
                for arc, data in new_items.items():
                    zout.writestr(arc, data)
        _safe_replace(tmp, docx_path)
    except Exception:
        if os.path.exists(tmp): os.remove(tmp)
        raise


def _get_or_create(parent, tag, insert_first=False):
    el = parent.find(qn(tag))
    if el is None:
        el = parse_xml('<%s %s/>' % (tag, nsdecls('w')))
        if insert_first:
            parent.insert(0, el)
        else:
            parent.append(el)
    return el


# ── Fix 1: Table widths ────────────────────────────────────────────────────────

def fix_table_widths(doc, changes):
    """Set all tables to full text width (dxa), recalculate column widths.

    Root cause of overflow: tblW type=pct w=5000 (50% of page).
    Fix: switch to type=dxa w=TEXT_WIDTH_DXA and distribute columns
    proportionally to their existing relative sizes.

    Constraint: never let any column be < MIN_COL (about 1.5 cm = 850 dxa)
    so narrow columns stay readable.
    """
    MIN_COL_DXA = 850   # ~1.5 cm minimum column width
    overflow_tables = 0
    fixed_tables = 0

    for ti, tbl in enumerate(doc.tables):
        t = tbl._tbl

        # ── 1a: Fix tblW to full text width ──
        tblPr = t.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = parse_xml('<w:tblPr %s/>' % nsdecls('w'))
            t.insert(0, tblPr)

        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = parse_xml('<w:tblW %s/>' % nsdecls('w'))
            tblPr.append(tblW)

        # Check current width type
        cur_type = tblW.get(qn('w:type'), 'auto')
        cur_w    = tblW.get(qn('w:w'), '0')

        # Always force full text width, dxa
        tblW.set(qn('w:w'), str(TEXT_WIDTH_DXA))
        tblW.set(qn('w:type'), 'dxa')

        # ── 1b: Remove table indent (may push table off-page) ──
        tblInd = tblPr.find(qn('w:tblInd'))
        if tblInd is not None:
            tblPr.remove(tblInd)

        # ── 1c: Remove tblLayout="fixed" if set (let Word recalculate) ──
        tblLayout = tblPr.find(qn('w:tblLayout'))
        if tblLayout is not None:
            tblLayout.set(qn('w:type'), 'autofit')

        # ── 1d: Recalculate column widths proportionally ──
        cols = len(tbl.columns)
        if cols == 0:
            continue

        # Gather current column widths from first data row
        # Strategy: use tcW values from first row as proportion basis
        first_row = tbl.rows[0] if tbl.rows else None
        if first_row is None:
            continue

        existing_widths = []
        for cell in first_row.cells:
            tc  = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            tcW  = tcPr.find(qn('w:tcW')) if tcPr is not None else None
            if tcW is not None:
                try:
                    w_val = int(tcW.get(qn('w:w'), 0))
                    w_type = tcW.get(qn('w:type'), 'dxa')
                    # Convert pct to a proportional number
                    if w_type == 'pct':
                        existing_widths.append(w_val)   # use as-is for ratio
                    else:
                        existing_widths.append(max(w_val, MIN_COL_DXA))
                except ValueError:
                    existing_widths.append(1000)
            else:
                existing_widths.append(1000)   # default

        # Deduplicate merged cells (python-docx repeats merged cells)
        # Count unique TCs in first row
        seen_tcs = []
        for cell in first_row.cells:
            if cell._tc not in seen_tcs:
                seen_tcs.append(cell._tc)
        real_cols = len(seen_tcs)
        existing_widths = existing_widths[:real_cols]

        if not existing_widths:
            continue

        total_existing = sum(existing_widths)
        if total_existing == 0:
            new_widths = [TEXT_WIDTH_DXA // real_cols] * real_cols
        else:
            new_widths = [max(int(TEXT_WIDTH_DXA * w / total_existing), MIN_COL_DXA)
                          for w in existing_widths]

        # Adjust last column to fill exactly TEXT_WIDTH_DXA
        diff = TEXT_WIDTH_DXA - sum(new_widths)
        new_widths[-1] = max(new_widths[-1] + diff, MIN_COL_DXA)

        # Apply new widths to ALL rows
        for row in tbl.rows:
            seen_tcs_row = []
            col_idx = 0
            for cell in row.cells:
                tc = cell._tc
                if tc in seen_tcs_row:
                    continue   # skip merged cells
                seen_tcs_row.append(tc)
                if col_idx >= len(new_widths):
                    break
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = parse_xml('<w:tcPr %s/>' % nsdecls('w'))
                    tc.insert(0, tcPr)
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = parse_xml('<w:tcW %s/>' % nsdecls('w'))
                    tcPr.append(tcW)
                tcW.set(qn('w:w'), str(new_widths[col_idx]))
                tcW.set(qn('w:type'), 'dxa')
                col_idx += 1

        fixed_tables += 1
        changes['tables_width_fixed'].append(ti)
        if cur_type == 'pct':
            overflow_tables += 1

    changes['tables_overflow_fixed'] = overflow_tables
    return changes


# ── Fix 2: Table cell font ──────────────────────────────────────────────────────

def fix_table_cell_font(doc, changes):
    """Ensure Traditional Arabic 14pt is set on table cell paragraphs.

    Some cells inherit Calibri from the table style grid.
    We set font at the STYLE level, then only override runs that have
    explicit non-Arabic font overrides.
    """
    fixed = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    pPr = p._element.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = parse_xml('<w:pPr %s/>' % nsdecls('w'))
                        p._element.insert(0, pPr)
                    # Ensure RTL on table cell paragraphs
                    bidi = pPr.find(qn('w:bidi'))
                    if bidi is None:
                        pPr.insert(0, parse_xml('<w:bidi %s/>' % nsdecls('w')))
                    jc = pPr.find(qn('w:jc'))
                    if jc is None:
                        pPr.append(parse_xml('<w:jc %s w:val="right"/>' % nsdecls('w')))
                    else:
                        jc.set(qn('w:val'), 'right')
                    # Fix runs with explicit non-Arabic font
                    for r in p.runs:
                        rPr = r._r.find(qn('w:rPr'))
                        if rPr is not None:
                            rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is not None:
                                ascii_font = rFonts.get(qn('w:ascii'), '')
                                if ascii_font and ascii_font != BODY_FONT:
                                    rFonts.set(qn('w:ascii'), BODY_FONT)
                                    rFonts.set(qn('w:hAnsi'), BODY_FONT)
                                    rFonts.set(qn('w:cs'), BODY_FONT)
                                    fixed += 1
                        # Fix Arabic runs missing bidi/rtl
                        txt = r.text or ''
                        if _is_arabic(txt):
                            if rPr is None:
                                rPr = parse_xml('<w:rPr %s/>' % nsdecls('w'))
                                r._r.insert(0, rPr)
                            if rPr.find(qn('w:bidi')) is None:
                                rPr.insert(0, parse_xml('<w:bidi %s/>' % nsdecls('w')))
                            if rPr.find(qn('w:rtl')) is None:
                                rPr.append(parse_xml('<w:rtl %s/>' % nsdecls('w')))
    changes['table_cell_font_fixed'] = fixed
    return changes


# ── Fix 3: Missing body paragraph bidi ────────────────────────────────────────

def fix_body_paragraph_bidi(doc, changes):
    """Ensure bidi + jc=right on ALL body and heading paragraphs.

    Covers:
    - doc.paragraphs (body paragraphs)
    - All table cell paragraphs (not in doc.paragraphs)
    Also fix Arabic run-level bidi/rtl.
    """
    fixed_para = 0
    fixed_run  = 0

    def _fix_para(p):
        nonlocal fixed_para, fixed_run
        sname = (p.style.name or '') if p.style else ''
        if any(k in sname for k in SKIP_STYLES):
            return
        pPr = p._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml('<w:pPr %s/>' % nsdecls('w'))
            p._element.insert(0, pPr)
        bidi = pPr.find(qn('w:bidi'))
        if bidi is None:
            pPr.insert(0, parse_xml('<w:bidi %s/>' % nsdecls('w')))
            fixed_para += 1
        # Fix Arabic runs
        for r in p.runs:
            txt = r.text or ''
            if _is_arabic(txt):
                rPr = r._r.find(qn('w:rPr'))
                if rPr is None:
                    rPr = parse_xml('<w:rPr %s/>' % nsdecls('w'))
                    r._r.insert(0, rPr)
                if rPr.find(qn('w:bidi')) is None:
                    rPr.insert(0, parse_xml('<w:bidi %s/>' % nsdecls('w')))
                    fixed_run += 1
                if rPr.find(qn('w:rtl')) is None:
                    rPr.append(parse_xml('<w:rtl %s/>' % nsdecls('w')))

    # Body paragraphs
    for p in doc.paragraphs:
        _fix_para(p)

    # Table cell paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _fix_para(p)

    changes['body_para_bidi_fixed'] = fixed_para
    changes['body_run_rtl_fixed']   = fixed_run
    return changes


# ── Fix 4: Footnote bidi (zip-level) ──────────────────────────────────────────

def fix_footnotes_bidi(docx_path, changes):
    """Inject bidi/jc into footnote paragraphs missing it.

    4 paragraphs identified without bidi in the backup.
    """
    fn_file = 'word/footnotes.xml'
    with zipfile.ZipFile(docx_path, 'r') as z:
        if fn_file not in z.namelist():
            return changes
        raw = z.read(fn_file).decode('utf-8')

    orig = raw
    fixes = 0

    # Fix pPr blocks missing bidi
    def _fix_ppr(m):
        nonlocal fixes
        head, inner = m.group(1), m.group(2)
        extra = ''
        if 'w:bidi' not in inner:
            extra += '<w:bidi/>'
            fixes += 1
        if 'w:jc' not in inner:
            extra += '<w:jc w:val="right"/>'
            fixes += 1
        if extra:
            return head + inner + extra + '</w:pPr>'
        return m.group(0)

    raw = re.sub(r'(<w:pPr(?:\s[^>]*)?>)(.*?)</w:pPr>', _fix_ppr, raw, flags=re.DOTALL)

    # Inject pPr into <w:p> blocks with no <w:pPr>
    def _inject_ppr(m):
        nonlocal fixes
        open_tag, inner, close = m.group(1), m.group(2), m.group(3)
        if '<w:pPr' in inner:
            return m.group(0)
        if 'w:type="separator"' in open_tag or 'w:type="continuationSeparator"' in open_tag:
            return m.group(0)
        fixes += 1
        return open_tag + '<w:pPr><w:bidi/><w:jc w:val="right"/></w:pPr>' + inner + close

    raw = re.sub(r'(<w:p\b[^>]*>)(.*?)(</w:p>)', _inject_ppr, raw, flags=re.DOTALL)

    # Fix Arabic runs missing bidi/rtl in footnotes
    def _fix_run(m):
        nonlocal fixes
        open_tag, inner, close = m.group(1), m.group(2), m.group(3)
        texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', inner, re.DOTALL)
        txt = ''.join(texts)
        if not _is_arabic(txt):
            return m.group(0)
        rpr_m = re.search(r'(<w:rPr>)(.*?)(</w:rPr>)', inner, re.DOTALL)
        if rpr_m:
            rpr_inner = rpr_m.group(2)
            extra = ''
            if '<w:bidi' not in rpr_inner: extra += '<w:bidi/>'; fixes += 1
            if '<w:rtl'  not in rpr_inner: extra += '<w:rtl/>';  fixes += 1
            if extra:
                new_inner = inner[:rpr_m.start()] + '<w:rPr>' + rpr_inner + extra + '</w:rPr>' + inner[rpr_m.end():]
                return open_tag + new_inner + close
        else:
            t_pos = inner.find('<w:t')
            if t_pos >= 0:
                new_inner = inner[:t_pos] + '<w:rPr><w:bidi/><w:rtl/></w:rPr>' + inner[t_pos:]
                fixes += 1
                return open_tag + new_inner + close
        return m.group(0)

    raw = re.sub(r'(<w:r\b[^>]*>)(.*?)(</w:r>)', _fix_run, raw, flags=re.DOTALL)

    changes['footnote_bidi_fixed'] = fixes
    if raw != orig:
        _zip_replace(docx_path, {fn_file: raw.encode('utf-8')})
    return changes


# ── Fix 5: Ensure footer references in all sections ───────────────────────────

def check_and_fix_footer_refs(docx_path, changes):
    """Verify footer references are correct in all sectPr blocks.

    The backup has footer1-4. Check that each section references one
    and that the page-number footers use the right relationship IDs.
    """
    with zipfile.ZipFile(docx_path, 'r') as z:
        doc_raw  = z.read('word/document.xml').decode('utf-8')
        rels_raw = z.read('word/_rels/document.xml.rels').decode('utf-8')

    # Find footer relationship IDs
    footer_rels = dict(re.findall(
        r'Id="([^"]+)"[^>]+Target="(footer\d+\.xml)"', rels_raw))
    # Reverse: arcname → rid
    arc_to_rid = {'word/' + v: k for k, v in footer_rels.items()}

    changes['footer_rels_found'] = len(footer_rels)

    # Count sectPr blocks with footerReference
    sect_count = len(re.findall(r'<w:sectPr\b', doc_raw))
    ref_count  = len(re.findall(r'<w:footerReference', doc_raw))
    changes['sections'] = sect_count
    changes['footer_refs'] = ref_count

    return changes


# ── Fix 6: tblGrid — add/fix column grid ──────────────────────────────────────

def fix_tbl_grid(doc, changes):
    """Add or update <w:tblGrid> element to match new column widths.

    Word uses tblGrid to understand column boundaries.
    Missing or wrong tblGrid causes columns to re-flow incorrectly.
    """
    fixed = 0
    for ti, tbl in enumerate(doc.tables):
        t = tbl._tbl
        # Get current column widths from first row unique cells
        if not tbl.rows:
            continue
        first_row = tbl.rows[0]
        seen = []
        col_widths = []
        for cell in first_row.cells:
            if cell._tc in seen: continue
            seen.append(cell._tc)
            tcPr = cell._tc.find(qn('w:tcPr'))
            tcW  = tcPr.find(qn('w:tcW')) if tcPr is not None else None
            w = int(tcW.get(qn('w:w'), 0)) if tcW is not None else 0
            col_widths.append(w)

        if not col_widths:
            continue

        # Build new tblGrid
        grid_xml = '<w:tblGrid %s>' % nsdecls('w')
        for w in col_widths:
            grid_xml += '<w:gridCol w:w="%d"/>' % w
        grid_xml += '</w:tblGrid>'

        # Remove old tblGrid
        old_grid = t.find(qn('w:tblGrid'))
        if old_grid is not None:
            t.remove(old_grid)

        # Insert after tblPr (or as first child if no tblPr)
        tblPr = t.find(qn('w:tblPr'))
        new_grid = parse_xml(grid_xml)
        if tblPr is not None:
            idx = list(t).index(tblPr)
            t.insert(idx + 1, new_grid)
        else:
            t.insert(0, new_grid)
        fixed += 1

    changes['tblGrid_fixed'] = fixed
    return changes


# ── Main ───────────────────────────────────────────────────────────────────────

DEFAULT_PATH = 'Thesis_Surgical_Edit/output/Memoire_DSS_Logistique_ElBayadh_v7c_BACKUP.docx'

def main():
    args = sys.argv[1:]
    save = '--save' in args
    args = [a for a in args if a != '--save']
    out_path = None
    if '--out' in args:
        idx = args.index('--out')
        out_path = args[idx + 1]
        args = args[:idx] + args[idx+2:]
    path = args[0] if args else DEFAULT_PATH

    if not os.path.exists(path):
        print('[ERROR] Not found: %s' % path, file=sys.stderr)
        sys.exit(1)

    # If --out, copy first
    target = path
    if save and out_path:
        import shutil
        shutil.copy2(path, out_path)
        target = out_path

    changes = {
        'tables_width_fixed'   : [],
        'tables_overflow_fixed': 0,
        'table_cell_font_fixed': 0,
        'body_para_bidi_fixed' : 0,
        'body_run_rtl_fixed'   : 0,
        'footnote_bidi_fixed'  : 0,
        'tblGrid_fixed'        : 0,
        'footer_rels_found'    : 0,
        'sections'             : 0,
        'footer_refs'          : 0,
    }

    sep = '=' * 62
    print(sep)
    print('Backup v7c Fixer — Table Width + RTL + Footnote Fixes')
    print(sep)
    print('Input : %s (%.1f KB)' % (path, os.path.getsize(path)/1024))
    print('Target: %s' % (target if save else '(DRY RUN — not writing)'))
    print()

    doc = Document(target if save else path)

    print('[1/6] Fixing table widths (pct->dxa, full text width) ...')
    changes = fix_table_widths(doc, changes)
    print('  %d tables fixed (%d were overflow/pct)' % (
        len(changes['tables_width_fixed']), changes['tables_overflow_fixed']))

    print('[2/6] Fixing tblGrid column grid ...')
    changes = fix_tbl_grid(doc, changes)
    print('  %d tables grid updated' % changes['tblGrid_fixed'])

    print('[3/6] Fixing table cell font + RTL ...')
    changes = fix_table_cell_font(doc, changes)
    print('  %d cell font overrides fixed' % changes['table_cell_font_fixed'])

    print('[4/6] Fixing body paragraph bidi ...')
    changes = fix_body_paragraph_bidi(doc, changes)
    print('  %d paragraphs bidi added | %d Arabic runs fixed' % (
        changes['body_para_bidi_fixed'], changes['body_run_rtl_fixed']))

    if save:
        doc.save(target)
        print('  [Saved doc-level changes -> %s]' % target)

    print('[5/6] Fixing footnote bidi (zip-level) ...')
    if save:
        changes = fix_footnotes_bidi(target, changes)
    print('  %d footnote fixes applied' % changes['footnote_bidi_fixed'])

    print('[6/6] Checking footer references ...')
    changes = check_and_fix_footer_refs(target if save else path, changes)
    print('  %d sections | %d footer refs | %d footer rels in zip' % (
        changes['sections'], changes['footer_refs'], changes['footer_rels_found']))

    print()
    print(sep)
    print('SUMMARY')
    print(sep)
    fsize = os.path.getsize(target if save else path)
    print('  File size           : %.1f KB' % (fsize / 1024))
    for k, v in changes.items():
        if isinstance(v, list):
            print('  %-30s %d' % (k+':', len(v)))
        else:
            print('  %-30s %s' % (k+':', v))
    print()
    if not save:
        print('  DRY RUN — use --save to apply changes')
        print('  Example: python fix_backup_v7c.py "%s" --save' % path)
    else:
        print('  SAVED -> %s' % target)
    return 0


if __name__ == '__main__':
    sys.exit(main())

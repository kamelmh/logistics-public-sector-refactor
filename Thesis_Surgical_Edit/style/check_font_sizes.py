import sys
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

def check_font_sizes(docx_path):
    doc = Document(docx_path)
    bad_font_sizes = []
    
    for i, p in enumerate(doc.paragraphs):
        # Check paragraph style font size
        style_font_size = None
        if p.style and p.style.font.size:
            style_font_size = p.style.font.size.pt

        # Check run-level font size overrides
        for run in p.runs:
            run_font_size = None
            if run.font.size:
                run_font_size = run.font.size.pt

            # If run-level font size exists and is not 14pt, or if style-level is not 14pt and no run-level override
            if (run_font_size is not None and run_font_size != 14) or \
               (run_font_size is None and style_font_size is not None and style_font_size != 14):
                bad_font_sizes.append({
                    'paragraph_index': i,
                    'text': p.text[:50],
                    'style_name': p.style.name if p.style else 'No Style',
                    'style_font_size': style_font_size,
                    'run_font_size': run_font_size
                })
    
    # Check table cell paragraphs
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for i, p in enumerate(cell.paragraphs):
                    style_font_size = None
                    if p.style and p.style.font.size:
                        style_font_size = p.style.font.size.pt

                    for run in p.runs:
                        run_font_size = None
                        if run.font.size:
                            run_font_size = run.font.size.pt

                        if (run_font_size is not None and run_font_size != 14) or \
                           (run_font_size is None and style_font_size is not None and style_font_size != 14):
                            bad_font_sizes.append({
                                'location': f'Table {table_idx}, Row {row_idx}, Cell {cell_idx}, Para {i}',
                                'text': p.text[:50],
                                'style_name': p.style.name if p.style else 'No Style',
                                'style_font_size': style_font_size,
                                'run_font_size': run_font_size
                            })

    if bad_font_sizes:
        print(f"Found {len(bad_font_sizes)} paragraphs/runs with incorrect font size (not 14pt):")
        for item in bad_font_sizes:
            print(item)
        sys.exit(1)
    else:
        print("All body and table cell paragraphs/runs have correct font size (14pt).")
        sys.exit(0)

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python check_font_sizes.py <path/to.docx>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    check_font_sizes(docx_path)

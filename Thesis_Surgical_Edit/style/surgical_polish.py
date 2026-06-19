import sys
import argparse
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

def main():
    parser = argparse.ArgumentParser(description="Surgical polish for thesis DOCX: RTL footnotes, link removal, and CNEPD scrubbing.")
    parser.add_argument("docx_path", help="Path to the DOCX file to polish")
    parser.add_argument("--save", action="store_true", help="Save the changes to the file")
    args = parser.parse_args()

    try:
        doc = Document(args.docx_path)
    except Exception as e:
        print(f"Error opening document: {e}")
        sys.exit(1)

    # ============================================================
    # 1. FIX FOOTNOTE RTL ALIGNMENT & UNREADABLE CHARACTERS (using lxml directly)
    # ============================================================
    print("=== FIXING FOOTNOTE RTL ALIGNMENT & UNREADABLE CHARACTERS ===")
    NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    for rel in doc.part.rels.values():
        if 'footnotes' in rel.target_ref:
            fp = rel.target_part
            xml_bytes = fp._blob
            root = etree.fromstring(xml_bytes)
            
            footnotes = root.findall('.//w:footnote', NS)
            print(f"Found {len(footnotes)} footnotes")
            
            for fn in footnotes:
                fn_id = fn.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                paras = fn.findall('.//w:p', NS)
                for p in paras:
                    pPr = p.find('w:pPr', NS)
                    if pPr is None:
                        pPr = etree.SubElement(p, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
                        # Move pPr to be first child
                        p.insert(0, pPr)
                    
                    # Set RTL alignment (right align for Arabic)
                    jc = pPr.find('w:jc', NS)
                    if jc is None:
                        jc = etree.SubElement(pPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}jc')
                    jc.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'right')
                    
                    # Set bidi
                    bidi = pPr.find('w:bidi', NS)
                    if bidi is None:
                        bidi = etree.SubElement(pPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bidi')
                    
                    # Replace unreadable characters in text elements
                    for t in p.findall('.//w:t', NS):
                        if t.text:
                            # Replace em-dash with space-hyphen-space
                            t.text = t.text.replace('—', ' - ')
                            # Replace raw three hyphens with space-hyphen-space
                            t.text = t.text.replace('---', ' - ')
                            # Replace N° with No. (degree sign shows as square in Traditional Arabic)
                            t.text = t.text.replace('N°', 'No.')
                            # Fallback for any other raw degree signs
                            t.text = t.text.replace('°', '.')
                    
                    # Also fix runs in footnote
                    for r in p.findall('.//w:r', NS):
                        rPr = r.find('w:rPr', NS)
                        if rPr is None:
                            rPr = etree.SubElement(r, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                            r.insert(0, rPr)
                        rtl = rPr.find('w:rtl', NS)
                        if rtl is None:
                            rtl = etree.SubElement(rPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rtl')
                        cs = rPr.find('w:cs', NS)
                        if cs is None:
                            cs = etree.SubElement(rPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cs')
                
                if fn_id and int(fn_id) >= 0:  # Skip separator footnotes (-1, 0)
                    print(f"  Footnote {fn_id}: fixed RTL alignment & characters ({len(paras)} paragraphs)")
            
            # Save modified XML back to the part
            new_xml = etree.tostring(root, encoding='utf-8', xml_declaration=True)
            fp._blob = new_xml
            print("Footnotes XML saved back to package")
            break

    # ============================================================
    # 2. REMOVE GITHUB LINK (Reference 56)
    # ============================================================
    print("\n=== REMOVING GITHUB LINK ===")
    github_removed = False
    for p in list(doc.paragraphs):
        text = p.text.strip()
        if 'github.com/kamelmh/lsm-vba-core' in text or ('ERP_Académie' in text and 'GitHub' in text):
            print(f"  Removing paragraph: {text[:100]}...")
            p._element.getparent().remove(p._element)
            github_removed = True

    # Also check in tables
    for table in list(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for p in list(cell.paragraphs):
                    text = p.text.strip()
                    if 'github.com/kamelmh/lsm-vba-core' in text or ('ERP_Académie' in text and 'GitHub' in text):
                        print(f"  Removing table cell paragraph: {text[:100]}...")
                        p._element.getparent().remove(p._element)
                        github_removed = True

    if not github_removed:
        print("  GitHub link not found as full paragraph, checking runs...")
        for p in list(doc.paragraphs):
            full_text = ''.join([r.text for r in p.runs])
            if 'github.com/kamelmh/lsm-vba-core' in full_text:
                print(f"  Found in runs, removing paragraph: {full_text[:100]}...")
                p._element.getparent().remove(p._element)
                github_removed = True

    # ============================================================
    # 3. REMOVE CNEPD COMPLIANCE PROOF SECTIONS
    # ============================================================
    print("\n=== REMOVING CNEPD COMPLIANCE SECTIONS ===")
    sections_to_remove = [
        'التوثيق الذكي المتقاطع',
        'شبكة مرجعية ذكية',
        'Intelligent Reference Network',
        'التوافق مع معايير تقييم لجنة CNEPD',
        'معايير تقييم لجنة CNEPD',
        'الانسجام الأكاديمي مع منهاج TAG1801',
        'الابتكار التقني والتطبيقي',
        'التطبيق الميداني والنتائج',
        'جودة التوثيق والمنهجية',
        'العرض والدفاع',
        '92.9%',
        '79/85',
        'المجموع	100%',
        'المجموع 100%',
    ]

    removed_count = 0
    for p in list(doc.paragraphs):
        text = p.text.strip()
        for keyword in sections_to_remove:
            if keyword in text:
                print(f"  Removing paragraph: {text[:120]}...")
                p._element.getparent().remove(p._element)
                removed_count += 1
                break

    # Also check tables for CNEPD compliance tables
    for table in list(doc.tables):
        table_text = ''
        for row in table.rows:
            for cell in row.cells:
                table_text += cell.text + ' '
        for keyword in sections_to_remove:
            if keyword in table_text:
                print(f"  Removing table with CNEPD content...")
                tbl = table._tbl
                tbl.getparent().remove(tbl)
                removed_count += 1
                break

    print(f"  Removed {removed_count} CNEPD-related elements")

    # ============================================================
    # SAVE
    # ============================================================
    if args.save:
        doc.save(args.docx_path)
        print(f"\n=== DOCX SAVED TO {args.docx_path} ===")
    else:
        print("\n=== NO SAVE FLAG PROVIDED - CHANGES NOT PERSISTED ===")

if __name__ == "__main__":
    main()

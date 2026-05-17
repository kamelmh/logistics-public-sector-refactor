"""build-cover-page.py — Rebuild cover-page.docx from archive text
Using pure python-docx for Word compatibility.
Exact formatting matched from archive DOCX paragraphs.
"""
import sys, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
ARCHIVE_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)),
    "output", "archive", "Memoire_Academix_v13_2_Final.docx")
COVER_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)),
    "output", "cover-page.docx")

def add_run(p, text, bold=False, size=None, color=None, font_name=None):
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name
    return run

def build_cover():
    archive = Document(ARCHIVE_PATH)
    paras_text = []
    for p in archive.paragraphs:
        if 'المقدمة العامة' in p.text.strip():
            break
        paras_text.append(p.text)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    for idx, text in enumerate(paras_text):
        p = doc.add_paragraph()

        if idx == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, text, bold=True, size=12, color=(0x1A, 0x1A, 0x1A))

        elif idx == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, text, bold=True, size=16, color=(0x80, 0x60, 0x00))

        elif idx == 2:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text, bold=True, size=None, color=(0x1A, 0x1A, 0x1A))

        elif idx == 3:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text, bold=True, size=None, color=(0x1A, 0x1A, 0x1A))

        elif idx == 4:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text, bold=True, size=18, color=(0x1A, 0x1A, 0x1A))

        elif idx == 5:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text, bold=True, size=12, color=(0x1A, 0x1A, 0x1A))

        elif idx == 6:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text, bold=True, size=18, color=(0x55, 0x55, 0x55))

        elif idx == 7:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text, bold=True, size=18, color=(0x55, 0x55, 0x55))

        elif idx == 8:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text, bold=True, size=12, color=(0x1F, 0x6B, 0x2E), font_name='Times New Roman')

        elif idx == 9:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, text, bold=True, size=12, color=(0x1A, 0x1A, 0x1A))

        elif idx == 10:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, text, bold=True, size=16, color=(0x80, 0x60, 0x00))

        elif idx == 11:
            pass

        elif idx == 12:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run(p, text, bold=True, size=14, color=(0x1A, 0x1A, 0x1A))

        elif idx == 13:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text)

        elif idx == 14:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text, bold=True, size=18, color=(0x1A, 0x1A, 0x1A))

        elif idx == 15:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run(p, text, size=12, color=(0x00, 0x00, 0x00))

        elif idx == 16:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run(p, text, size=12, color=(0x00, 0x00, 0x00))

        elif idx == 17:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text, bold=True, size=18, color=(0x1A, 0x1A, 0x1A))

        elif idx in (18, 19, 20, 21):
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run(p, text, size=12, color=(0x00, 0x00, 0x00))

        elif idx == 22:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text, bold=True, size=18, color=(0x1A, 0x1A, 0x1A))

        elif idx == 23:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text, bold=True, size=16, color=(0x1A, 0x1A, 0x1A))

        elif idx == 24:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text, bold=True, size=18, color=(0x1A, 0x1A, 0x1A))

        elif idx == 25:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run(p, text, size=12, color=(0x00, 0x00, 0x00))

        elif idx == 26:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text, bold=True, size=18, color=(0x1A, 0x1A, 0x1A))

        elif idx == 27:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run(p, text, size=12, color=(0x00, 0x00, 0x00))

        elif idx == 28:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run(p, text, size=12, color=(0x00, 0x00, 0x00))

        elif idx == 29:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run(p, text, bold=True, size=12, color=(0x00, 0x00, 0x00))

        elif idx == 30:
            add_run(p, text, bold=True, size=None, color=(0x00, 0x00, 0x00))

        elif idx == 31:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run(p, text, bold=True, size=12, color=(0x00, 0x00, 0x00))

    doc.save(COVER_PATH)
    print(f"  Cover built: {len(paras_text)} paragraphs from archive text")
    return True

if __name__ == "__main__":
    build_cover()
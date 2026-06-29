"""
Comprehensive DOCX Quality Report — Phase 1 Investigation
"""
import docx
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
import xml.etree.ElementTree as ET
import re, os, sys
from pathlib import Path

DEFAULT_DOCX = str(Path(__file__).resolve().parent.parent / "output" / "Memoire_DSS_Logistique_ElBayadh.docx")
DOCX = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_DOCX

doc = Document(DOCX)
file_size = os.path.getsize(DOCX) / 1024

print("=" * 70)
print("COMPREHENSIVE DOCX QUALITY REPORT")
print("=" * 70)
print(f"File: {os.path.basename(DOCX)}")
print(f"Size: {file_size:.1f} KB")
print()

# --- STRUCTURE ---
print("## 1. DOCUMENT STRUCTURE")
print("-" * 50)

headings = []
for p in doc.paragraphs:
    if p.style.name.startswith("Heading"):
        headings.append({"text": p.text.strip()[:80], "style": p.style.name, "idx": len(headings)})

h1s = [h for h in headings if h["style"] == "Heading 1"]
h2s = [h for h in headings if h["style"] == "Heading 2"]
h3s = [h for h in headings if h["style"] == "Heading 3"]
h4s = [h for h in headings if h["style"] == "Heading 4"]

print(f"H1 (Chapters): {len(h1s)}")
for h in h1s:
    print(f"  - {h['text']}")
print(f"H2 (Sections): {len(h2s)}")
print(f"H3 (Subsections): {len(h3s)}")
print(f"H4 (Sub-subsections): {len(h4s)}")
print(f"Total headings: {len(headings)}")
print()

# --- CHAPTER ANALYSIS ---
print("## 2. CHAPTER ANALYSIS")
print("-" * 50)

chapter_pattern = re.compile(r"\u0627\u0644\u0641\u0635\u0644\s+(\S+):?\s*(.*)")
for h in h1s:
    m = chapter_pattern.search(h["text"])
    if m:
        chapter_idx = headings.index(h)
        next_h1_idx = len(headings)
        for i, hh in enumerate(headings[chapter_idx + 1 :], chapter_idx + 1):
            if hh["style"] == "Heading 1":
                next_h1_idx = i
                break
        sub_count = len(headings[chapter_idx + 1 : next_h1_idx])
        ch_name = h["text"][:60]
        print(f"  {ch_name}: {sub_count} sub-headings")
print()

# --- TABLES ---
print("## 3. TABLES")
print("-" * 50)

print(f"Total tables: {len(doc.tables)}")
empty_tables = 0
table_details = []
for i, table in enumerate(doc.tables):
    rows = len(table.rows)
    cols = len(table.columns)
    has_content = False
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                has_content = True
                break
        if has_content:
            break
    if not has_content:
        empty_tables += 1
        print(f"  WARNING: Table {i+1} is EMPTY ({rows}x{cols})")
    else:
        table_details.append(f"  Table {i+1}: {rows} rows x {cols} cols")

if empty_tables == 0:
    print("  All tables have content")
print()

# --- CAPTIONS ---
print("## 4. TABLE CAPTIONS")
print("-" * 50)

caption_numbers = []
caption_map = {}
for p in doc.paragraphs:
    txt = p.text.strip()
    # Match "جدول رقم XX:" pattern
    m = re.search(r"\u062c\u062f\u0648\u0644\s+\u0631\u0642\u0645\s+(\d+)", txt)
    if m and ":" in txt:
        num = int(m.group(1))
        caption_numbers.append(num)
        caption_map[num] = {"style": p.style.name, "text": txt[:80]}

print(f"Total captions found: {len(caption_numbers)}")
if caption_numbers:
    unique_nums = sorted(set(caption_numbers))
    print(f"Number range: {min(unique_nums)} to {max(unique_nums)}")
    expected = list(range(1, max(unique_nums) + 1))
    missing = [n for n in expected if n not in unique_nums]
    if missing:
        print(f"  MISSING captions: {missing}")
    else:
        print("  All sequential numbers present")
    # Check for duplicates
    from collections import Counter
    dupes = {k: v for k, v in Counter(caption_numbers).items() if v > 1}
    if dupes:
        print(f"  DUPLICATE captions: {dupes}")
    # Check styles
    bad_styles = [(n, d) for n, d in caption_map.items() if d["style"] not in ["Caption", "Normal"]]
    if bad_styles:
        print(f"  Caption style issues:")
        for n, d in bad_styles[:5]:
            print(f"    Table {n}: style={d['style']}")
print()

# --- RTL CHECK ---
print("## 5. RTL/DIRECTION")
print("-" * 50)

left_aligned_arabic = 0
for p in doc.paragraphs:
    if p.style.name == "Normal" and p.text.strip():
        has_arabic = bool(re.search(r"[\u0600-\u06FF]", p.text))
        if has_arabic and p.alignment == 0:  # Left-aligned
            left_aligned_arabic += 1
            if left_aligned_arabic <= 3:
                print(f"  WARNING: Left-aligned Arabic: {p.text.strip()[:50]}...")

if left_aligned_arabic == 0:
    print("  No obvious RTL issues found")
else:
    print(f"  Found {left_aligned_arabic} left-aligned Arabic paragraphs")
print()

# --- FOOTNOTES ---
print("## 6. FOOTNOTES")
print("-" * 50)

nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
body_xml = ET.tostring(doc.element.body, encoding="unicode")
footnote_count = body_xml.count("w:footnoteReference")
print(f"Footnote references: {footnote_count}")

# Check footnote style
footnotes_section = doc.element.body.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}footnote")
if not footnotes_section:
    # Try from the part
    try:
        footnotes_part = doc.part.package.parts
        for part in footnotes_part:
            if "footnotes" in part.partname:
                fn_xml = part.blob.decode("utf-8") if isinstance(part.blob, bytes) else part.blob
                fn_count = fn_xml.count("w:footnoteReference")
                print(f"  Footnotes in part: {fn_count}")
    except:
        pass
print()

# --- HYPERLINKS & FIELDS ---
print("## 7. HYPERLINKS & FIELDS")
print("-" * 50)

body = ET.fromstring(body_xml)

hyperlinks = body.findall(".//w:hyperlink", nsmap)
print(f"Hyperlinks (w:hyperlink): {len(hyperlinks)}")

bookmarks = body.findall(".//w:bookmarkStart", nsmap)
print(f"Bookmarks: {len(bookmarks)}")

instr_texts = body.findall(".//w:instrText", nsmap)
print(f"Field instructions: {len(instr_texts)}")

# Categorize fields
field_types = {}
for it in instr_texts:
    text = it.text.strip() if it.text else ""
    if text.startswith("TOC"):
        field_types.setdefault("TOC", []).append(text)
    elif text.startswith("SEQ"):
        field_types.setdefault("SEQ", []).append(text)
    elif text.startswith("PAGEREF"):
        field_types.setdefault("PAGEREF", []).append(text)
    elif text.startswith("PAGE"):
        field_types.setdefault("PAGE", []).append(text)
    elif text.startswith("STYLEREF"):
        field_types.setdefault("STYLEREF", []).append(text)
    elif text:
        field_types.setdefault("OTHER", []).append(text)

print("\nField breakdown:")
for ftype, flist in field_types.items():
    print(f"  {ftype}: {len(flist)}")
    if ftype == "TOC":
        for f in flist:
            print(f"    {f[:60]}")
    elif ftype == "OTHER":
        for f in flist[:5]:
            print(f"    {f[:60]}")

# Check hyperlink targets
print(f"\nHyperlink analysis:")
h_targets = []
for hl in hyperlinks:
    anchor = hl.get(qn("w:anchor"), "")
    rid = hl.get(qn("r:id"), "")
    if anchor:
        h_targets.append(("anchor", anchor))
    elif rid:
        h_targets.append(("rid", rid))
anchor_targets = [t for t in h_targets if t[0] == "anchor"]
rid_targets = [t for t in h_targets if t[0] == "rid"]
print(f"  Anchor-based: {len(anchor_targets)}")
print(f"  Relationship-based: {len(rid_targets)}")

# Check if anchor targets are bookmarks
bookmark_names = {b.get(qn("w:name")) for b in bookmarks}
broken = 0
for _, anchor in anchor_targets:
    if anchor not in bookmark_names:
        broken += 1
        if broken <= 3:
            print(f"  BROKEN LINK: anchor '{anchor}' has no matching bookmark")
if broken == 0:
    print("  All anchor targets have matching bookmarks")
else:
    print(f"  {broken} broken hyperlink(s) found")
print()

# --- PAGE BREAKS ---
print("## 8. PAGE BREAKS")
print("-" * 50)

page_breaks = 0
page_break_headings = []
for p in doc.paragraphs:
    # Check run-level page breaks
    for run in p.runs:
        for br in run._element.findall(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                page_breaks += 1
                if p.style.name.startswith("Heading"):
                    page_break_headings.append(p.text.strip()[:60])
    # Check paragraph-level page break before
    pPr = p._element.find(qn("w:pPr"))
    if pPr is not None:
        pbb = pPr.find(qn("w:pageBreakBefore"))
        if pbb is not None:
            page_breaks += 1
            if p.style.name.startswith("Heading"):
                page_break_headings.append(f"[before] {p.text.strip()[:60]}")

print(f"Page breaks: {page_breaks}")
if page_break_headings:
    print("Chapter/section page breaks:")
    for h in page_break_headings[:10]:
        print(f"  {h}")
print()

# --- SECTION PROPERTIES ---
print("## 9. SECTION PROPERTIES")
print("-" * 50)

sections = doc.sections
print(f"Sections: {len(sections)}")
for i, sec in enumerate(sections):
    w = sec.page_width.cm if sec.page_width else 0
    h = sec.page_height.cm if sec.page_height else 0
    top = sec.top_margin.cm if sec.top_margin else 0
    bot = sec.bottom_margin.cm if sec.bottom_margin else 0
    left = sec.left_margin.cm if sec.left_margin else 0
    right = sec.right_margin.cm if sec.right_margin else 0
    orient = sec.orientation
    print(f"  Section {i+1}: {w:.1f}x{h:.1f}cm, margins T={top:.1f} B={bot:.1f} L={left:.1f} R={right:.1f}")
    # Check for header/footer
    try:
        has_header = sec.header is not None and len(sec.header.paragraphs) > 0
        header_text = "".join(p.text for p in sec.header.paragraphs).strip() if has_header else ""
        has_footer = sec.footer is not None and len(sec.footer.paragraphs) > 0
        footer_text = "".join(p.text for p in sec.footer.paragraphs).strip() if has_footer else ""
        print(f"    Header: {'yes' if header_text else 'empty'}")
        print(f"    Footer: {'yes' if footer_text else 'empty'}")
    except:
        print(f"    Header/Footer: could not inspect")
print()

# --- STYLE ANALYSIS ---
print("## 10. STYLE USAGE")
print("-" * 50)

from collections import Counter
style_counter = Counter()
for p in doc.paragraphs:
    style_counter[p.style.name] += 1

print("Style distribution (top 15):")
for style, count in style_counter.most_common(15):
    print(f"  {style}: {count}")
print()

# --- BODY TEXT QUALITY ---
print("## 11. BODY TEXT QUALITY")
print("-" * 50)

empty_paras = 0
very_short_paras = 0
single_char_paras = 0
for p in doc.paragraphs:
    txt = p.text.strip()
    if not txt:
        empty_paras += 1
    elif len(txt) == 1:
        single_char_paras += 1
    elif len(txt) < 10 and p.style.name == "Normal":
        very_short_paras += 1

print(f"Empty paragraphs: {empty_paras}")
print(f"Single-char paragraphs: {single_char_paras}")
print(f"Very short (<10 chars) body paragraphs: {very_short_paras}")
print()

# --- SUMMARY ---
print("=" * 70)
print("SUMMARY")
print("=" * 70)
print(f"File size:    {file_size:.1f} KB")
print(f"Paragraphs:   {len(doc.paragraphs)}")
print(f"Tables:       {len(doc.tables)}")
print(f"Headings:     {len(headings)} (H1:{len(h1s)} H2:{len(h2s)} H3:{len(h3s)} H4:{len(h4s)})")
print(f"Footnotes:    {footnote_count}")
print(f"Hyperlinks:   {len(hyperlinks)} ({broken} broken)" if broken else f"Hyperlinks: {len(hyperlinks)} (all valid)")
print(f"Bookmarks:    {len(bookmarks)}")
print(f"Fields:       {len(instr_texts)}")
print(f"Page breaks:  {page_breaks}")
print(f"Empty paras:  {empty_paras}")
print()

# --- PRIORITY ISSUES ---
print("=" * 70)
print("PRIORITY ISSUES")
print("=" * 70)

issues = []
if broken > 0:
    issues.append(("HIGH", f"{broken} broken hyperlink(s) — anchor target has no bookmark"))
if empty_tables > 0:
    issues.append(("HIGH", f"{empty_tables} empty table(s)"))
if missing:
    issues.append(("MEDIUM", f"Missing captions: {missing}"))
if left_aligned_arabic > 5:
    issues.append(("MEDIUM", f"{left_aligned_arabic} left-aligned Arabic paragraphs"))
if single_char_paras > 3:
    issues.append(("LOW", f"{single_char_paras} single-character paragraphs (possible artifacts)"))

if not issues:
    issues.append(("OK", "No critical issues found"))

for severity, msg in issues:
    icon = {"HIGH": "[!]", "MEDIUM": "[~]", "LOW": "[.]", "OK": "[OK]"}[severity]
    print(f"  {icon} {severity}: {msg}")

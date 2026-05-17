from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_wrapper():
    doc = Document()

    # --- Styling ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- Title ---
    title = doc.add_heading('PROJECT ACADEMIX v13.2', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Final Project Completion & Executive Summary')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)

    doc.add_paragraph('\n') # Spacer

    # --- Executive Summary ---
    doc.add_heading('1. Executive Summary', level=1)
    summary_text = (
        "Project Academix v13.2 represents a comprehensive digital transformation of public sector "
        "inventory management. By implementing a Decision Support System (DSS) based on a hybrid "
        "ABC-Wilson optimization model, the project has successfully moved the administrative "
        "logistics of the Education Directorate of El Bayadh from a reactive, manual process to "
        "a proactive, data-driven operation. The system is designed with an 'offline-first' "
        "philosophy, ensuring 100% portability and zero-cost deployment in resource-constrained environments."
    )
    doc.add_paragraph(summary_text)

    # --- Core Deliverables Table ---
    doc.add_heading('2. Final Deliverables Matrix', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Table Headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Asset'
    hdr_cells[1].text = 'Verification Status'
    hdr_cells[2].text = 'Key Result'

    # Data Rows
    deliverables = [
        ('ERP Workbook (v13.2)', '174/174 PASS', '99.7% Operational Performance'),
        ('BTS Thesis', '36/36 PASS', 'Full CNEPD Compliance'),
        ('Research Paper', 'Peer-Review Ready', 'IMRaD Structure Complete'),
        ('VBA Source Code', 'Audit Verified', '0 Compile Errors / ACID Safe')
    ]

    for asset, status, result in deliverables:
        row_cells = table.add_row().cells
        row_cells[0].text = asset
        row_cells[1].text = status
        row_cells[2].text = result

    doc.add_paragraph('\n') # Spacer

    # --- Technical Highlights ---
    doc.add_heading('3. Technical Achievements', level=1)
    highlights = [
        "Integrated Wilson EOQ Model: Automated optimal order quantity (Q*) calculation.",
        "Dynamic ROP Engine: Eliminated stockout risks via real-time Reorder Point alerts.",
        "Pure VBA Architecture: Zero-dependency deployment ensuring permanent sustainability.",
        "Surgical Edit Pipeline: 137-point verification ensuring 'Golden' code stability."
    ]
    for item in highlights:
        doc.add_paragraph(item, style='List Bullet')

    # --- Final Verdict ---
    doc.add_paragraph('\n')
    verdict_box = doc.add_paragraph()
    verdict_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = verdict_box.add_run('FINAL VERDICT: GOLDEN / PRINT-READY')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 128, 0) # Green

    # --- Footer ---
    doc.add_paragraph('\n\n')
    footer = doc.add_paragraph('Date: May 17, 2026\nAuthor: Mahi Kamel Abdelghani')
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Save
    filename = "Academix_v13_2_Executive_Summary.docx"
    doc.save(filename)
    print(f"Successfully generated: {filename}")

if __name__ == "__main__":
    create_wrapper()

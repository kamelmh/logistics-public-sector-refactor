#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Academix v13.2 — Main DOCX Builder
Generates: output/Memoire_Academix_v13_2_POLISHED.docx
"""

import sys, os, re
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from build_docx import *
from build_front import *
from render_body import *
from parse_md import parse_markdown

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.abspath(os.path.join(SCRIPT_DIR, '..'))
THESIS_DIR = os.path.join(PROJECT_ROOT, 'Thesis_Surgical_Edit')
OUTPUT = os.path.join(THESIS_DIR, 'output', 'Memoire_Academix_v13_2_POLISHED.docx')
THESIS_MD = os.path.join(THESIS_DIR, 'Memoire_DSS_Logistique_ElBayadh.md')

# ── BLOCK SPLITTER ────────────────────────────────────────────────────────────
def extract_front_matter(blocks):
    """
    Split blocks into sections:
    returns dict with keys: cover, bismi, hadiya, shukr,
    abstract_ar, abstract_fr, abstract_en, keywords, abbrev, toc, figures, body
    """
    result = {
        'cover': [], 'bismi': [], 'hadiya': [], 'shukr': [],
        'abstract_ar': [], 'abstract_fr': [], 'abstract_en': [],
        'keywords': [], 'abbrev': [], 'toc': [], 'figures': [],
        'body': []
    }

    current = 'cover'
    cover_done = False

    for blk in blocks:
        t = blk['type']
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', blk.get('text', '')).strip()

        # Detect section boundaries
        if text == 'بسم الله الرحمن الرحيم':
            current = 'bismi'
            cover_done = True
            continue
        if text == 'إهداء':
            current = 'hadiya'
            continue
        if text == 'شكر وتقدير':
            current = 'shukr'
            continue
        if text == 'ملخص' and current in ('shukr', 'bismi', 'hadiya'):
            current = 'abstract_ar'
            continue
        if text == 'Résumé':
            current = 'abstract_fr'
            continue
        if text == 'Abstract':
            current = 'abstract_en'
            continue
        if text in ('Mots-clés / Keywords', 'Mots-clés / Keywords:'):
            current = 'keywords'
            continue
        if text == 'جدول المختصرات والرموز التقنية':
            current = 'abbrev'
            continue
        if text == 'فهرس المحتويات':
            current = 'toc'
            continue
        if text == 'قائمة الجداول والأشكال':
            current = 'figures'
            continue
        if text == 'المقدمة العامة' or t == 'chapter':
            current = 'body'

        if current == 'body' or t in ('chapter', 'mabhath', 'matlab',
                                       'awalan', 'conclusion', 'footnote',
                                       'separator', 'code'):
            if not cover_done and current != 'body':
                result['cover'].append(blk)
            else:
                result['body'].append(blk)
        else:
            result[current].append(blk)

    return result

def extract_abbrev_rows(blocks):
    """Extract abbreviation table rows from blocks"""
    rows = []
    for blk in blocks:
        if blk['type'] == 'table':
            for row in blk.get('rows', []):
                rows.append(row)
        elif blk['type'] == 'body':
            text = blk.get('text', '').strip()
            # Parse lines like: ERP | Enterprise Resource Planning | نظام تخطيط موارد المؤسسة
            if '|' in text:
                parts = [p.strip() for p in text.split('|')]
                if len(parts) >= 2:
                    rows.append(parts[:3])
    return rows

def extract_figures_rows(blocks):
    """Extract table of figures from blocks"""
    rows = []
    for blk in blocks:
        if blk['type'] == 'table':
            for row in blk.get('rows', []):
                rows.append(row)
    return rows

# ── CONFIGURE DOCUMENT RTL ────────────────────────────────────────────────────
def set_doc_defaults(doc):
    """Set document-level RTL and Arabic defaults"""
    settings = doc.settings.element
    # Add RTL theme
    docDefaults = settings.find('.//' + qn('w:docDefaults'))
    # Document body direction
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is None:
        sectPr = OxmlElement('w:sectPr')
        body.append(sectPr)
    # Set bidi on body sectPr
    bidi = sectPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        sectPr.insert(0, bidi)
    # Page setup
    pgSz = sectPr.find(qn('w:pgSz'))
    if pgSz is None:
        pgSz = OxmlElement('w:pgSz')
        sectPr.append(pgSz)
    pgSz.set(qn('w:w'), str(int(21 * 1440 / 2.54)))
    pgSz.set(qn('w:h'), str(int(29.7 * 1440 / 2.54)))
    pgSz.set(qn('w:orient'), 'portrait')

    pgMar = sectPr.find(qn('w:pgMar'))
    if pgMar is None:
        pgMar = OxmlElement('w:pgMar')
        sectPr.append(pgMar)
    m = str(int(3.5 * 1440 / 2.54))
    pgMar.set(qn('w:top'), m)
    pgMar.set(qn('w:right'), m)
    pgMar.set(qn('w:bottom'), m)
    pgMar.set(qn('w:left'), m)

    # Default font
    try:
        docDefaults = doc.element.find('.//' + qn('w:docDefaults'))
        if docDefaults is not None:
            rPrDefault = docDefaults.find('.//' + qn('w:rPrDefault'))
            if rPrDefault is None:
                rPrDefault = OxmlElement('w:rPrDefault')
                docDefaults.append(rPrDefault)
            rPr = rPrDefault.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                rPrDefault.append(rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:ascii'), FONT_AR)
            rFonts.set(qn('w:hAnsi'), FONT_AR)
            rFonts.set(qn('w:cs'), FONT_AR)
            sz = rPr.find(qn('w:sz'))
            if sz is None:
                sz = OxmlElement('w:sz')
                rPr.append(sz)
            sz.set(qn('w:val'), '28')  # 14pt = 28 half-points
    except Exception:
        pass

def add_section_properties(doc, fmt='arabicAbjad', start=1, show_footer=True):
    """
    Insert a section break paragraph that starts a new section
    with specified page number format.
    """
    # Add a paragraph that carries the sectPr for the PREVIOUS section
    p = doc.add_paragraph()
    p._p.clear()
    pPr = OxmlElement('w:pPr')

    sectPr = OxmlElement('w:sectPr')

    # Page size
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), str(int(21 * 1440 / 2.54)))
    pgSz.set(qn('w:h'), str(int(29.7 * 1440 / 2.54)))
    sectPr.append(pgSz)

    # Margins
    pgMar = OxmlElement('w:pgMar')
    m = str(int(3.5 * 1440 / 2.54))
    pgMar.set(qn('w:top'), m)
    pgMar.set(qn('w:right'), m)
    pgMar.set(qn('w:bottom'), m)
    pgMar.set(qn('w:left'), m)
    pgMar.set(qn('w:footer'), '720')
    sectPr.append(pgMar)

    # RTL
    bidi = OxmlElement('w:bidi')
    sectPr.append(bidi)

    # Page number format
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), fmt)
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)

    # Footer reference (simple page number)
    if show_footer:
        ftrRef = OxmlElement('w:footerReference')
        ftrRef.set(qn('w:type'), 'default')
        ftrRef.set(qn('w:id'), 'rId99')  # placeholder
        # Skip this - will add footer manually

    pPr.append(sectPr)
    p._p.append(pPr)
    return p

# ── REFERENCES SECTION ────────────────────────────────────────────────────────
def render_references(doc, blocks):
    """Render references in proper format"""
    heading1(doc, 'قائمة المصادر والمراجع')
    sub_title = None
    for blk in blocks:
        text = clean_line(blk.get('text', ''))
        t = blk['type']
        if t == 'blank':
            continue
        if t == 'chapter':
            heading2(doc, text)
        elif t == 'body':
            # Check if it's a subsection title
            clean = re.sub(r'\*\*(.*?)\*\*', r'\1', text).strip()
            if clean.startswith('القسم ') or clean.startswith('## '):
                heading2(doc, clean.replace('## ', ''))
            elif re.match(r'^\[\d+\]', clean):
                # Reference entry
                p = doc.add_paragraph()
                set_para_rtl(p, 'right')
                set_line_spacing(p, 1.5)
                set_space_before_after(p, 40, 40)
                fn = FONT_EN if is_english(clean) else FONT_AR
                add_run(p, clean, size=12, font=fn)
            else:
                body_para(doc, text, size=12, align='right')

# ── GLOSSARY ──────────────────────────────────────────────────────────────────
def render_glossary(doc, blocks):
    """Render glossary (قائمة المصطلحات)"""
    heading1(doc, 'قائمة المصطلحات')
    for blk in blocks:
        if blk['type'] == 'table':
            headers = blk.get('headers', [])
            rows = blk.get('rows', [])
            if headers:
                # 4-column glossary table
                col_w = [3.0, 3.0, 3.0, 5.0]
                build_table(doc, headers, rows, col_w)
        elif blk['type'] not in ('blank', 'separator'):
            text = clean_line(blk.get('text', ''))
            if text:
                body_para(doc, text, size=12)

# ── APPENDICES ────────────────────────────────────────────────────────────────
def render_appendices(doc, blocks):
    """Render appendices section"""
    heading1(doc, 'الملاحق')
    for blk in blocks:
        t = blk['type']
        text = clean_line(blk.get('text', ''))
        if t == 'blank':
            continue
        elif t == 'table':
            headers = blk.get('headers', [])
            rows = blk.get('rows', [])
            if headers:
                n = len(headers)
                col_w = [14.0/n] * n
                try:
                    build_table(doc, headers, rows, col_w)
                except:
                    pass
        elif t == 'chapter':
            heading2(doc, text)
        elif t == 'body':
            clean = re.sub(r'\*\*(.*?)\*\*', r'\1', text).strip()
            if clean.startswith('##'):
                heading2(doc, clean.replace('##', '').strip())
            elif clean.startswith('#'):
                heading3(doc, clean.replace('#', '').strip())
            elif re.match(r'^### ', text) or re.match(r'^#### ', text):
                heading3(doc, clean)
            else:
                fn = FONT_EN if is_english(clean) else FONT_AR
                body_para(doc, text, size=12, font=fn)
        elif t == 'code':
            p = doc.add_paragraph()
            set_space_before_after(p, 60, 60)
            run = p.add_run(blk.get('text', ''))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

# ── BODY CHAPTER RENDERER ─────────────────────────────────────────────────────
def render_body(doc, body_blocks):
    """Render all main body blocks"""
    footnote_blocks = []
    refs_started = False
    glossary_started = False
    appendix_started = False
    refs_blocks = []
    glossary_blocks = []
    appendix_blocks = []

    i = 0
    while i < len(body_blocks):
        blk = body_blocks[i]
        t = blk['type']
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', blk.get('text', '')).strip()

        # Detect major section boundaries
        if text in ('قائمة المصادر والمراجع', '# قائمة المصادر والمراجع'):
            refs_started = True
            glossary_started = False
            appendix_started = False
            i += 1
            continue
        if refs_started and text in ('قائمة المصطلحات', 'قائمة المصطلحات (Glossaire / قائمة المصطلحات)'):
            glossary_started = True
            refs_started = False
            i += 1
            continue
        if (refs_started or glossary_started) and text in ('الملاحق',):
            appendix_started = True
            refs_started = False
            glossary_started = False
            i += 1
            continue

        if refs_started:
            refs_blocks.append(blk)
        elif glossary_started:
            glossary_blocks.append(blk)
        elif appendix_started:
            appendix_blocks.append(blk)
        elif t == 'footnote':
            footnote_blocks.append(blk)
        else:
            try:
                render_body_block(doc, blk)
            except Exception as e:
                print(f"  Warning block {i}: {e}")
        i += 1

    # Render special end sections
    if refs_blocks:
        add_page_break(doc)
        render_references(doc, refs_blocks)
    if glossary_blocks:
        add_page_break(doc)
        render_glossary(doc, glossary_blocks)
    if appendix_blocks:
        add_page_break(doc)
        render_appendices(doc, appendix_blocks)

    # Footnotes section at end
    if footnote_blocks:
        add_page_break(doc)
        p = doc.add_paragraph()
        set_para_rtl(p, 'right')
        add_run(p, 'الهوامش والمراجع', size=14, bold=True, color=C_BLUE)
        for fb in footnote_blocks:
            render_body_block(doc, fb)

# ── MAIN BUILD ────────────────────────────────────────────────────────────────
def main():
    print("═" * 60)
    print("Academix v13.2 — DOCX Builder")
    print("═" * 60)

    # 1. Parse markdown
    print("\n[1] Parsing thesis markdown...")
    blocks = parse_markdown(THESIS_MD)
    print(f"    → {len(blocks)} blocks parsed")

    # 2. Split into sections
    print("[2] Splitting sections...")
    sections = extract_front_matter(blocks)
    print(f"    → body blocks: {len(sections['body'])}")
    print(f"    → hadiya blocks: {len(sections['hadiya'])}")
    print(f"    → shukr blocks: {len(sections['shukr'])}")
    print(f"    → abstract_ar: {len(sections['abstract_ar'])}")

    # 3. Create document
    print("[3] Creating document...")
    doc = Document()
    # Remove default empty paragraph
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # Configure document defaults
    set_doc_defaults(doc)

    # ── SECTION 1: COVER (no page number) ────────────────────────────────
    print("[4] Building cover page...")
    build_cover(doc)

    # ── SECTION BREAK → Front matter (arabicAbjad page numbers) ─────────
    # Insert section break paragraph with arabicAbjad numbering
    sb1 = add_section_properties(doc, fmt='none', start=1, show_footer=False)

    # ── بسم الله ─────────────────────────────────────────────────────────
    print("[5] Building بسم الله...")
    build_bismi(doc)

    # ── إهداء (page أ) ───────────────────────────────────────────────────
    print("[6] Building إهداء...")
    if sections['hadiya']:
        build_hadiya(doc, sections['hadiya'])
    else:
        add_page_break(doc)

    # ── شكر وتقدير (page ب) ──────────────────────────────────────────────
    print("[7] Building شكر وتقدير...")
    if sections['shukr']:
        build_shukr(doc, sections['shukr'])
    else:
        add_page_break(doc)

    # ── ملخص (page ج) ────────────────────────────────────────────────────
    print("[8] Building abstracts...")
    if sections['abstract_ar']:
        abstract_section(doc, 'ملخص', sections['abstract_ar'], lang='ar')
    # ── Résumé (page د) ───────────────────────────────────────────────────
    if sections['abstract_fr']:
        abstract_section(doc, 'Résumé', sections['abstract_fr'], lang='fr')
    # ── Abstract (page هـ) ────────────────────────────────────────────────
    if sections['abstract_en']:
        abstract_section(doc, 'Abstract', sections['abstract_en'], lang='en')

    # ── جدول المختصرات ────────────────────────────────────────────────────
    print("[9] Building abbreviations table...")
    abbrev_rows = extract_abbrev_rows(sections['abbrev'])
    if abbrev_rows:
        build_abbrev_table(doc, abbrev_rows)
    else:
        # Try keywords section
        abbrev_rows = extract_abbrev_rows(sections['keywords'])
        if abbrev_rows:
            build_abbrev_table(doc, abbrev_rows)

    # ── فهرس المحتويات ────────────────────────────────────────────────────
    print("[10] Building TOC placeholder...")
    build_toc_placeholder(doc)

    # ── قائمة الجداول ─────────────────────────────────────────────────────
    print("[11] Building figures list...")
    fig_rows = extract_figures_rows(sections['figures'])
    if fig_rows:
        build_figures_list(doc, fig_rows)

    # ── SECTION BREAK → Main body (decimal page numbers from 1) ──────────
    sb2 = add_section_properties(doc, fmt='decimal', start=1, show_footer=True)

    # ── MAIN BODY ─────────────────────────────────────────────────────────
    print("[12] Rendering main body...")
    render_body(doc, sections['body'])

    # ── SAVE ──────────────────────────────────────────────────────────────
    print(f"\n[13] Saving to {OUTPUT}...")
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    size_kb = os.path.getsize(OUTPUT) / 1024
    print(f"    ✓ Saved: {size_kb:.1f} KB")
    print("\n═" * 60)
    print("Done! Memoire_Academix_v13_2_POLISHED.docx generated.")

if __name__ == '__main__':
    main()

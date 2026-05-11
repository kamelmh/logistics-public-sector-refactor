#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Body content renderer for Academix v13.2
Converts parsed markdown blocks into python-docx elements.
"""

import re, sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from build_docx import *
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_EN = "Times New Roman"

def clean_line(text):
    """Clean markdown artifacts from a line"""
    # Remove trailing backslash
    text = text.rstrip('\\').strip()
    # Remove \\ markdown escape
    text = text.replace('\\\\', '')
    # Fix --- to —
    text = text.replace('---', '—')
    text = text.replace('--', '—')
    # Fix \- to -
    text = text.replace('\\-', '-')
    # Remove extra asterisks
    return text

def is_english(text):
    """Detect if text is primarily Latin/English"""
    ascii_count = sum(1 for c in text if c.isascii() and c.isalpha())
    total = max(len(text.replace(' ', '')), 1)
    return ascii_count / total > 0.5

def guess_font(text):
    return FONT_EN if is_english(text) else FONT_AR

def render_body_block(doc, blk):
    """Render a single body block into the document"""
    t = blk['type']
    raw = blk.get('text', '')
    text = clean_line(raw)

    if not text and t not in ('table', 'code'):
        return

    # ── CHAPTER heading ──────────────────────────────────────────────────
    if t == 'chapter':
        if text in ('قائمة المصادر والمراجع', 'قائمة المصطلحات', 'الملاحق'):
            heading1(doc, text)
        else:
            add_page_break(doc)
            heading1(doc, text)

    # ── MABHATH ──────────────────────────────────────────────────────────
    elif t == 'mabhath':
        heading2(doc, text)

    # ── MATLAB ───────────────────────────────────────────────────────────
    elif t == 'matlab':
        if text == 'تمهيد':
            heading3(doc, text)
        else:
            heading3(doc, text)

    # ── AWALAN ───────────────────────────────────────────────────────────
    elif t == 'awalan':
        heading4(doc, text)

    # ── CONCLUSION ───────────────────────────────────────────────────────
    elif t == 'conclusion':
        p = doc.add_paragraph()
        set_para_rtl(p, 'right')
        set_line_spacing(p, 1.3)
        set_space_before_after(p, 160, 80)
        add_run(p, text, size=15, bold=True, color=C_BLUE)

    # ── TABLE ─────────────────────────────────────────────────────────────
    elif t == 'table':
        headers = blk.get('headers', [])
        rows = blk.get('rows', [])
        # Estimate column widths
        n = max(len(headers), 1)
        total_cm = 14.0  # A4 - 2*3.5 margins
        col_w = [total_cm / n] * n
        # For tables with many columns, narrow them
        if n >= 5:
            col_w = [total_cm / n] * n
        elif n == 3:
            col_w = [3.0, 6.0, 5.0]
        elif n == 4:
            col_w = [2.5, 4.5, 3.5, 3.5]
        elif n == 7:
            col_w = [2.0, 3.0, 1.5, 1.5, 1.5, 1.5, 3.0]
        elif n == 8:
            col_w = [1.5, 2.5, 1.5, 1.5, 1.5, 1.5, 1.5, 2.0]
        try:
            tbl = build_table(doc, headers, rows, col_w)
        except Exception as e:
            # Fallback: simple table
            p = doc.add_paragraph()
            add_run(p, f'[جدول: {headers}]', size=11, color=C_GRAY)
        spacer(doc, 4)

    # ── CODE BLOCK ────────────────────────────────────────────────────────
    elif t == 'code':
        p = doc.add_paragraph()
        # Code in monospace, LTR
        pPr = p._p.get_or_add_pPr()
        set_space_before_after(p, 80, 80)
        code_text = blk.get('text', '')
        if code_text.strip():
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

    # ── BODY PARAGRAPH ───────────────────────────────────────────────────
    elif t == 'body':
        # Detect special inline headings
        stripped_no_bold = re.sub(r'\*\*(.*?)\*\*', r'\1', text)

        # Skip pure cover page lines that already rendered
        if _is_cover_line(stripped_no_bold):
            return

        # Detect section title lines (short, bold, no period)
        if is_section_title(text):
            p = doc.add_paragraph()
            set_para_rtl(p, 'center')
            set_line_spacing(p, 1.5)
            set_space_before_after(p, 120, 80)
            clean = re.sub(r'\*\*(.*?)\*\*', r'\1', text).strip()
            add_run(p, clean, size=16, bold=True)
            return

        # Normal body paragraph
        fn = guess_font(text)
        # Handle italic *...* (French/English)
        text2 = re.sub(r'\*(.*?)\*', r'\1', text)  # strip italic markers
        body_para(doc, text2, size=14, align='both', font=fn)

    # ── FOOTNOTE ─────────────────────────────────────────────────────────
    elif t == 'footnote':
        # Render as small text at bottom section
        p = doc.add_paragraph()
        set_para_rtl(p, 'right')
        set_space_before_after(p, 20, 20)
        num = blk.get('num', '?')
        fn_text = blk.get('text', '')
        fn_text = clean_line(fn_text)
        add_run(p, f'[{num}] ', size=9, bold=True)
        add_run(p, fn_text, size=9,
                font=FONT_EN if is_english(fn_text) else FONT_AR)

    # ── SEPARATOR ─────────────────────────────────────────────────────────
    elif t == 'separator':
        # Many are just \ in the cover; skip or add small space
        pass

    # ── BLANK ─────────────────────────────────────────────────────────────
    elif t == 'blank':
        pass  # handled by caller


COVER_LINES = {
    'الجمهورية الجزائرية الديمقراطية الشعبية',
    'وزارة التكوين والتعليم المهنيين',
    'المعهد الوطني المتخصص في التكوين المهني',
    'بن سعيدي عبد العاطي — البيض',
    'بن سعيدي عبد العاطي --- البيض',
    'مذكرة تخرج',
    'لنيل شهادة تقني سامي في تسيير المخزونات واللوجستيك بعنوان',
    'بناء نظام دعم قرار لإتمام دورة الشراء والتخزين',
    'دراسة حالة في ظل محدودية الموارد وانعدام الاتصال بالإنترنت',
    'دراسة حالة: مديرية التربية -- البيض',
    'دراسة حالة: مديرية التربية -- البيض',
    'Designing a Decision Support System for Integrating Purchasing and Warehousing: A Case',
    'Study of El Bayadh Directorate of Education',
    'إعداد الطالب: ماحي كمال عبد الغني',
    'إشراف الأستاذة: دهيني ميمونة',
    'دورة أكتوبر 2025',
}

def _is_cover_line(text):
    return text.strip() in COVER_LINES or any(text.strip().startswith(cl) for cl in COVER_LINES)


SECTION_TITLES = {
    'بسم الله الرحمن الرحيم', 'إهداء', 'شكر وتقدير', 'ملخص',
    'Résumé', 'Abstract', 'Mots-clés / Keywords',
    'جدول المختصرات والرموز التقنية', 'فهرس المحتويات',
    'قائمة الجداول والأشكال', 'قائمة المصادر والمراجع',
    'قائمة المصطلحات', 'الملاحق',
    'القسم الأول: المراجع الوطنية والمؤسسية',
    'القسم الثاني: المراجع الأجنبية والكتب المتخصصة',
    'القسم الثالث: المراجع الرقمية والمواقع الإلكترونية',
}

def is_section_title(text):
    clean = re.sub(r'\*\*(.*?)\*\*', r'\1', text).strip()
    # Bold standalone on its line → possible section title
    is_bold_only = text.strip().startswith('**') and text.strip().endswith('**')
    in_titles = clean in SECTION_TITLES
    # Short + bold + no sentence-ending period → section title candidate
    is_short_bold = (is_bold_only and len(clean) < 60 and
                     not clean.endswith('.') and not clean.endswith(':') and
                     '\n' not in clean)
    return in_titles or (is_short_bold and len(clean) > 3)


print("Body renderer module loaded.")

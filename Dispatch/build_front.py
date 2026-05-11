#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Cover page + front matter builder for Academix v13.2
"""

import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from build_docx import *
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_footer_pagenum(section, fmt='arabicAbjad', start=1, show=True):
    """Add page number footer to a section"""
    footer = section.footer
    footer.is_linked_to_previous = False
    # Clear existing
    for p in footer.paragraphs:
        p.clear()
    if not show:
        return
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    set_para_rtl(fp, 'center')
    # Page number field
    run = fp.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.size = Pt(10)
    run.font.name = FONT_AR

def set_section_pagenum(sectPr, fmt='arabicAbjad', start=1):
    """Configure page number format in sectPr"""
    existing = sectPr.find(qn('w:pgNumType'))
    if existing is not None:
        sectPr.remove(existing)
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), fmt)
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)

def get_or_add_sectPr(doc_body):
    """Get or create the body-level sectPr"""
    sectPr = doc_body.find(qn('w:sectPr'))
    if sectPr is None:
        sectPr = OxmlElement('w:sectPr')
        doc_body.append(sectPr)
    return sectPr

# ── COVER PAGE ────────────────────────────────────────────────────────────────
def build_cover(doc):
    """Build the cover page — recreating cover-page.docx design"""

    def cov(text, size=14, bold=False, color=None, align='center', font=None,
            before=60, after=60):
        p = doc.add_paragraph()
        set_para_rtl(p, align)
        set_line_spacing(p, 1.3)
        set_space_before_after(p, before, after)
        add_run(p, text, size=size, bold=bold, color=color, font=font or FONT_AR)
        return p

    # Top: republic header
    cov('الجمهورية الجزائرية الديمقراطية الشعبية',
        size=16, bold=True, before=120)

    cov('وزارة التكوين والتعليم المهنيين', size=13, before=20, after=20)

    cov('المعهد الوطني المتخصص في التكوين المهني',
        size=14, bold=True, before=20, after=20)

    cov('بن سعيدي عبد العاطي — البيض',
        size=14, bold=True, before=10, after=60)

    spacer(doc, 18)

    cov('مذكرة تخرج', size=18, bold=True, before=40, after=20)

    cov('لنيل شهادة تقني سامي في تسيير المخزونات واللوجستيك',
        size=12, before=10, after=40)

    spacer(doc, 12)

    cov('بناء نظام دعم قرار لإتمام دورة الشراء والتخزين',
        size=16, bold=True, before=60, after=30)

    cov('دراسة حالة: مديرية التربية – البيض -',
        size=18, bold=True, color=C_GRAY, before=20, after=30)

    spacer(doc, 10)

    # English title
    cov('Designing a Decision Support System for Integrating Purchasing and Warehousing:\n'
        'A Case Study of El Bayadh Directorate of Education',
        size=12, bold=True, color=C_GREEN, font=FONT_EN, before=20, after=40)

    spacer(doc, 20)

    # Author + supervisor line (RTL tab layout)
    p_auth = doc.add_paragraph()
    set_para_rtl(p_auth, 'both')
    set_line_spacing(p_auth, 1.3)
    set_space_before_after(p_auth, 80, 40)
    # Use tab to separate the two parts
    r1 = add_run(p_auth, 'إعداد الطالب: ماحي كمال عبد الغني', size=12, bold=True)
    r_tab = p_auth.add_run('\t')
    r_tab.font.size = Pt(12)
    r2 = add_run(p_auth, 'إشراف الأستاذة: دهيني ميمونة', size=12, bold=True)

    spacer(doc, 20)

    cov('دورة أكتوبر 2025', size=16, bold=True,
        color=C_GOLD, before=60, after=120)

    # Page break after cover (no page number on cover)
    add_page_break(doc)

# ── FRONT MATTER SECTION BUILDER ──────────────────────────────────────────────
def build_bismi(doc):
    spacer(doc, 36)
    p = doc.add_paragraph()
    set_para_rtl(p, 'center')
    set_line_spacing(p, 2.0)
    set_space_before_after(p, 240, 240)
    add_run(p, 'بسم الله الرحمن الرحيم', size=20, bold=True)
    add_page_break(doc)

def front_section_title(doc, text, size=18):
    p = doc.add_paragraph()
    set_para_rtl(p, 'center')
    set_line_spacing(p, 1.5)
    set_space_before_after(p, 120, 120)
    add_run(p, text, size=size, bold=True)
    return p

def build_hadiya(doc, blocks):
    """إهداء page"""
    front_section_title(doc, 'إهداء')
    for blk in blocks:
        if blk['type'] not in ('blank', 'separator'):
            body_para(doc, blk['text'], size=14, align='both')
    add_page_break(doc)

def build_shukr(doc, blocks):
    """شكر وتقدير page"""
    front_section_title(doc, 'شكر وتقدير')
    for blk in blocks:
        if blk['type'] not in ('blank', 'separator'):
            body_para(doc, blk['text'], size=14, align='both')
    add_page_break(doc)

def abstract_section(doc, title, blocks, lang='ar'):
    """Build one abstract section"""
    p = doc.add_paragraph()
    align = 'center'
    set_para_rtl(p, align)
    set_line_spacing(p, 1.3)
    set_space_before_after(p, 120, 80)
    if lang == 'ar':
        add_run(p, title, size=18, bold=True)
    else:
        add_run(p, title, size=18, bold=True, font=FONT_EN)
    for blk in blocks:
        if blk['type'] not in ('blank', 'separator'):
            fn = FONT_EN if lang in ('fr', 'en') else FONT_AR
            body_para(doc, blk['text'], size=13, align='both', font=fn)
    add_page_break(doc)

def build_abbrev_table(doc, abbrev_rows):
    """Build abbreviation table (3-column)"""
    front_section_title(doc, 'جدول المختصرات والرموز التقنية')
    if abbrev_rows:
        headers = ['الاختصار', 'الاسم الكامل', 'المعنى']
        build_table(doc, headers, abbrev_rows)
    add_page_break(doc)

def build_toc_placeholder(doc):
    """فهرس المحتويات - simple placeholder for Word's TOC"""
    front_section_title(doc, 'فهرس المحتويات')
    p = doc.add_paragraph()
    set_para_rtl(p, 'center')
    set_line_spacing(p, 1.5)
    set_space_before_after(p, 80, 80)
    add_run(p, '[ يُضاف الفهرس التلقائي هنا في Microsoft Word ]',
            size=11, color=RGBColor(0x88, 0x88, 0x88))
    add_page_break(doc)

def build_figures_list(doc, figures):
    """قائمة الجداول والأشكال"""
    front_section_title(doc, 'قائمة الجداول والأشكال')
    if figures:
        headers = ['الرقم', 'عنوان الجدول', 'الموقع']
        build_table(doc, headers, figures)
    add_page_break(doc)

print("Cover/front matter module loaded.")

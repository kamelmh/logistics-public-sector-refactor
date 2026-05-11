#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Academix v13.2 — DOCX Builder
Generates Memoire_Academix_v13_2_POLISHED.docx
"""

import os, re, copy
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# ── COLORS ──────────────────────────────────────────────────────────────────
C_DARK  = RGBColor(0x1B, 0x26, 0x31)   # chapter headings
C_BLUE  = RGBColor(0x0C, 0x44, 0x7C)   # mabhath/matlab, table headers
C_GREEN = RGBColor(0x1F, 0x6B, 0x2E)   # English cover title
C_GOLD  = RGBColor(0x80, 0x60, 0x00)   # date
C_GRAY  = RGBColor(0x55, 0x55, 0x55)   # subtitle
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK = RGBColor(0x00, 0x00, 0x00)
HEX_HEADER = "0C447C"
HEX_ALT    = "EBF5FB"
HEX_WHITE  = "FFFFFF"

FONT_AR = "Traditional Arabic"
FONT_EN = "Times New Roman"

MARGIN = Cm(3.5)  # 3.5 cm margins

# ── RTL / BIDI HELPERS ───────────────────────────────────────────────────────
def make_rtl_elem():
    e = OxmlElement('w:bidi')
    return e

def set_para_rtl(para, align='right'):
    pPr = para._p.get_or_add_pPr()
    # bidi
    bidi_list = pPr.findall(qn('w:bidi'))
    if not bidi_list:
        bidi = OxmlElement('w:bidi')
        pPr.insert(0, bidi)
    # justification
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), align)

def set_run_rtl(run):
    rPr = run._r.get_or_add_rPr()
    rtl = rPr.find(qn('w:rtl'))
    if rtl is None:
        rtl = OxmlElement('w:rtl')
        rPr.append(rtl)

def set_run_fonts(run, font_name):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)

def set_line_spacing(para, factor=1.5):
    pPr = para._p.get_or_add_pPr()
    pPr_spacing = pPr.find(qn('w:spacing'))
    if pPr_spacing is None:
        pPr_spacing = OxmlElement('w:spacing')
        pPr.append(pPr_spacing)
    # 1.5 line = 360 twips (240 = single)
    pPr_spacing.set(qn('w:line'), str(int(240 * factor)))
    pPr_spacing.set(qn('w:lineRule'), 'auto')

def set_space_before_after(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        pPr.append(sp)
    if before is not None:
        sp.set(qn('w:before'), str(before))
    if after is not None:
        sp.set(qn('w:after'), str(after))

def add_run(para, text, size=14, bold=False, italic=False,
            color=None, font=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    fn = font or FONT_AR
    run.font.name = fn
    run.font.size = Pt(size)
    set_run_fonts(run, fn)
    if color:
        run.font.color.rgb = color
    set_run_rtl(run)
    return run

# ── DOCUMENT CREATION ────────────────────────────────────────────────────────
def new_document():
    doc = Document()
    # Remove default empty paragraph
    doc.element.body.clear()
    return doc

def configure_section(section):
    """Apply A4 + 3.5cm margin + RTL to a section"""
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = MARGIN
    section.right_margin= MARGIN
    section.top_margin  = MARGIN
    section.bottom_margin = MARGIN
    # Set RTL layout for the section
    sectPr = section._sectPr
    bidi = sectPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        sectPr.append(bidi)

def add_section_break(doc, break_type='nextPage'):
    """Add a section break paragraph"""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), str(int(21 * 1440 / 2.54)))   # A4 width in twips
    pgSz.set(qn('w:h'), str(int(29.7 * 1440 / 2.54)))  # A4 height in twips
    pgSz.set(qn('w:orient'), 'portrait')
    sectPr.append(pgSz)
    pgMar = OxmlElement('w:pgMar')
    m = str(int(3.5 * 1440 / 2.54))
    pgMar.set(qn('w:top'), m)
    pgMar.set(qn('w:right'), m)
    pgMar.set(qn('w:bottom'), m)
    pgMar.set(qn('w:left'), m)
    sectPr.append(pgMar)
    bidi = OxmlElement('w:bidi')
    sectPr.append(bidi)
    pPr.append(sectPr)
    return para

# ── TABLE HELPERS ─────────────────────────────────────────────────────────────
def hex_shading(cell, hex_color, hex_fill="auto"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)

def set_cell_border(cell, color="0C447C"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def cell_para(cell, text, size=12, bold=False, color=None, align='right', font=None):
    """Format cell content with RTL"""
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    para.clear()
    set_para_rtl(para, align)
    set_line_spacing(para, 1.15)
    set_space_before_after(para, 40, 40)
    run = add_run(para, text, size=size, bold=bold, color=color, font=font)
    return para

def build_table(doc, headers, rows, col_widths_cm=None):
    """Build a styled table with header row and alternating body rows"""
    n_cols = len(headers)
    table = doc.add_table(rows=1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Calculate widths
    avail_cm = 21 - 2 * 3.5   # 14 cm content width
    if col_widths_cm:
        widths = [Cm(w) for w in col_widths_cm]
    else:
        w = avail_cm / n_cols
        widths = [Cm(w)] * n_cols

    # Set table width
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    total_dxa = sum(int(w.twips) for w in widths)
    tblW.set(qn('w:w'), str(total_dxa))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # Header row
    hrow = table.rows[0]
    for i, hdr in enumerate(headers):
        if i >= len(hrow.cells):
            break
        cell = hrow.cells[i]
        hex_shading(cell, HEX_HEADER)
        set_cell_border(cell)
        # Set column width
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        wi = widths[i] if i < len(widths) else widths[-1]
        tcW.set(qn('w:w'), str(int(wi.twips)))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
        cell_para(cell, hdr, size=12, bold=True, color=C_WHITE)

    # Body rows
    for r_idx, row_data in enumerate(rows):
        row = table.add_row()
        fill = HEX_ALT if r_idx % 2 == 0 else "FFFFFF"
        for i, cell_text in enumerate(row_data):
            if i >= len(row.cells):
                break
            cell = row.cells[i]
            hex_shading(cell, fill)
            set_cell_border(cell)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            wi = widths[i] if i < len(widths) else widths[-1]
            tcW.set(qn('w:w'), str(int(wi.twips)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            cell_para(cell, str(cell_text), size=11)

    return table

# ── PAGE BREAK ────────────────────────────────────────────────────────────────
def page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_break_type())
    return para

def docx_break_type():
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.text import WD_BREAK
    return WD_BREAK.PAGE

def add_page_break(doc):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pb = OxmlElement('w:pageBreakBefore')
    pPr.append(pb)
    set_para_rtl(para)
    return para

# ── PARAGRAPH HELPERS ─────────────────────────────────────────────────────────
def body_para(doc, text, size=14, bold=False, color=None,
              align='both', spacing=1.5, before=80, after=80,
              font=None, indent_cm=None):
    """Standard Arabic body paragraph"""
    para = doc.add_paragraph()
    set_para_rtl(para, align)
    set_line_spacing(para, spacing)
    set_space_before_after(para, before, after)
    if indent_cm:
        pPr = para._p.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:right'), str(int(Cm(indent_cm).twips)))
        pPr.append(ind)
    # Parse inline bold: **text**
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        is_bold = (idx % 2 == 1) or bold
        add_run(para, part, size=size, bold=is_bold,
                color=color, font=font or FONT_AR)
    return para

def heading1(doc, text):
    """Chapter heading — Traditional Arabic 22pt, dark, bold, centered"""
    para = doc.add_paragraph()
    set_para_rtl(para, 'center')
    set_line_spacing(para, 1.3)
    set_space_before_after(para, 240, 120)
    add_run(para, text, size=22, bold=True, color=C_DARK)
    return para

def heading2(doc, text):
    """مبحث — Traditional Arabic 18pt, blue, bold"""
    para = doc.add_paragraph()
    set_para_rtl(para, 'right')
    set_line_spacing(para, 1.3)
    set_space_before_after(para, 200, 80)
    add_run(para, text, size=18, bold=True, color=C_BLUE)
    return para

def heading3(doc, text):
    """مطلب — Traditional Arabic 16pt, blue, bold"""
    para = doc.add_paragraph()
    set_para_rtl(para, 'right')
    set_line_spacing(para, 1.3)
    set_space_before_after(para, 160, 60)
    add_run(para, text, size=16, bold=True, color=C_BLUE)
    return para

def heading4(doc, text):
    """أولاً/ثانياً — Traditional Arabic 14pt, bold, underlined"""
    para = doc.add_paragraph()
    set_para_rtl(para, 'right')
    set_line_spacing(para, 1.5)
    set_space_before_after(para, 120, 40)
    run = add_run(para, text, size=14, bold=True)
    run.underline = True
    return para

def spacer(doc, size=6):
    para = doc.add_paragraph()
    set_para_rtl(para)
    set_space_before_after(para, 0, 0)
    run = para.add_run('')
    run.font.size = Pt(size)
    return para

print("Helper module loaded.")
